# %%
# =============================================================================
# PARTE 1: CONFIGURAÇÕES GERAIS E BIBLIOTECAS
# =============================================================================
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import matplotlib.ticker as ticker
import os
from datetime import datetime
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert
from IPython.display import display

# --- CONFIGURAÇÕES GERAIS (MODIFIQUE AQUI) ---

# 1. Nome do Fundo para Análise
# ! ATENÇÃO: O nome deve ser exatamente igual ao que está no arquivo de dados.
FUNDO_ANALISADO = 'FIDC FCT II'

# 2. Caminhos dos Arquivos de Entrada
CAMINHO_RENTABILIDADE = r"C:\Users\Leo\Desktop\Porto_Real\portoauto\src\vortx_estoques\data\135972-Rentabilidade_Sintetica.csv"
TEMPLATE_DOCX = r"C:\Users\Leo\Desktop\Porto_Real\portoreal\notebooks\template.docx"
LOGO_EMPRESA = r"C:\Users\Leo\Desktop\Porto_Real\portoreal\images\logo.png"

# 3. Caminhos dos Arquivos de Saída
PASTA_SAIDA_IMAGENS = r"C:\Users\Leo\Desktop\Porto_Real\portoauto\src\vortx_estoques\output\lamina_imagens"
PASTA_SAIDA_RELATORIOS = r"C:\Users\Leo\Desktop\Porto_Real\portoreal\notebooks"
# ----------------------------------------------------

# Cria as pastas de saída se não existirem
os.makedirs(PASTA_SAIDA_IMAGENS, exist_ok=True)
os.makedirs(PASTA_SAIDA_RELATORIOS, exist_ok=True)

# Paleta de cores para os gráficos
color1 = '#4BC1D7'  # Azul claro
color2 = '#286D82'  # Azul escuro
color3 = '##00B1DE' # Azul ciano
color4 = '##445D6D' # Cinza azulado
colorA = '##e2eef0' # Cor de fundo para cabeçalhos de tabela
colorB = '##107082' # Cor de texto para cabeçalhos de tabela

print(f"Configurações carregadas para o fundo: {FUNDO_ANALISADO}")
print(f"Imagens serão salvas em: {PASTA_SAIDA_IMAGENS}")
print(f"Relatórios finais serão salvos em: {PASTA_SAIDA_RELATORIOS}")


# %%
# =============================================================================
# PARTE 2: ANÁLISE DE DADOS E GERAÇÃO DAS IMAGENS
# =============================================================================

# --- Leitura e Preparação dos Dados ---
print("\nIniciando a leitura e preparação dos dados...")
try:
    df_long = pd.read_csv(
        CAMINHO_RENTABILIDADE, sep=';', encoding='latin1', header=None,
        names=['carteira', 'nomeFundo', 'tipoCarteira', 'administrador', 'periodo', 
               'emissao', 'data', 'valorCota', 'quantidade', 'numeroCotistas', 
               'variacaoDia', 'variacaoPeriodo', 'captacoes', 'resgate', 'eventos', 'pl', 'coluna_extra']
    )
    df_long['data'] = pd.to_datetime(df_long['data'], format='%d/%m/%Y', errors='coerce')
    for col in ['valorCota', 'pl']:
        df_long[col] = pd.to_numeric(df_long[col].astype(str).str.replace('.', '', regex=False).str.replace(',', '.', regex=False), errors='coerce')
    df_long.dropna(subset=['data', 'valorCota'], inplace=True)
    print("Dados de rentabilidade carregados com sucesso.")
except FileNotFoundError:
    raise Exception(f"ERRO: Arquivo de rentabilidade não encontrado em: {CAMINHO_RENTABILIDADE}")

# --- Transformação e Adição de Benchmarks ---
df_cota = df_long.pivot_table(index='data', columns='nomeFundo', values='valorCota').fillna(method='ffill')
df_patr = df_long.pivot_table(index='data', columns='nomeFundo', values='pl').fillna(method='ffill')

# Placeholder para Benchmarks (substituir pela sua fonte de dados real)
datas_base = pd.date_range(start=df_cota.columns.min(), end=df_cota.columns.max(), freq='B')
cdi_series = (1 + 0.12 / 252) ** np.arange(len(datas_base)) * 1000
df_benchmarks = pd.DataFrame({'CDI': cdi_series}, index=datas_base)
df_cota = pd.concat([df_cota, df_benchmarks.T], axis=0).T.fillna(method='ffill').T
df_cota.index = df_cota.index.str.strip()
df_patr.index = df_patr.index.str.strip()
print("Benchmarks adicionados e dados transformados.")

# --- Validação e Configuração da Análise ---
if FUNDO_ANALISADO not in df_cota.index:
    raise ValueError(f"ERRO: Fundo '{FUNDO_ANALISADO}' não foi encontrado nos dados!")
    
ref_date = df_cota.columns.max()
comparables = [FUNDO_ANALISADO, 'CDI']
print(f"Data de referência para a análise: {ref_date.strftime('%d/%m/%Y')}")

# --- Dicionário para armazenar os caminhos das imagens geradas ---
paths_imagens_geradas = {}

# %%
# =============================================================================
# CÉLULA 6: CÁLCULO DA TABELA DE RENTABILIDADE MENSAL
# =============================================================================
if not df_cota.empty:
    df_cota_fundo = df_cota.loc[fname].replace(0, np.nan).dropna()
    start_date = df_cota_fundo.index[0]
    years = sorted(list(set(df_cota_fundo.index.year)))
    rows = []
    name_dict = {'Ibovespa': 'IBOV', 'CDI': 'CDI'}

    def format_pct(x):
        if isinstance(x, str) or np.isnan(x): return '-'
        return f'{x:.2%}'.replace('.', ',')

    for year in years:
        for comp in comparables:
            df_aux = df_cota.loc[comp].replace(0, np.nan).dropna()
            df_aux = df_aux[df_aux.index >= start_date]
            name = name_dict.get(comp, 'Fundo')
            ret_mensal = df_aux.resample('M').last() / df_aux.resample('M').last().shift(1) - 1
            ret_mensal.index = ret_mensal.index.to_period('M')
            ytd_val = df_aux[df_aux.index.year == year].iloc[-1] / df_aux[df_aux.index.year == year].iloc[0] - 1
            row = ret_mensal[ret_mensal.index.year == year].to_frame().T
            row.index = pd.MultiIndex.from_tuples([(year, name)])
            row.columns = row.columns.strftime('%b')
            row['YTD'] = ytd_val
            rows.append(row)

    df_return = pd.concat(rows).fillna(np.nan)
    cols_ordem = ['Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun', 'Jul', 'Aug', 'Sep', 'Oct', 'Nov', 'Dec', 'YTD']
    df_return = df_return.reindex(columns=cols_ordem)
    
    for year in years:
        fundo_ret = df_return.loc[(year, 'Fundo')]
        bench_ret = df_return.loc[(year, name_dict.get(comparables[1]))]
        pct_bench = fundo_ret / bench_ret
        pct_bench.name = (year, f'% {name_dict.get(comparables[1])}')
        df_return = pd.concat([df_return, pct_bench.to_frame().T])

    df_return = df_return.sort_index()
    df_return_formatted = df_return.applymap(format_pct)
    df_return_formatted.columns = df_return_formatted.columns.map({'Jan':'Jan', 'Feb':'Fev', 'Mar':'Mar', 'Apr':'Abr', 'May':'Mai', 'Jun':'Jun', 'Jul':'Jul', 'Aug':'Ago', 'Sep':'Set', 'Oct':'Out', 'Nov':'Nov', 'Dec':'Dez', 'YTD':'Ano'})
    display(df_return_formatted)

# %%
# =============================================================================
# CÉLULA 7: RENDERIZAÇÃO DA TABELA DE RENTABILIDADE (COMO IMAGEM)
# =============================================================================
def render_mpl_table(data, col_width=1.5, row_height=0.625, font_size=14,
                     header_color=colorA, row_colors=['#f1f1f2', 'w'], edge_color='w',
                     bbox=[0, 0, 1, 1], ax=None, **kwargs):
    if ax is None:
        size = (np.array(data.shape[::-1]) + np.array([0, 1])) * np.array([col_width, row_height])
        fig, ax = plt.subplots(figsize=size)
        ax.axis('off')
    mpl_table = ax.table(cellText=data.values, bbox=bbox, colLabels=data.columns,
                         rowLabels=[f"{idx[0]} {idx[1]}" for idx in data.index],
                         loc='center', **kwargs)
    mpl_table.auto_set_font_size(False)
    mpl_table.set_fontsize(font_size)
    for k, cell in mpl_table._cells.items():
        if k[0] == 0:
            cell.set_text_props(weight='bold', color=colorB)
            cell.set_facecolor(header_color)
        else:
            cell.set_facecolor(row_colors[(k[0] - 1) % len(row_colors)])
        if k[1] == -1:
            cell.set_text_props(weight='bold', color=colorB)
            cell.get_text().set_ha('right')
        cell.set_edgecolor(edge_color)
        cell.set_linewidth(0)
    return ax.get_figure(), ax

if 'df_return_formatted' in locals() and not df_return_formatted.empty:
    fig_table, ax_table = render_mpl_table(df_return_formatted, col_width=1.0, row_height=0.5, font_size=12)
    
    # --- PONTO DE MODIFICAÇÃO ---
    nome_arquivo_tabela = os.path.join(pasta_saida, f"tabela_rentabilidade_{fname.replace(' ', '_')}.png")
    fig_table.savefig(nome_arquivo_tabela, dpi=300, bbox_inches='tight', pad_inches=0.1)
    print(f"Tabela de rentabilidade salva em: {nome_arquivo_tabela}")
    # -----------------------------
    
    plt.show()

# %%
# =============================================================================
# CÉLULA 8: GRÁFICO DE RETORNO ACUMULADO
# =============================================================================
if not df_cota.empty:
    fig, ax = plt.subplots(figsize=(12, 6))
    df_fund_evol = df_cota.loc[fname].replace(0, np.nan).dropna()
    start = df_fund_evol.index[0]
    df_fund_evol_norm = df_fund_evol / df_fund_evol.iloc[0] * 100
    ax.plot(df_fund_evol_norm, color=color3, lw=2.5, label='Fundo')
    ax.text(df_fund_evol_norm.index[-1], df_fund_evol_norm.iloc[-1], f' {df_fund_evol_norm.iloc[-1]:.2f}', color=color3, fontsize=12, va='center')

    for comp in comparables[1:]:
        df_bench_evol = df_cota.loc[comp].replace(0, np.nan).dropna()
        df_bench_evol = df_bench_evol[df_bench_evol.index >= start]
        df_bench_evol_norm = df_bench_evol / df_bench_evol.iloc[0] * 100
        ax.plot(df_bench_evol_norm, color='k', linestyle='--', alpha=0.7, lw=2, label=comp)
        ax.text(df_bench_evol_norm.index[-1], df_bench_evol_norm.iloc[-1], f' {df_bench_evol_norm.iloc[-1]:.2f}', color='k', alpha=0.8, fontsize=12, va='center')

    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.tick_params(axis='x', rotation=45)
    ax.yaxis.set_major_formatter(ticker.FormatStrFormatter('%.0f'))
    ax.grid(axis='y', linestyle='--', alpha=0.7)
    ax.set_title(f'Retorno Acumulado - {fname}', fontsize=16, pad=20)
    ax.legend(loc='upper left', frameon=False, fontsize=12)
    plt.tight_layout()

    # --- PONTO DE MODIFICAÇÃO ---
    nome_arquivo_retorno = os.path.join(pasta_saida, f"grafico_retorno_{fname.replace(' ', '_')}.png")
    fig.savefig(nome_arquivo_retorno, dpi=300, bbox_inches='tight')
    print(f"Gráfico de retorno salvo em: {nome_arquivo_retorno}")
    # -----------------------------

    plt.show()

# %%
# =============================================================================
# CÉLULA 9: GRÁFICO DE VOLATILIDADE (JANELA MÓVEL DE 22 DIAS)
# =============================================================================
if not df_cota.empty:
    fig, ax = plt.subplots(figsize=(12, 6))
    window = 22
    ret_fundo = df_cota.loc[fname].pct_change()
    vol_fundo = ret_fundo.rolling(window=window).std() * np.sqrt(252)
    ax.plot(vol_fundo, color=color3, lw=2, label='Fundo')

    for comp in comparables[1:]:
        ret_bench = df_cota.loc[comp].pct_change()
        vol_bench = ret_bench.rolling(window=window).std() * np.sqrt(252)
        ax.plot(vol_bench, color='k', linestyle='--', alpha=0.7, lw=2, label=comp)

    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.tick_params(axis='x', rotation=45)
    ax.yaxis.set_major_formatter(ticker.PercentFormatter(1.0))
    ax.grid(axis='y', linestyle='--', alpha=0.7)
    ax.set_title(f'Volatilidade Anualizada (Janela Móvel de {window}d) - {fname}', fontsize=16, pad=20)
    ax.legend(loc='upper left', frameon=False, fontsize=12)
    plt.tight_layout()

    # --- PONTO DE MODIFICAÇÃO ---
    nome_arquivo_vol = os.path.join(pasta_saida, f"grafico_volatilidade_{fname.replace(' ', '_')}.png")
    fig.savefig(nome_arquivo_vol, dpi=300, bbox_inches='tight')
    print(f"Gráfico de volatilidade salvo em: {nome_arquivo_vol}")
    # -----------------------------

    plt.show()

# %%
# =============================================================================
# CÉLULA 10: GRÁFICO DE DRAWDOWN
# =============================================================================
if not df_cota.empty:
    fig, ax = plt.subplots(figsize=(12, 6))
    df_fund_dd = df_cota.loc[fname].dropna()
    roll_max = df_fund_dd.cummax()
    daily_dd = df_fund_dd / roll_max - 1.0
    ax.plot(daily_dd, color=color3, lw=1, label='Fundo')
    ax.fill_between(daily_dd.index, daily_dd, 0, color=color3, alpha=0.3)

    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.tick_params(axis='x', rotation=45)
    ax.yaxis.set_major_formatter(ticker.PercentFormatter(1.0))
    ax.grid(axis='y', linestyle='--', alpha=0.7)
    ax.set_title(f'Drawdown Histórico - {fname}', fontsize=16, pad=20)
    plt.tight_layout()
    
    # --- PONTO DE MODIFICAÇÃO ---
    nome_arquivo_dd = os.path.join(pasta_saida, f"grafico_drawdown_{fname.replace(' ', '_')}.png")
    fig.savefig(nome_arquivo_dd, dpi=300, bbox_inches='tight')
    print(f"Gráfico de drawdown salvo em: {nome_arquivo_dd}")
    # -----------------------------
    
    plt.show()

# %%
# =============================================================================
# CÉLULA 11: GRÁFICO DE EVOLUÇÃO DO PATRIMÔNIO LÍQUIDO (PL)
# =============================================================================
if not df_patr.empty and fname in df_patr.index:
    fig, ax = plt.subplots(figsize=(12, 6))
    pl_fundo = df_patr.loc[fname].dropna()
    pl_fundo = pl_fundo[pl_fundo > 0]

    ax.plot(pl_fundo, color=color2, lw=2)
    ax.fill_between(pl_fundo.index, pl_fundo, 0, color=color2, alpha=0.3)

    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.tick_params(axis='x', rotation=45)
    ax.yaxis.set_major_formatter(ticker.FuncFormatter(lambda x, p: f'R$ {x/1e6:.1f}M'))
    ax.grid(axis='y', linestyle='--', alpha=0.7)
    ax.set_title(f'Evolução do Patrimônio Líquido (PL) - {fname}', fontsize=16, pad=20)
    ax.text(pl_fundo.index[-1], pl_fundo.iloc[-1], f' R$ {pl_fundo.iloc[-1]/1e6:.2f}M', color=color2, fontsize=12, va='bottom')
    
    plt.tight_layout()

    # --- PONTO DE MODIFICAÇÃO ---
    nome_arquivo_pl = os.path.join(pasta_saida, f"grafico_pl_{fname.replace(' ', '_')}.png")
    fig.savefig(nome_arquivo_pl, dpi=300, bbox_inches='tight')
    print(f"Gráfico de PL salvo em: {nome_arquivo_pl}")
    # -----------------------------

    plt.show()
else:
    print(f"Não foram encontrados dados de Patrimônio Líquído (PL) para o fundo '{fname}'.")



paths_imagens_geradas['grafico_pl'] = os.path.join(PASTA_SAIDA_IMAGENS, f"grafico_pl_{nome_arquivo_limpo}.png")
# fig_pl.savefig(paths_imagens_geradas['grafico_pl'], ...)
print(f"IMAGEM GERADA: Gráfico de Evolução do PL")

print("\nTodas as imagens foram geradas e salvas com sucesso.")


# %%
# =============================================================================
# PARTE 3: MONTAGEM DO RELATÓRIO EM WORD E PDF
# =============================================================================
print("\nIniciando a montagem do relatório...")

# --- Abertura do Template e Configuração ---
doc = Document(TEMPLATE_DOCX)
section = doc.sections[0]
section.top_margin, section.bottom_margin, section.left_margin, section.right_margin = [Cm(1), Cm(1), Cm(1.6), Cm(1)]

# --- Criação do Cabeçalho Dinâmico ---
header = section.header
header_table = header.tables[0] if header.tables else header.add_table(rows=1, cols=2, width=section.page_width - section.left_margin - section.right_margin)

# Célula da Esquerda: Nome do fundo e data
p_left = header_table.cell(0, 0).paragraphs[0]
run_name = p_left.add_run(FUNDO_ANALISADO.upper())
run_name.font.name, run_name.font.size, run_name.bold = "Gill Sans MT", Pt(14), True
run_name.font.color.rgb = RGBColor(16, 112, 130)
p_left.add_run("\n")
meses_pt = ["Janeiro", "Fevereiro", "Março", "Abril", "Maio", "Junho", "Julho", "Agosto", "Setembro", "Outubro", "Novembro", "Dezembro"]
data_formatada = f"{meses_pt[ref_date.month - 1]} {ref_date.year}"
run_month = p_left.add_run(data_formatada)
run_month.font.name, run_month.font.size, run_month.italic = "Arial", Pt(12), True

# Célula da Direita: Logo
p_right = header_table.cell(0, 1).paragraphs[0]
p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_right.add_run().add_picture(LOGO_EMPRESA, width=Cm(4.93))

# --- Função Auxiliar para Títulos ---
def add_section_title(document, text):
    p = document.add_paragraph()
    p.paragraph_format.space_before, p.paragraph_format.space_after = Pt(18), Pt(4)
    run = p.add_run(f"\t{text}")
    run.font.name, run.font.size, run.bold = "Gill Sans MT", Pt(11), True
    run.font.color.rgb = RGBColor(89, 89, 89)

# --- Inserção das Imagens no Corpo do Documento ---
add_section_title(doc, "Rentabilidade Mensal")
doc.add_picture(paths_imagens_geradas['tabela_rentabilidade'], width=Cm(18))

add_section_title(doc, "Retorno Acumulado")
doc.add_picture(paths_imagens_geradas['grafico_retorno'], width=Cm(18))

# Tabela para gráficos de risco lado a lado
risk_table = doc.add_table(rows=1, cols=2)
risk_table.columns[0].width, risk_table.columns[1].width = Cm(9), Cm(9)
cell_vol = risk_table.cell(0, 0)
add_section_title(cell_vol, "Volatilidade Anualizada")
cell_vol.add_paragraph().add_run().add_picture(paths_imagens_geradas['grafico_volatilidade'], width=Cm(8.8))
cell_dd = risk_table.cell(0, 1)
add_section_title(cell_dd, "Drawdown Histórico")
cell_dd.add_paragraph().add_run().add_picture(paths_imagens_geradas['grafico_drawdown'], width=Cm(8.8))

add_section_title(doc, "Evolução do Patrimônio Líquido (PL)")
doc.add_picture(paths_imagens_geradas['grafico_pl'], width=Cm(18))

print("Estrutura do documento montada com todas as imagens.")

# --- Salvamento do Documento Final (DOCX e PDF) ---
try:
    nome_final = f"Lamina_{nome_arquivo_limpo}_{ref_date.strftime('%Y-%m')}"
    output_path_docx = os.path.join(PASTA_SAIDA_RELATORIOS, f"{nome_final}.docx")
    output_path_pdf = os.path.join(PASTA_SAIDA_RELATORIOS, f"{nome_final}.pdf")

    doc.save(output_path_docx)
    print(f"Relatório DOCX gerado com sucesso em: {output_path_docx}")
    
    convert(output_path_docx, output_path_pdf)
    print(f"Relatório PDF gerado com sucesso em: {output_path_pdf}")

except Exception as e:
    print(f"ERRO FINAL: Ocorreu um erro ao salvar o relatório: {e}")

print("\nProcesso concluído!")