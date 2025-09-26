# %%
# =============================================================================
# CÉLULA 1: BIBLIOTECAS E CONFIGURAÇÕES GERAIS
# ==============================n===============================================
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import matplotlib.ticker as ticker
import matplotlib.dates as dates
from IPython.display import display
import os # <--- ADICIONADO para criar a pasta de saída

from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert
from datetime import datetime

# --- PONTO DE MODIFICAÇÃO ---
# ! ATENÇÃO: Defina aqui a pasta onde as imagens dos gráficos serão salvas.
pasta_saida = r"C:\Users\Leo\Desktop\Porto_Real\portoauto\src\vortx_estoques\output\lamina_imagens"

# Cria a pasta de saída se ela não existir
os.makedirs(pasta_saida, exist_ok=True)
print(f"As imagens geradas serão salvas em: {pasta_saida}")
# -----------------------------

# Configurações de visualização do Pandas
pd.options.display.max_rows = 100
pd.options.display.max_columns = 50

# Dicionário para tradução de meses (para os gráficos)
month_map_br = {
    '01': 'jan', '02': 'fev', '03': 'mar', '04': 'abr', '05': 'mai', '06': 'jun',
    '07': 'jul', '08': 'ago', '09': 'set', '10': 'out', '11': 'nov', '12': 'dez'
}

# Paleta de cores para os gráficos
"""
color1 = '#4BC1D7'  # Azul claro
color2 = '#286D82'  # Azul escuro
color3 = '#00B1DE'  # Azul ciano
color4 = '#445D6D'  # Cinza azulado
colorA = '#e2eef0'  # Cor de fundo para cabeçalhos de tabela
colorB = '#107082'  # Cor de texto para cabeçalhos de tabela
"""
color1 = '#76C6C5' # '#4BC1D7'
color2 = '#0E5D5F' # '#286D82'
color3 = '#76C6C5' #'#00B1DE'
color4 = '#163F3F'  # '#445D6D' 

colorA = '#F1F9F9'      #'#e2eef0'

colorB = '#0E5D5F'         # '#107082'

# %%
# =============================================================================
# CÉLULA 2: LEITURA E PREPARAÇÃO DOS DADOS DE RENTABILIDADE (CORRIGIDA v3)
# =============================================================================
# ! ATENÇÃO: Confirme se o caminho do arquivo CSV está correto.
caminho_rentabilidade = r"C:\Users\Leo\Desktop\Porto_Real\portoauto\src\vortx_estoques\data\135972-Rentabilidade_Sintetica.csv"

colunas_corretas = [
    'carteira', 'nomeFundo', 'tipoCarteira', 'administrador', 'periodo', 
    'emissao', 'data', 'valorCota', 'quantidade', 'numeroCotistas', 
    'variacaoDia', 'variacaoPeriodo', 'captacoes', 'resgate', 'eventos', 
    'pl', 'coluna_extra'
]

try:
    df_long = pd.read_csv(
        caminho_rentabilidade, 
        sep=';', 
        encoding='latin1',
        header=None,
        names=colunas_corretas
    )
    print("Arquivo 'Rentabilidade_Sintetica.csv' carregado e colunas renomeadas com sucesso!")
except FileNotFoundError:
    print(f"ERRO: Arquivo não encontrado em: {caminho_rentabilidade}")
    df_long = pd.DataFrame()

if not df_long.empty:
    df_long.columns = df_long.columns.str.strip()
    colunas_numericas = ['valorCota', 'pl', 'quantidade', 'numeroCotistas']
    for col in colunas_numericas:
        if col in df_long.columns:
            df_long[col] = df_long[col].astype(str).str.replace('.', '', regex=False).str.replace(',', '.', regex=False)
            df_long[col] = pd.to_numeric(df_long[col], errors='coerce')

    df_long['data'] = pd.to_datetime(df_long['data'], format='%d/%m/%Y', errors='coerce')
    df_long.dropna(subset=['data', 'valorCota'], inplace=True)
    
    print("\n--- Dados após limpeza e conversão ---")
    print(df_long.info())
    display(df_long.head())

# %%
# =============================================================================
# CÉLULA 3: TRANSFORMAÇÃO DOS DADOS (PIVOT)
# =============================================================================
if not df_long.empty:
    print("Transformando dados do formato 'longo' para o 'largo' (pivot)...")

    df_cota = df_long.pivot_table(index='data', columns='nomeFundo', values='valorCota').fillna(method='ffill')
    df_patr = df_long.pivot_table(index='data', columns='nomeFundo', values='pl').fillna(method='ffill')

    df_cota = df_cota.T
    df_patr = df_patr.T
    
    df_cota.index = df_cota.index.str.strip()
    df_patr.index = df_patr.index.str.strip()

    print("\n--- Formato final de df_cota ---")
    display(df_cota.head())
else:
    print("DataFrame de entrada está vazio. Pulando a transformação.")
    df_cota = pd.DataFrame()
    df_patr = pd.DataFrame()

# %%
# =============================================================================
# CÉLULA 4: ADIÇÃO DE DADOS DE BENCHMARK (COM DADOS REAIS DO BACEN)
# =============================================================================

def consulta_bc(codigo_bcb):
    """
    Busca uma série temporal da API do Banco Central do Brasil (Bacen).
    """
    try:
        # URL atualizada com a data inicial para otimizar a consulta
        url = f'https://api.bcb.gov.br/dados/serie/bcdata.sgs.{codigo_bcb}/dados?formato=json&dataInicial=28/03/2017'
        df = pd.read_json(url)
        df['data'] = pd.to_datetime(df['data'], dayfirst=True)
        df.set_index('data', inplace=True)
        df['valor'] = pd.to_numeric(df['valor'], errors='coerce')
        print(f"Série {codigo_bcb} do Bacen carregada com sucesso!")
        return df
    except Exception as e:
        print(f"ERRO: Falha ao buscar dados da API do Bacen para a série {codigo_bcb}.")
        print(f"Detalhe do erro: {e}")
        return pd.DataFrame()

# --- LÓGICA PRINCIPAL DA CÉLULA ---

if not df_cota.empty:
    df_cdi_diario = consulta_bc(12)

    if not df_cdi_diario.empty:
        df_cdi_acumulado = (1 + df_cdi_diario['valor'] / 100).cumprod()
        df_cota.loc['CDI'] = df_cdi_acumulado
        
        # --- MUDANÇA AQUI ---
        # Trocamos o método antigo pelo novo, recomendado pelo Pandas.
        df_cota.loc['CDI'] = df_cota.loc['CDI'].ffill()
        # ---------------------

        print("\nBenchmark CDI (dados reais) adicionado e alinhado com sucesso.")
        
    else:
        print("\nAVISO: Não foi possível carregar os dados do CDI. A análise de performance pode falhar.")
else:
    print("AVISO: DataFrame de cotas está vazio. Pulando o carregamento do benchmark.")
    
# %%
# =============================================================================
# CÉLULA 5: CONFIGURAÇÃO DA ANÁLISE E SELEÇÃO DO FUNDO (CORRIGIDA)
# =============================================================================
fname = 'FIDC FCT II'
comparables = []
interval = 3

if not df_cota.empty:
    if fname not in df_cota.index:
        print(f"ERRO: Fundo '{fname}' não encontrado! Escolha um da lista abaixo:")
        print(df_cota.index.tolist())
        raise ValueError("Fundo não encontrado")

    comp_dict = {
        'FIDC FCT II SR2': ['CDI'],
        'FIDC FCT II':     ['CDI'],
    }
    ref_date = df_cota.columns.max()
    print(f"Análise para o fundo: '{fname}'")
    print(f"Data de referência (último dado): {ref_date.strftime('%d/%m/%Y')}")
    comparables = [fname] + comp_dict.get(fname, ['CDI'])
else:
    print("Não há dados de cota para configurar a análise.")

# %%
# =============================================================================
# CÉLULA 6 (VERSÃO MODIFICADA): EXIBINDO VALORES NEGATIVOS PARA % CDI
# =============================================================================
# Esta versão foi ajustada para exibir o valor calculado do % CDI mesmo
# quando o fundo tem rentabilidade negativa, contanto que o CDI seja positivo.

def render_mpl_table_detailed(data, col_width=1.5, row_height=0.625, font_size=10,
                              header_color=colorA, row_colors=['#f1f1f2', 'w'], edge_color='w',
                              bbox=[0, 0, 1, 1], ax=None, **kwargs):
    if ax is None:
        size = (np.array(data.shape[::-1]) + np.array([0, 1])) * np.array([col_width, row_height])
        fig, ax = plt.subplots(figsize=size)
        ax.axis('off')
    
    mpl_table = ax.table(cellText=data.values, cellLoc='center', bbox=bbox, colLabels=data.columns, 
                         loc='center', colLoc='center', **kwargs)
    
    mpl_table.auto_set_font_size(False)
    mpl_table.set_fontsize(font_size)

    for k, cell in mpl_table._cells.items():
        cell.set_edgecolor(edge_color)
        cell.set_linewidth(0)
        
        if k[0] == 0: # Header Row
            cell.set_text_props(weight='bold', color=colorB)
            cell.set_facecolor(header_color)
        else: # Data Rows
            if (k[0] - 1) % 3 == 2:
                 cell.set_facecolor('#F0F0F0') 
            else:
                 cell.set_facecolor(row_colors[k[0] % len(row_colors)])

        if k[1] <= 0: # Header Columns
            cell.set_text_props(weight='bold', color=colorB)
            cell.set_facecolor(header_color)
            
    return ax.get_figure(), ax


if not df_cota.empty:
    print("Iniciando geração da tabela de performance (exibindo valores negativos)...")
    
    df_cotas_t = df_cota.T
    df_cotas_pct = df_cotas_t.pct_change().dropna()

    years = sorted(set(df_cotas_t.index.year))
    months_str = ['Jan', 'Fev', 'Mar', 'Abr', 'Maio', 'Jun', 'Jul', 'Ago', 'Set', 'Out', 'Nov', 'Dez']
    all_months_num = [f'{i:02d}' for i in range(1, 13)]
    bench_name = comparables[1]
    
    def format_pct_smart(x):
        if isinstance(x, str) or pd.isna(x): return '-'
        return f'{x:.2%}'.replace('.', ',')

    def format_pct_of_bench(x):
        if isinstance(x, str) or pd.isna(x): return '-'
        return f'{x*100:.1f}%'.replace('.', ',')

    rows = []
    for year in years:
        row_fundo = [year, fname]
        row_bench = [year, bench_name]
        row_pct_bench = [year, f"% {bench_name}"]
        
        series_fundo = df_cotas_pct[fname]
        series_bench = df_cotas_pct[bench_name]
        
        for month in all_months_num:
            ym = f'{year}-{month}'
            mask = df_cotas_pct.index.strftime('%Y-%m') == ym
            if mask.sum() == 0:
                row_fundo.append('-')
                row_bench.append('-')
                row_pct_bench.append('-')
            else:
                ret_fundo = (series_fundo[mask] + 1).product() - 1
                ret_bench = (series_bench[mask] + 1).product() - 1
                
                row_fundo.append(format_pct_smart(ret_fundo))
                row_bench.append(format_pct_smart(ret_bench))
                
                # --- MUDANÇA CRÍTICA AQUI ---
                # Agora, calcula o % CDI contanto que o CDI seja positivo.
                if ret_bench > 0:
                    pct_cdi = ret_fundo / ret_bench
                    row_pct_bench.append(format_pct_of_bench(pct_cdi))
                else:
                    # Mantém 'N/A' apenas se o CDI for negativo ou zero.
                    row_pct_bench.append('N/A')

        # Mesma lógica para o Ano (YTD)
        mask_ano = df_cotas_pct.index.year == year
        ret_fundo_ano = (series_fundo[mask_ano] + 1).product() - 1
        ret_bench_ano = (series_bench[mask_ano] + 1).product() - 1
        row_fundo.append(format_pct_smart(ret_fundo_ano))
        row_bench.append(format_pct_smart(ret_bench_ano))
        if ret_bench_ano > 0:
            row_pct_bench.append(format_pct_of_bench(ret_fundo_ano / ret_bench_ano))
        else:
            row_pct_bench.append('N/A')

        # Mesma lógica para "Desde o Início"
        start_date = df_cota.loc[fname].dropna().index[0]
        mask_inicio = (df_cotas_pct.index.year <= year) & (df_cotas_pct.index >= start_date)
        ret_fundo_inicio = (series_fundo[mask_inicio] + 1).product() - 1
        ret_bench_inicio = (series_bench[mask_inicio] + 1).product() - 1
        row_fundo.append(format_pct_smart(ret_fundo_inicio))
        row_bench.append(format_pct_smart(ret_bench_inicio))
        if ret_bench_inicio > 0:
             row_pct_bench.append(format_pct_of_bench(ret_fundo_inicio / ret_bench_inicio))
        else:
            row_pct_bench.append('N/A')
        
        rows.append(row_fundo)
        rows.append(row_bench)
        rows.append(row_pct_bench)

    table_cols = ['Ano', 'Cota'] + months_str + ['YTD', 'Desde o\nInício']
    df_table = pd.DataFrame(rows, columns=table_cols)
    
    anos = df_table['Ano'].tolist()
    anos_formatados = []
    for i, ano in enumerate(anos):
        if i % 3 == 0:
            anos_formatados.append(ano)
        else:
            anos_formatados.append('')
    df_table['Ano'] = anos_formatados

    fig, ax = render_mpl_table_detailed(df_table, col_width=1.2, row_height=0.5, font_size=9, row_colors=['#FFFFFF', '#FFFFFF'])
    
    nome_arquivo_limpo = fname.replace(' ', '_')
    path_tabela_rentabilidade = os.path.join(pasta_saida, f"tabela_rentabilidade_{nome_arquivo_limpo}.png")
    
    fig.savefig(path_tabela_rentabilidade, dpi=300, bbox_inches='tight', pad_inches=0.1)
    print(f"Tabela de rentabilidade completa salva em: {path_tabela_rentabilidade}")
    plt.show()

else:
    print("Não há dados de cota para gerar a tabela de rentabilidade.")
    
# %%
# =============================================================================
# CÉLULA 8: GRÁFICO DE RETORNO ACUMULADO
# =============================================================================
if not df_cota.empty:
    fig, ax = plt.subplots(figsize=(12, 6))
    df_fund_evol = df_cota.loc[fname].replace(0, np.nan).dropna()
    start = df_fund_evol.index[0]
    df_fund_evol_norm = df_fund_evol / df_fund_evol.iloc[0] * 100
    ax.plot(df_fund_evol_norm, color=color3, lw=2.5, label='Fundo')
    ax.text(df_fund_evol_norm.index[-1], df_fund_evol_norm.iloc[-1], f' {df_fund_evol_norm.iloc[-1]:.2f}', color=color3, fontsize=12, va='center')

    for comp in comparables[1:]:
        df_bench_evol = df_cota.loc[comp].replace(0, np.nan).dropna()
        df_bench_evol = df_bench_evol[df_bench_evol.index >= start]
        df_bench_evol_norm = df_bench_evol / df_bench_evol.iloc[0] * 100
        ax.plot(df_bench_evol_norm, color='k', linestyle='--', alpha=0.7, lw=2, label=comp)
        ax.text(df_bench_evol_norm.index[-1], df_bench_evol_norm.iloc[-1], f' {df_bench_evol_norm.iloc[-1]:.2f}', color='k', alpha=0.8, fontsize=12, va='center')

    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.tick_params(axis='x', rotation=45)
    ax.yaxis.set_major_formatter(ticker.FormatStrFormatter('%.0f'))
    ax.grid(axis='y', linestyle='--', alpha=0.7)
    ax.set_title(f'Retorno Acumulado - {fname}', fontsize=16, pad=20)
    ax.legend(loc='upper left', frameon=False, fontsize=12)
    plt.tight_layout()

    # --- PONTO DE MODIFICAÇÃO ---
    nome_arquivo_retorno = os.path.join(pasta_saida, f"grafico_retorno_{fname.replace(' ', '_')}.png")
    fig.savefig(nome_arquivo_retorno, dpi=300, bbox_inches='tight')
    print(f"Gráfico de retorno salvo em: {nome_arquivo_retorno}")
    # -----------------------------

    plt.show()

# %%
# =============================================================================
# CÉLULA 9: GRÁFICO DE VOLATILIDADE (JANELA MÓVEL DE 22 DIAS)
# =============================================================================
if not df_cota.empty:
    fig, ax = plt.subplots(figsize=(12, 6))
    window = 22
    ret_fundo = df_cota.loc[fname].pct_change()
    vol_fundo = ret_fundo.rolling(window=window).std() * np.sqrt(252)
    ax.plot(vol_fundo, color=color3, lw=2, label='Fundo')

    for comp in comparables[1:]:
        ret_bench = df_cota.loc[comp].pct_change()
        vol_bench = ret_bench.rolling(window=window).std() * np.sqrt(252)
        ax.plot(vol_bench, color='k', linestyle='--', alpha=0.7, lw=2, label=comp)

    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.tick_params(axis='x', rotation=45)
    ax.yaxis.set_major_formatter(ticker.PercentFormatter(1.0))
    ax.grid(axis='y', linestyle='--', alpha=0.7)
    ax.set_title(f'Volatilidade Anualizada (Janela Móvel de {window}d) - {fname}', fontsize=16, pad=20)
    ax.legend(loc='upper left', frameon=False, fontsize=12)
    plt.tight_layout()

    # --- PONTO DE MODIFICAÇÃO ---
    nome_arquivo_vol = os.path.join(pasta_saida, f"grafico_volatilidade_{fname.replace(' ', '_')}.png")
    fig.savefig(nome_arquivo_vol, dpi=300, bbox_inches='tight')
    print(f"Gráfico de volatilidade salvo em: {nome_arquivo_vol}")
    # -----------------------------

    plt.show()

# %%
# =============================================================================
# CÉLULA 10: GRÁFICO DE DRAWDOWN
# =============================================================================
if not df_cota.empty:
    fig, ax = plt.subplots(figsize=(12, 6))
    df_fund_dd = df_cota.loc[fname].dropna()
    roll_max = df_fund_dd.cummax()
    daily_dd = df_fund_dd / roll_max - 1.0
    ax.plot(daily_dd, color=color3, lw=1, label='Fundo')
    ax.fill_between(daily_dd.index, daily_dd, 0, color=color3, alpha=0.3)

    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.tick_params(axis='x', rotation=45)
    ax.yaxis.set_major_formatter(ticker.PercentFormatter(1.0))
    ax.grid(axis='y', linestyle='--', alpha=0.7)
    ax.set_title(f'Drawdown Histórico - {fname}', fontsize=16, pad=20)
    plt.tight_layout()
    
    # --- PONTO DE MODIFICAÇÃO ---
    nome_arquivo_dd = os.path.join(pasta_saida, f"grafico_drawdown_{fname.replace(' ', '_')}.png")
    fig.savefig(nome_arquivo_dd, dpi=300, bbox_inches='tight')
    print(f"Gráfico de drawdown salvo em: {nome_arquivo_dd}")
    # -----------------------------
    
    plt.show()

# %%
# =============================================================================
# CÉLULA 11: GRÁFICO DE EVOLUÇÃO DO PATRIMÔNIO LÍQUIDO (PL)
# =============================================================================
if not df_patr.empty and fname in df_patr.index:
    fig, ax = plt.subplots(figsize=(12, 6))
    pl_fundo = df_patr.loc[fname].dropna()
    pl_fundo = pl_fundo[pl_fundo > 0]

    ax.plot(pl_fundo, color=color2, lw=2)
    ax.fill_between(pl_fundo.index, pl_fundo, 0, color=color2, alpha=0.3)

    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.tick_params(axis='x', rotation=45)
    ax.yaxis.set_major_formatter(ticker.FuncFormatter(lambda x, p: f'R$ {x/1e6:.1f}M'))
    ax.grid(axis='y', linestyle='--', alpha=0.7)
    ax.set_title(f'Evolução do Patrimônio Líquido (PL) - {fname}', fontsize=16, pad=20)
    ax.text(pl_fundo.index[-1], pl_fundo.iloc[-1], f' R$ {pl_fundo.iloc[-1]/1e6:.2f}M', color=color2, fontsize=12, va='bottom')
    
    plt.tight_layout()

    # --- PONTO DE MODIFICAÇÃO ---
    nome_arquivo_pl = os.path.join(pasta_saida, f"grafico_pl_{fname.replace(' ', '_')}.png")
    fig.savefig(nome_arquivo_pl, dpi=300, bbox_inches='tight')
    print(f"Gráfico de PL salvo em: {nome_arquivo_pl}")
    # -----------------------------

    plt.show()
else:
    print(f"Não foram encontrados dados de Patrimônio Líquído (PL) para o fundo '{fname}'.")

#%%
# estoque
import os
import pandas as pd
import glob
import numpy as np
import matplotlib.pyplot as plt

# --- 1. FUNÇÃO DE PROCESSAMENTO CORRIGIDA ---

def process_new_data(base_path):
    """
    Lê e processa os arquivos de estoque com a codificação correta (UTF-16).
    """
    list_files = glob.glob(os.path.join(base_path, "*.csv"))

    if not list_files:
        raise ValueError(f"Nenhum arquivo CSV encontrado em {base_path}")

    float_cols = [
        'Valor Aquisicao', 'Valor Nominal', 'Valor Presente', 'PDD Vencido',
        'PDD Total', 'Taxa Operada Originador', 'CET Mensal', 'Taxa CCB',
        'Taxa Originador Split', 'Taxa Split FIDC'
    ]
    date_cols = ['Data Aquisicao', 'Data Vencimento', 'Data Referencia', 'Data de Nascimento']
    id_cols = ['CCB', 'SEU NUMERO', 'PARCELA']
    
    all_dfs = []
    for file in list_files:
        print(f"Lendo arquivo: {os.path.basename(file)}")
        try:
            # --- CORREÇÃO FINAL AQUI ---
            # Trocamos a codificação para 'utf-16'
            df = pd.read_csv(
                file, 
                encoding='utf-16', # A MUDANÇA ESSENCIAL
                sep='\t', 
                engine='python',
                on_bad_lines='warn',
                header=0
            )
            
            # Com a codificação correta, a limpeza de colunas pode ser mais simples
            df.columns = df.columns.str.strip()

            if not df.empty:
                for col in float_cols:
                    if col in df.columns:
                        # Convertendo para string antes para garantir que .str funcione
                        df[col] = df[col].astype(str).str.replace(',', '.').astype(float)
                
                for col in date_cols:
                     if col in df.columns:
                        df[col] = pd.to_datetime(df[col], errors='coerce', dayfirst=True)

                all_dfs.append(df)
            else:
                 print(f"AVISO: O DataFrame para o arquivo {os.path.basename(file)} está vazio após a leitura.")

        except Exception as e:
            print(f"Erro CRÍTICO ao processar o arquivo {file}: {e}")
            continue 

    if not all_dfs:
        raise ValueError("Nenhum dado foi carregado com sucesso.")

    df_final_estoque = pd.concat(all_dfs, ignore_index=True)
    
    # Garante que as colunas de ID existam antes de usá-las
    for col in id_cols:
        if col not in df_final_estoque.columns:
            raise KeyError(f"Erro Crítico: A coluna de ID '{col}' não foi encontrada após processar todos os arquivos. Verifique os CSVs.")

    # Verificação de duplicatas
    _1 = df_final_estoque[id_cols].shape[0]
    _2 = df_final_estoque[id_cols].drop_duplicates().shape[0]
    if _1 != _2:
        print(f"Aviso: Foram encontradas duplicatas ou inconsistências. Linhas totais: {_1}, Linhas únicas: {_2}")

    print("Dados de estoque consolidados e processados com sucesso!")
    return df_final_estoque


# ... (Cole aqui as suas funções de plotagem: vencimento_mensal, vencimento_anual, etc. Elas não precisam de alteração) ...
def vencimento_mensal(df_aberto, col_valor, col_vcto, cutoff=6):
    month_map = {'01': 'jan', '02': 'fev', '03': 'mar', '04': 'abr', '05': 'maio', '06': 'jun', 
                 '07': 'jul', '08': 'ago', '09': 'set', '10': 'out', '11': 'nov', '12': 'dez'}
    df2 = df_aberto.groupby(col_vcto)[col_valor].sum()
    df3 = (df2/df2.sum()).cumsum()
    df2 = df2[:cutoff]
    df3 = df3[:cutoff]
    if df2.empty:
        print("Não há dados para gerar o gráfico de vencimento mensal.")
        return plt.subplots(figsize=(8,4.5))
    ymin1 = 0
    ymax1 = df2.max() * 2
    ymax2 = 1.3
    ymin2 = df3.min() * 2.5 - 1.5 * ymax2
    fig, ax = plt.subplots(figsize=(8,4.5))
    ax2 = ax.twinx()
    bars = ax.bar(df2.index, np.clip(df2.values, 0.01*df2.max(), df2.max()), color=color2)
    ax2.plot(df3.index, df3.values, color=color4, lw=2, marker='o')
    xs = [bar.get_x() for bar in bars]
    w = bars[0].get_width()
    xmin = min(xs) - 0.05 * (max(xs) - min(xs))
    xmax = max(xs) + 0.05 * (max(xs) - min(xs)) + w
    def map_month(x):
        return f'{month_map[x[-2:]]}/{x[2:4]}'
    ax.set_xlabel('')
    ax.set_xticks(df3.index.tolist())
    ax.set_xticklabels(df2.index.map(map_month))
    ax.set_xlim(xmin, xmax)
    ax.set_ylim(ymin1, ymax1)
    ax2.set_ylim(ymin2, ymax2)
    ax.set_yticks([])
    ax2.set_yticks([])
    ax.tick_params(bottom=False)
    ax.set_frame_on(False)
    ax2.spines['top'].set_visible(False)
    ax2.spines['right'].set_visible(False)
    ax2.spines['bottom'].set_visible(False)
    ax2.spines['left'].set_visible(False)
    max_height = max([bar.get_height() for bar in bars])
    for i, bar in enumerate(bars):
        h = bar.get_height() + max_height * 0.04
        xc = bar.get_x() + bar.get_width()/2
        value = df2.values[i]/1000
        value = f'{value:,.0f}'.replace(',', '.')
        ax.text(xc, h, f'{value}', ha='center')
    for x, y in df3.items():
        value = f'{y:.0%}' if ((y<.990) or (y==1)) else '>99%'
        ax2.text(x, y + (ymax2 - ymin2) * 0.04, value, ha='center')
    return fig, ax

def vencimento_anual(df_aberto, col_valor, col_vcto):
    df2 = df_aberto.groupby(col_vcto)[col_valor].sum()
    df3 = df2.cumsum()/df2.sum()
    if df2.empty:
        print("Não há dados para gerar o gráfico de vencimento anual.")
        return plt.subplots(figsize=(8,4.5))
    dx = 1 
    xmin = df3.index.min() - dx
    xmax = df3.index.max() + dx
    ymin1 = 0
    ymax1 = df2.max() * 2
    ymax2 = 1.15
    ymin2 = df3.min() * 2.5 - 1.5 * ymax2
    fig, ax = plt.subplots(figsize=(8,4.5))
    ax2 = ax.twinx()
    bars = ax.bar(df2.index, np.clip(df2.values, 0.01*df2.max(), df2.max()), color=color2)
    ax2.plot(df3.index, df3.values, color=color4, lw=2, marker='o')
    ax.set_xlabel('')
    ax.set_xticks(df3.index.tolist())
    ax.set_xlim(xmin, xmax)
    ax.set_ylim(ymin1, ymax1)
    ax2.set_ylim(ymin2, ymax2)
    ax.set_yticks([])
    ax2.set_yticks([])
    ax.tick_params(bottom=False)
    ax.set_frame_on(False)
    ax2.spines['top'].set_visible(False)
    ax2.spines['right'].set_visible(False)
    ax2.spines['bottom'].set_visible(False)
    ax2.spines['left'].set_visible(False)
    max_height = max([bar.get_height() for bar in bars])
    for i, bar in enumerate(bars):
        h = bar.get_height() + max_height * 0.04
        xc = bar.get_x() + bar.get_width()/2
        value = df2.values[i]/1000
        value = f'{value:,.0f}'.replace(',', '.')
        ax.text(xc, h, f'{value}', ha='center')
    for x, y in df3.items():
        value = f'{y:.0%}' if ((y<.990) or (np.isclose(y, 1))) else '>99%'
        ax2.text(x, y+0.06, value, ha='center')
    return fig, ax

def concentracao_ente(df_aberto, col_valor, col_class, cutoff):
    df11 = df_aberto.groupby(col_class, observed=False)[col_valor].sum().sort_values(ascending=False)
    ncut = min(len(df11), cutoff)
    if ncut >= cutoff:
        x = df11[:cutoff].index.tolist() + ['Outros']
        y = list(df11[:cutoff].values) + [df11.iloc[cutoff:].sum()]
    else:
        x = df11.index.tolist()
        y = list(df11.values)
    fig, ax = plt.subplots(figsize=(6,4))
    bars = ax.barh(x, y, color=color2)
    plt.gca().invert_yaxis()
    ax.set_xlabel('')
    ax.set_xticks([])
    ax.tick_params(left=False, bottom=False)
    ax.set_frame_on(False)
    max_width = max([bar.get_width() for bar in bars])
    for i, bar in enumerate(bars):
        bw = bar.get_width()
        byc = bar.get_y() + bar.get_height()/2
        ax.text(bw + 0.01*max_width, byc, f'{y[i]/1000:,.0f}'.replace(',', '.'), va='center', ha='left')
    return fig, ax

def dist_capag(df_in, class_col, col_valor):
    df_capag = df_in.dropna(subset=[class_col]).groupby(class_col, observed=True)[col_valor].sum()
    fig, ax = plt.subplots(figsize=(8,4.5))
    bars = ax.bar(df_capag.index, np.clip(df_capag.values, 0.01*df_capag.max(), df_capag.max()), color=color2)
    ax.set_xlabel('')
    ax.set_xticks(df_capag.index.tolist())
    ax.set_yticks([])
    ax.tick_params(bottom=False)
    ax.set_frame_on(False)
    for spine in ax.spines.values():
        spine.set_visible(False)
    max_height = max([bar.get_height() for bar in bars])
    for i, bar in enumerate(bars):
        h = bar.get_height() + max_height * 0.04
        xc = bar.get_x() + bar.get_width()/2
        value = df_capag.values[i]/1000
        value = f'{value:,.0f}'.replace(',', '.')
        ax.text(xc, h, f'{value}', ha='center')
    return fig, ax


# --- 3. EXECUÇÃO PRINCIPAL E GERAÇÃO DOS GRÁFICOS ---

downloads_base_path = r"C:\Users\Leo\Desktop\Porto_Real\portoauto\src\vortx_estoques\data\estoque_consolidado_agosto"
output_path = r"C:\Users\Leo\Desktop\Porto_Real\portoauto\src\vortx_estoques\output\lamina_imagens"

# Processa os dados
df_estoque = process_new_data(downloads_base_path)

# Filtra o dataframe para 'A vencer'
df_avencer = df_estoque[df_estoque['Status'] == 'A vencer'].copy()

# --- Gráfico de Vencimento Mensal ---
df_avencer['VCTO_MES'] = df_avencer['Data Vencimento'].dt.strftime('%Y-%m')
fig_m, ax_m = vencimento_mensal(df_avencer, col_valor='Valor Presente', col_vcto='VCTO_MES', cutoff=10)
temp_venc_mensal = os.path.join(output_path, "vencimento_mensal.png")
fig_m.savefig(temp_venc_mensal, bbox_inches='tight')
plt.close(fig_m)
print(f"Gráfico de vencimento mensal salvo em: {temp_venc_mensal}")

# --- Gráfico de Vencimento Anual ---
df_avencer['VCTO_ANO'] = df_avencer['Data Vencimento'].dt.year
fig_a, ax_a = vencimento_anual(df_avencer, col_valor='Valor Presente', col_vcto='VCTO_ANO')
temp_venc_anual = os.path.join(output_path, "vencimento_anual.png")
fig_a.savefig(temp_venc_anual, bbox_inches='tight')
plt.close(fig_a)
print(f"Gráfico de vencimento anual salvo em: {temp_venc_anual}")

# --- Gráfico de Concentração por Ente (UF) ---
fig_ente, ax_ente = concentracao_ente(df_avencer, col_valor='Valor Presente', col_class='UF', cutoff=15)
temp_conc_ente = os.path.join(output_path, "concentracao_uf.png")
fig_ente.savefig(temp_conc_ente, bbox_inches='tight')
plt.close(fig_ente)
print(f"Gráfico de concentração por UF salvo em: {temp_conc_ente}")

# --- Gráfico de Distribuição por CAPAG ---
fig_capag, ax_capag = dist_capag(df_estoque, class_col='CAPAG', col_valor='Valor Presente')
temp_dist_capag = os.path.join(output_path, "distribuicao_capag.png")
fig_capag.savefig(temp_dist_capag, bbox_inches='tight')
plt.close(fig_capag)
print(f"Gráfico de distribuição por CAPAG salvo em: {temp_dist_capag}")

print("\nTodos os gráficos foram gerados com sucesso.")



# %%
# =============================================================================
# CÉLULA 1: BIBLIOTECAS E CONFIGURAÇÕES DE CAMINHOS
# =============================================================================


# --- CONFIGURAÇÕES DO RELATÓRIO ---
report_name = 'Lâmina de Performance e Estoque'
ref_date = datetime.now()

# --- PONTO DE MODIFICAÇÃO PRINCIPAL ---
# ! ATENÇÃO: Coloque aqui o nome exato do fundo usado para gerar os gráficos de rentabilidade.
fname_rentabilidade = 'FIDC FCT II'
# -----------------------------------------

# --- CAMINHOS PARA AS PASTAS ---
# Assumindo que todas as imagens estão na mesma pasta para simplificar
pasta_imagens = r"C:\Users\Leo\Desktop\Porto_Real\portoauto\src\vortx_estoques\output\lamina_imagens"

# Caminhos para template e logos
template_path = r"C:\Users\Leo\Desktop\Porto_Real\portoreal\notebooks\template.docx"
logo_path = r"C:\Users\Leo\Desktop\Porto_Real\portoreal\images\logo.png"

# Caminho final do relatório
output_path_base = r"C:\Users\Leo\Desktop\Porto_Real\portoauto\src\vortx_estoques\output"
nome_arquivo_saida = f"Lamina_Completa_{ref_date.strftime('%Y-%m')}.docx"
output_path_docx = os.path.join(output_path_base, nome_arquivo_saida)


# --- CAMINHOS PARA TODAS AS IMAGENS (COM LÓGICA CORRIGIDA) ---

# 1. Constrói a parte dinâmica do nome do arquivo de rentabilidade
nome_arquivo_limpo = fname_rentabilidade.replace(' ', '_')

# 2. Define os caminhos para as imagens de RENTABILIDADE (Dinâmico)
path_tabela_rentabilidade = os.path.join(pasta_imagens, f"tabela_rentabilidade_{nome_arquivo_limpo}.png")
path_grafico_retorno = os.path.join(pasta_imagens, f"grafico_retorno_{nome_arquivo_limpo}.png")
path_grafico_volatilidade = os.path.join(pasta_imagens, f"grafico_volatilidade_{nome_arquivo_limpo}.png")
path_grafico_drawdown = os.path.join(pasta_imagens, f"grafico_drawdown_{nome_arquivo_limpo}.png")
path_grafico_pl = os.path.join(pasta_imagens, f"grafico_pl_{nome_arquivo_limpo}.png")

# 3. Define os caminhos para as imagens de ESTOQUE (Fixo)
path_vencimento_mensal = os.path.join(pasta_imagens, "vencimento_mensal.png")
path_vencimento_anual = os.path.join(pasta_imagens, "vencimento_anual.png")
path_concentracao_uf = os.path.join(pasta_imagens, "concentracao_uf.png")
path_distribuicao_capag = os.path.join(pasta_imagens, "distribuicao_capag.png")


# %%
# =============================================================================
# CÉLULA 2: CRIAÇÃO E CONFIGURAÇÃO DO DOCUMENTO (Sem alterações)
# =============================================================================
doc = Document(template_path)
if doc.paragraphs and doc.paragraphs[0].text.strip() == "":
    p = doc.paragraphs[0]._element
    p.getparent().remove(p)

section = doc.sections[0]
section.top_margin = Cm(1)
section.bottom_margin = Cm(1)
section.left_margin = Cm(1.6)
section.right_margin = Cm(1)
section.header_distance = Cm(1.14)

header = section.header
header_table = header.tables[0] if header.tables else header.add_table(rows=1, cols=2, width=Cm(18))

left_cell = header_table.cell(0, 0)
left_cell.text = ""
p_left = left_cell.paragraphs[0]
run_name = p_left.add_run(report_name.upper())
run_name.font.name = "Gill Sans MT"
run_name.font.size = Pt(14)
run_name.bold = True
run_name.font.color.rgb = RGBColor(16, 112, 130)
p_left.add_run("\n")
meses_pt = ["Janeiro", "Fevereiro", "Março", "Abril", "Maio", "Junho", "Julho", "Agosto", "Setembro", "Outubro", "Novembro", "Dezembro"]
data_formatada = f"{meses_pt[ref_date.month - 1]} {ref_date.year}"
run_month = p_left.add_run(data_formatada)
run_month.font.name = "Arial"
run_month.font.size = Pt(12)
run_month.italic = True
run_month.font.color.rgb = RGBColor(38, 38, 38)

right_cell = header_table.cell(0, 1)
right_cell.text = ""
p_right = right_cell.paragraphs[0]
p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_right.add_run().add_picture(logo_path, width=Cm(4.93))


# %%
# =============================================================================
# CÉLULA 3: MONTAGEM DO CORPO DO RELATÓRIO (Sem alterações)
# =============================================================================

def add_section_title(cell, text):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = "Gill Sans MT"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(89, 89, 89)
    run.bold = True

title_p = doc.add_paragraph()
title_p.paragraph_format.space_before = Pt(12)
run = title_p.add_run("\tRentabilidade Mensal")
run.font.name = "Gill Sans MT"
run.font.size = Pt(11)
run.bold = True
doc.add_picture(path_tabela_rentabilidade, width=Cm(17.8))

graphs_table = doc.add_table(rows=4, cols=2)
graphs_table.autofit = False
graphs_table.allow_autofit = False

col_width = Cm(8.9)
for col in graphs_table.columns:
    col.width = col_width

img_width = Cm(8.5)

cell_retorno = graphs_table.cell(0, 0)
add_section_title(cell_retorno, "Retorno Acumulado")
cell_retorno.add_paragraph().add_run().add_picture(path_grafico_retorno, width=img_width)

cell_pl = graphs_table.cell(0, 1)
add_section_title(cell_pl, "Evolução do Patrimônio Líquido (PL)")
cell_pl.add_paragraph().add_run().add_picture(path_grafico_pl, width=img_width)

cell_vol = graphs_table.cell(1, 0)
add_section_title(cell_vol, "Volatilidade Anualizada")
cell_vol.add_paragraph().add_run().add_picture(path_grafico_volatilidade, width=img_width)

cell_dd = graphs_table.cell(1, 1)
add_section_title(cell_dd, "Drawdown Histórico")
cell_dd.add_paragraph().add_run().add_picture(path_grafico_drawdown, width=img_width)

cell_vcto_m = graphs_table.cell(2, 0)
add_section_title(cell_vcto_m, "Vencimento Mensal")
cell_vcto_m.add_paragraph().add_run().add_picture(path_vencimento_mensal, width=img_width)

cell_vcto_a = graphs_table.cell(2, 1)
add_section_title(cell_vcto_a, "Vencimento Anual")
cell_vcto_a.add_paragraph().add_run().add_picture(path_vencimento_anual, width=img_width)

cell_uf = graphs_table.cell(3, 0)
add_section_title(cell_uf, "Concentração por UF")
cell_uf.add_paragraph().add_run().add_picture(path_concentracao_uf, width=img_width)

cell_capag = graphs_table.cell(3, 1)
add_section_title(cell_capag, "Distribuição por CAPAG")
cell_capag.add_paragraph().add_run().add_picture(path_distribuicao_capag, width=img_width)


# %%
# =============================================================================
# CÉLULA 4: RODAPÉ E SALVAMENTO DO ARQUIVO (Sem alterações)
# =============================================================================
try:
    os.makedirs(output_path_base, exist_ok=True)
    doc.save(output_path_docx)
    print(f"Relatório DOCX gerado com sucesso em: {output_path_docx}")

    print("Iniciando conversão para PDF...")
    pdf_output_path = output_path_docx.replace(".docx", ".pdf")
    convert(output_path_docx, pdf_output_path)
    print(f"Relatório PDF gerado com sucesso em: {pdf_output_path}")

except Exception as e:
    print(f"Ocorreu um erro ao salvar o relatório: {e}")
# %%
