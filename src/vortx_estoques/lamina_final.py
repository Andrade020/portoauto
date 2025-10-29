# %%

#!    --    -->     TODO:    <--    --     !#
# mes no nome do relatorio
# remover 2 gráficos iniciais
# nome do fundo
# 

# =============================================================================
# CÉLULA 1: BIBLIOTECAS E CONFIGURAÇÕES GERAIS
# =============================================================================

import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import matplotlib.ticker as ticker
import os
import glob
from datetime import datetime
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert
from IPython.display import display
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re
from dateutil.parser import parse
import time
# Adicione estes imports junto com os outros
import matplotlib.dates as mdates
import locale

# Adicione esta linha após os imports para configurar o idioma
locale.setlocale(locale.LC_TIME, 'pt_BR.UTF-8')

# --- PARÂMETROS DE ENTRADA E SAÍDA ---
# ! ATENÇÃO: Configure os caminhos e nomes de arquivos aqui.

# FUNDO PRINCIPAL PARA ANÁLISE
FUNDO_ALVO = 'FIDC FCT II SR2'

# CAMINHOS DOS DADOS DE ENTRADA
PATH_RENTABILIDADE = r"C:\Users\Leo\Desktop\Porto_Real\portoauto\src\vortx_estoques\data\Rentabilidades\151265-Rentabilidade_Sintetica.csv"
# ### ALTERAÇÃO ### - Caminho do estoque atualizado para a pasta principal.
PATH_ESTOQUE = r"C:\Users\Leo\Desktop\Porto_Real\portoauto\src\vortx_estoques\data\Estoques"
PATH_TEMPLATE_DOCX = r"C:\Users\Leo\Desktop\Porto_Real\portoreal\notebooks\template.docx"
PATH_LOGO = r"C:\Users\Leo\Desktop\Porto_Real\portoauto\images\logo.png"
PATH_ANBIMA_LOGO = r"C:\Users\Leo\Desktop\Porto_Real\portoauto\images\anbima.png"
PATH_LOGO_FICTOR = r"C:\Users\Leo\Desktop\Porto_Real\portoauto\images\logo_fictor.png"
PATH_LOGO_VORTX = r"C:\Users\Leo\Desktop\Porto_Real\portoauto\images\logo_vortx.png"

# CAMINHOS DE SAÍDA
PASTA_SAIDA_IMAGENS = r"C:\Users\Leo\Desktop\Porto_Real\portoauto\src\vortx_estoques\output\lamina_imagens"
PASTA_SAIDA_RELATORIOS = r"C:\Users\Leo\Desktop\Porto_Real\portoauto\src\vortx_estoques\output"

# --- CONFIGURAÇÕES VISUAIS ---
pd.options.display.max_rows = 100
pd.options.display.max_columns = 50

# PALETA DE CORES
COLOR_PALETTE = {
    'primary_light': '#76C6C5',
    'primary_dark': '#0E5D5F',
    'secondary': '#163F3F',
    'header_bg': '#F1F9F9',
    'header_text': '#0E5D5F'
}

# Cria as pastas de saída se não existirem
os.makedirs(PASTA_SAIDA_IMAGENS, exist_ok=True)
os.makedirs(PASTA_SAIDA_RELATORIOS, exist_ok=True)
print(f"Imagens serão salvas em: {PASTA_SAIDA_IMAGENS}")
print(f"Relatórios serão salvos em: {PASTA_SAIDA_RELATORIOS}")


# %%
# =============================================================================
# CÉLULA 2: FUNÇÕES DE PROCESSAMENTO DE DADOS
# =============================================================================

def carregar_dados_rentabilidade(caminho_arquivo):
    """Lê, limpa e pré-processa o arquivo de rentabilidade sintética."""
    print("Carregando dados de rentabilidade...")
    try:
        colunas = [
            'carteira', 'nomeFundo', 'tipoCarteira', 'administrador', 'periodo',
            'emissao', 'data', 'valorCota', 'quantidade', 'numeroCotistas',
            'variacaoDia', 'variacaoPeriodo', 'captacoes', 'resgate', 'eventos',
            'pl', 'coluna_extra'
        ]
        df = pd.read_csv(
            caminho_arquivo, sep=';', encoding='latin1', header=None, names=colunas, skiprows=1
        )
        df.columns = df.columns.str.strip()

        colunas_numericas = ['valorCota', 'pl', 'quantidade', 'numeroCotistas', 'variacaoDia']
        for col in colunas_numericas:
            if col in df.columns:
                df[col] = df[col].astype(str).str.replace('.', '', regex=False).str.replace(',', '.', regex=False)
                df[col] = pd.to_numeric(df[col], errors='coerce')

        df['data'] = pd.to_datetime(df['data'], format='%d/%m/%Y', errors='coerce')
        df.dropna(subset=['data', 'valorCota'], inplace=True)
        df['nomeFundo'] = df['nomeFundo'].str.strip()

        print("Dados de rentabilidade carregados e limpos com sucesso.")
        return df
    except FileNotFoundError:
        print(f"ERRO: Arquivo de rentabilidade não encontrado em: {caminho_arquivo}")
        return pd.DataFrame()

def processar_dados_estoque(caminho_pasta):
    """Lê o conjunto de arquivos de estoque mais recente e calcula o status de vencimento."""
    print("\nCarregando dados do novo estoque...")
    todos_arquivos = glob.glob(os.path.join(caminho_pasta, '**/*.csv'), recursive=True)
    if not todos_arquivos:
        print(f"AVISO: Nenhum arquivo CSV de estoque encontrado em {caminho_pasta}")
        return pd.DataFrame()

    data_recente = None
    regex_data = re.compile(r'(\d{2}\.\d{2}\.\d{2})')

    for arquivo in todos_arquivos:
        match = regex_data.search(os.path.basename(arquivo))
        if match:
            data_obj = datetime.strptime(match.group(1), '%d.%m.%y')
            if data_recente is None or data_obj > data_recente:
                data_recente = data_obj

    if data_recente is None:
        print("AVISO: Não foi possível determinar a data do conjunto de estoque mais recente.")
        return pd.DataFrame()

    data_recente_str = data_recente.strftime('%d.%m.%y')
    print(f"Usando o conjunto de estoque da data (data de referência): {data_recente_str}")

    arquivos_para_ler = [f for f in todos_arquivos if data_recente_str in os.path.basename(f)]
    lista_dfs_partes = []
    for arquivo in arquivos_para_ler:
        try:
            df_parte = pd.read_csv(arquivo, sep=';', decimal=',', encoding='utf-8', on_bad_lines='warn')
            lista_dfs_partes.append(df_parte)
        except Exception as e:
            print(f"Aviso: Não foi possível ler o arquivo {arquivo}. Erro: {e}")
            continue

    if not lista_dfs_partes:
        print("Nenhum dado de estoque foi carregado com sucesso.")
        return pd.DataFrame()

    df_consolidado = pd.concat(lista_dfs_partes, ignore_index=True)

    # --- INÍCIO DA CORREÇÃO ---
    # 1. Converte a coluna original 'DataVencimento' (sem espaço) para datetime.
    df_consolidado['DataVencimento'] = pd.to_datetime(df_consolidado['DataVencimento'], format='%d/%m/%Y', errors='coerce')

    # 2. Converte a data de referência para Timestamp do pandas e a normaliza (zera o tempo).
    #    Esta linha corrige o erro 'AttributeError'.
    data_ref_estoque = pd.Timestamp(data_recente).normalize()

    # 3. Cria a coluna 'Situacao' com base na comparação das datas.
    df_consolidado['Situacao'] = np.where(
        df_consolidado['DataVencimento'] > data_ref_estoque,
        'A vencer',
        'Vencido'
    )
    # --- FIM DA CORREÇÃO ---

    # 4. Agora, renomeia as colunas para o padrão esperado pelo resto do script (com espaços).
    df_consolidado.rename(columns={
        'DataVencimento': 'Data Vencimento',
        'ValorPresente': 'Valor Presente',
        'ValorNominal': 'Valor Nominal'
    }, inplace=True)

    # Converte colunas de valores numéricos
    df_consolidado['Valor Presente'] = pd.to_numeric(df_consolidado['Valor Presente'], errors='coerce')
    df_consolidado['Valor Nominal'] = pd.to_numeric(df_consolidado['Valor Nominal'], errors='coerce')

    # Remove linhas onde a conversão de data ou valor falhou
    df_consolidado.dropna(subset=['Data Vencimento', 'Valor Presente'], inplace=True)

    print(f"Dados do novo estoque consolidados: {len(df_consolidado)} linhas.")
    print("Distribuição de Status calculada:")
    print(df_consolidado['Situacao'].value_counts())
    
    return df_consolidado

def obter_dados_bacen(codigo_bcb, data_inicio='28/03/2017', tentativas=3, espera=3):
    """Busca uma série temporal da API do Banco Central do Brasil, com lógica de retentativa."""
    url = f'https://api.bcb.gov.br/dados/serie/bcdata.sgs.{codigo_bcb}/dados?formato=json&dataInicial={data_inicio}'

    for tentativa in range(1, tentativas + 1):
        try:
            print(f"Tentando buscar dados do Bacen (Série {codigo_bcb}), tentativa {tentativa}/{tentativas}...")
            df = pd.read_json(url)
            df['data'] = pd.to_datetime(df['data'], dayfirst=True)
            df.set_index('data', inplace=True)
            df['valor'] = pd.to_numeric(df['valor'], errors='coerce')
            print(f"Série {codigo_bcb} do Bacen carregada com sucesso!")
            return df
        except Exception as e:
            print(f"ERRO na tentativa {tentativa}: {e}")
            if tentativa < tentativas:
                print(f"Aguardando {espera} segundos antes de tentar novamente...")
                time.sleep(espera)
            else:
                print(f"ERRO FINAL: Não foi possível buscar dados do Bacen (Série {codigo_bcb}) após {tentativas} tentativas.")

    return pd.DataFrame()

def preparar_df_retornos(df_rentabilidade, df_cota_com_cdi):
    """Cria o DataFrame de retornos diários a partir da 'variacaoDia' e adiciona o CDI."""
    df_retornos = df_rentabilidade.pivot_table(
        index='data',
        columns='nomeFundo',
        values='variacaoDia'
    ) / 100
    if 'CDI' in df_cota_com_cdi.index:
        df_retornos['CDI'] = df_cota_com_cdi.loc['CDI'].pct_change()
    return df_retornos


# %%
# =============================================================================
# CÉLULA 3: FUNÇÕES DE GERAÇÃO DE GRÁFICOS E TABELAS
# =============================================================================
# Adicione estas bibliotecas no início do seu script, caso não as tenha
# Adicione estas bibliotecas no início do seu script, caso não as tenha
import numpy as np
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
import matplotlib.ticker as ticker
import os

def plotar_grafico_subordinacao(df_patrimonio, pasta_saida):
    """
    Calcula e plota o histórico do índice de subordinação do fundo.
    MODIFICADO: Adicionada anotação de "Limite Mínimo" na linha pontilhada.
    """
    print("Gerando gráfico de subordinação...")
    try:
        # ... (código de cálculo) ...
        pl_subordinado = df_patrimonio.loc['FIDC FCT II']
        pl_senior = df_patrimonio.loc['FIDC FCT II SR2']
        pl_total = pl_senior + pl_subordinado
        indice_subordinacao = (pl_subordinado / pl_total.replace(0, np.nan)).dropna() * 100
        if indice_subordinacao.empty: return None

        # --- PLOTAGEM ---
        fig, ax = plt.subplots(figsize=(8, 4.5))
        ax.plot(indice_subordinacao.index, indice_subordinacao.values, color=COLOR_PALETTE['primary_dark'], lw=2.5)
        ax.axhline(y=20, color='gray', linestyle='--', linewidth=1.2, alpha=0.7)
        
        # --- Formatação do Gráfico ---
        ax.axhline(0, color='lightgray', linestyle='-', linewidth=1, zorder=0)
        ax.spines[['top', 'right', 'left', 'bottom']].set_visible(False)
        ax.tick_params(
            axis='both', which='both', length=0, 
            labelsize=10, labelcolor='#333333'
        )
        ax.xaxis.set_major_locator(mdates.AutoDateLocator(minticks=5, maxticks=7))
        ax.xaxis.set_major_formatter(mdates.DateFormatter('%b/%y'))
        fig.autofmt_xdate()
        ax.yaxis.set_major_formatter(ticker.PercentFormatter(xmax=100.0))
        ax.set_ylim(bottom=0)
        
        # --- Rótulos e Anotações ---
        ultimo_valor = indice_subordinacao.iloc[-1]
        ax.text(indice_subordinacao.index[-1], ultimo_valor, f' {ultimo_valor:.2f}%'.replace('.',','), 
                color=COLOR_PALETTE['primary_dark'], fontsize=12, va='center', ha='right')
        
        # ALTERAÇÃO: Anotação para a linha de limite mínimo
        ax.annotate(
            'Limite Mínimo',
            xy=(indice_subordinacao.index[-1], 20), # Posição da seta
            xytext=(8, 0),                          # Deslocamento do texto
            textcoords='offset points',
            ha='left',
            va='center',
            fontsize=9,
            color='gray'
        )

        path_saida_grafico = os.path.join(pasta_saida, "grafico_subordinacao.png")
        fig.savefig(path_saida_grafico, dpi=300, bbox_inches='tight')
        plt.close(fig)
        
        print(f"Gráfico de subordinação (versão final) salvo em: {path_saida_grafico}")
        return path_saida_grafico

    except KeyError as e:
        print(f"AVISO: Não foi possível gerar o gráfico de subordinação. Cota não encontrada: {e}")
        return None
    except Exception as e:
        print(f"ERRO inesperado ao gerar gráfico de subordinação: {e}")
        return None


def _render_mpl_table(data, ax, col_width=1.5, row_height=0.625, font_size=10):
    """Função auxiliar para renderizar uma tabela Matplotlib com estilo."""
    header_color = COLOR_PALETTE['header_bg']
    header_font_color = COLOR_PALETTE['header_text']
    edge_color = 'w'
    row_colors = ['#FFFFFF', '#FFFFFF']

    ax.axis('off')
    mpl_table = ax.table(cellText=data.values, cellLoc='center', bbox=[0, 0, 1, 1],
                         colLabels=data.columns, loc='center', colLoc='center')
    mpl_table.auto_set_font_size(False)
    mpl_table.set_fontsize(font_size)

    for k, cell in mpl_table._cells.items():
        cell.set_edgecolor(edge_color)
        cell.set_linewidth(0)
        if k[0] == 0:  # Linha de cabeçalho
            cell.set_text_props(weight='bold', color=header_font_color)
            cell.set_facecolor(header_color)
        else:
            if (k[0] - 1) % 3 == 2:
                 cell.set_facecolor('#F0F0F0') # Linha % CDI
            else:
                 cell.set_facecolor(row_colors[k[0] % len(row_colors)])

        if k[1] <= 0:  # Colunas de cabeçalho (Ano, Cota)
            cell.set_text_props(weight='bold', color=header_font_color)
            cell.set_facecolor(header_color)
    return ax

def criar_tabela_performance(df_rentabilidade_processada, fundos_a_exibir, df_cota_base, benchmark='CDI'):
    """
    Gera a tabela de performance com a lógica de comparação do benchmark (CDI) corrigida
    para sempre usar o mesmo período de dias do fundo.
    """
    # --- ETAPA 1: PREPARAR DADOS DIÁRIOS ---
    df_calc = df_rentabilidade_processada[['data', 'nomeFundo', 'variacaoDia']].copy()
    df_calc['retorno_diario'] = df_calc['variacaoDia'] / 100
    
    cdi_disponivel = benchmark in df_cota_base.index
    retornos_cdi_diarios = pd.Series(dtype=float)
    if cdi_disponivel:
        retornos_cdi_diarios = df_cota_base.loc[benchmark].pct_change().dropna()
    else:
        print(f"AVISO: Benchmark '{benchmark}' não encontrado. Tabela será gerada sem essa comparação.")

    # --- ETAPA 2: MONTAR A ESTRUTURA DA TABELA ---
    all_rows = []
    years = sorted(df_calc['data'].dt.year.dropna().unique())

    def format_pct(x):
        return f'{x:.2%}'.replace('.', ',') if isinstance(x, (int, float)) and pd.notna(x) else '-'
    def format_pct_of_bench(x):
        return f'{x*100:.1f}%'.replace('.', ',') if isinstance(x, (int, float)) and pd.notna(x) else '-'

    for year in years:
        for nome_fundo, nome_exibicao in fundos_a_exibir.items():
            df_fundo_ano = df_calc[(df_calc['nomeFundo'] == nome_fundo) & (df_calc['data'].dt.year == year)]
            if df_fundo_ano.empty:
                continue

            row_fundo = [year, nome_exibicao]
            if cdi_disponivel:
                row_pct_bench = [year, f"% {benchmark}"]

            # --- CÁLCULO MENSAL CORRIGIDO ---
            for month in range(1, 13):
                df_fundo_mes = df_fundo_ano[df_fundo_ano['data'].dt.month == month]
                
                if df_fundo_mes.empty:
                    row_fundo.append('-')
                    if cdi_disponivel: row_pct_bench.append('-')
                else:
                    # Calcula o retorno do fundo para o mês (parcial ou completo)
                    ret_fundo = (1 + df_fundo_mes['retorno_diario']).prod() - 1
                    row_fundo.append(format_pct(ret_fundo))

                    if cdi_disponivel:
                        # Calcula o retorno do CDI nos mesmos dias do fundo
                        datas_do_fundo_no_mes = df_fundo_mes['data']
                        ret_bench = (1 + retornos_cdi_diarios.reindex(datas_do_fundo_no_mes).dropna()).prod() - 1
                        pct_cdi = ret_fundo / ret_bench if ret_bench != 0 else pd.NA
                        row_pct_bench.append(format_pct_of_bench(pct_cdi))

            # --- CÁLCULO YTD (ANO) CORRIGIDO ---
            ret_fundo_ano = (1 + df_fundo_ano['retorno_diario']).prod() - 1
            ret_bench_ano = pd.NA
            if cdi_disponivel:
                datas_do_fundo_no_ano = df_fundo_ano['data']
                ret_bench_ano = (1 + retornos_cdi_diarios.reindex(datas_do_fundo_no_ano).dropna()).prod() - 1
            
            # --- CÁLCULO DESDE O INÍCIO CORRIGIDO ---
            df_fundo_total = df_calc[(df_calc['nomeFundo'] == nome_fundo) & (df_calc['data'].dt.year <= year)]
            ret_fundo_inicio = (1 + df_fundo_total['retorno_diario']).prod() - 1
            ret_bench_inicio = pd.NA
            if cdi_disponivel:
                datas_do_fundo_total = df_fundo_total['data']
                ret_bench_inicio = (1 + retornos_cdi_diarios.reindex(datas_do_fundo_total).dropna()).prod() - 1
            
            # Adiciona os valores calculados na linha
            row_fundo.extend([format_pct(ret_fundo_ano), format_pct(ret_fundo_inicio)])
            all_rows.append(row_fundo)
            
            if cdi_disponivel:
                pct_cdi_ano = ret_fundo_ano / ret_bench_ano if ret_bench_ano != 0 else pd.NA
                pct_cdi_inicio = ret_fundo_inicio / ret_bench_inicio if ret_bench_inicio != 0 else pd.NA
                row_pct_bench.extend([format_pct_of_bench(pct_cdi_ano), format_pct_of_bench(pct_cdi_inicio)])
                all_rows.append(row_pct_bench)

    # --- ETAPA 3: RENDERIZAR E SALVAR A TABELA (sem alterações) ---
    if not all_rows: return None
    
    table_cols = ['Ano', 'Cota', 'Jan', 'Fev', 'Mar', 'Abr', 'Maio', 'Jun', 'Jul', 'Ago', 'Set', 'Out', 'Nov', 'Dez', 'YTD', 'Desde o\nInício']
    df_table = pd.DataFrame(all_rows, columns=table_cols)
    
    num_rows_per_year = len(all_rows) // len(years)
    anos_formatados = [row['Ano'] if i % num_rows_per_year == 0 else '' for i, row in df_table.iterrows()]
    df_table['Ano'] = anos_formatados

    def render_mpl_table_final(data, ax, font_size=8.5):
        header_color = COLOR_PALETTE['header_bg']; header_font_color = COLOR_PALETTE['header_text']
        font_color = '#4F4F4F'; font_name = 'Gill Sans MT'
        ax.axis('off')
        mpl_table = ax.table(cellText=data.values, cellLoc='center', bbox=[0, 0, 1, 1], colLabels=data.columns, loc='center', colLoc='center')
        mpl_table.auto_set_font_size(False); mpl_table.set_fontsize(font_size)
        for k, cell in mpl_table._cells.items():
            cell.set_edgecolor('w'); cell.set_linewidth(0)
            props = {'fontfamily': font_name}
            if k[0] == 0:
                props.update({'weight': 'bold', 'color': header_font_color}); cell.set_facecolor(header_color)
            else:
                props.update({'color': font_color})
                is_pct_cdi_row = cdi_disponivel and (k[0] % 2 == 0)
                if is_pct_cdi_row: cell.set_facecolor('#F0F0F0')
                else: cell.set_facecolor('#FFFFFF')
            if k[1] <= 0:
                props.update({'weight': 'bold', 'color': header_font_color})
                if k[0] != 0: cell.set_facecolor(header_color)
            cell.set_text_props(**props)
        return ax

    fig, ax = plt.subplots(figsize=(12, len(all_rows) * 0.35))
    render_mpl_table_final(df_table, ax)
    
    # Usa o primeiro fundo da lista para nomear o arquivo
    primeiro_fundo_nome = list(fundos_a_exibir.keys())[0].replace(' ', '_')
    path_tabela = os.path.join(PASTA_SAIDA_IMAGENS, f"tabela_rentabilidade_{primeiro_fundo_nome}.png")
    fig.savefig(path_tabela, dpi=300, bbox_inches='tight', pad_inches=0.1)
    plt.close(fig)
    print(f"Tabela de rentabilidade (lógica do CDI corrigida) salva em: {path_tabela}")
    return path_tabela



# Adicione estas bibliotecas no início do seu script, caso não as tenha
import matplotlib.dates as mdates
import matplotlib.ticker as ticker

def plotar_graficos_rentabilidade(df_cota, df_patr, df_retornos, fundos_a_exibir, benchmark='CDI'):
    """
    Cria e salva todos os gráficos de performance.
    MODIFICADO: Ajustes finos nos eixos e rótulos do gráfico de retorno.
    """
    # ... (código inicial da função) ...
    fundo_alvo_principal = list(fundos_a_exibir.keys())[0]
    nome_exib_principal = list(fundos_a_exibir.values())[0]
    nome_arquivo_limpo = fundo_alvo_principal.replace(' ', '_')

    # --- 1. Gráfico de Retorno Acumulado ---
    fig, ax = plt.subplots(figsize=(8, 4.5))
    
    # ... (código de cálculo das linhas) ...
    inception_dates = []
    fundos_para_loop = list(fundos_a_exibir.keys())
    for f in fundos_para_loop:
        first_valid_index = df_retornos[f].dropna().index.min()
        inception_dates.append(first_valid_index)
    common_start_date = max(inception_dates)
    cor_fundo = COLOR_PALETTE['primary_light']

    # Plotagem do Fundo
    for fundo_nome, fundo_label in fundos_a_exibir.items():
        # ... (cálculo do retorno do fundo) ...
        retornos_diarios = df_retornos[fundo_nome].dropna()[common_start_date:]
        if retornos_diarios.empty: continue
        df_fund_evol_pct = ((1 + retornos_diarios).cumprod() - 1) * 100
        ax.plot(df_fund_evol_pct, color=cor_fundo, lw=2.5, label=fundo_label)
        last_val = df_fund_evol_pct.iloc[-1]
        
        # ALTERAÇÃO: Troca ponto por vírgula no rótulo
        label_text_br = f' +{last_val:.2f}%'.replace('.', ',')
        ax.text(df_fund_evol_pct.index[-1], last_val, label_text_br, 
                color=cor_fundo, fontsize=12, va='center')

    # Plotagem do Benchmark
    if benchmark in df_retornos.columns:
        # ... (cálculo do retorno do benchmark) ...
        retornos_cdi = df_retornos[benchmark].dropna()[common_start_date:]
        if not retornos_cdi.empty:
            df_bench_evol_pct = ((1 + retornos_cdi).cumprod() - 1) * 100
            ax.plot(df_bench_evol_pct, color='k', linestyle='--', alpha=0.7, lw=2, label=benchmark)
            last_val_bench = df_bench_evol_pct.iloc[-1]
            
            # ALTERAÇÃO: Troca ponto por vírgula no rótulo
            label_text_bench_br = f' +{last_val_bench:.2f}%'.replace('.', ',')
            ax.text(df_bench_evol_pct.index[-1], last_val_bench, label_text_bench_br, 
                    color='k', alpha=0.8, fontsize=12, va='center')
    
    # --- Formatação do Gráfico ---
    ax.axhline(0, color='lightgray', linestyle='-', linewidth=1, zorder=0)
    ax.spines[['top', 'right', 'left', 'bottom']].set_visible(False)
    ax.tick_params(
        axis='both', which='both', length=0, 
        labelsize=10, labelcolor='#333333'
    )
    
    # ALTERAÇÃO: Remove as casas decimais do eixo Y
    ax.yaxis.set_major_formatter(ticker.PercentFormatter(xmax=100, decimals=0))
    
    # ALTERAÇÃO: Define intervalos de 2 em 2 meses no eixo X
    ax.xaxis.set_major_locator(mdates.MonthLocator(interval=2)) 
    ax.xaxis.set_major_formatter(mdates.DateFormatter('%b/%y'))
    fig.autofmt_xdate()
    
    ax.legend(loc='upper left', frameon=False, fontsize=10)
    plt.tight_layout()
    path_retorno = os.path.join(PASTA_SAIDA_IMAGENS, f"grafico_retorno_{nome_arquivo_limpo}.png")
    fig.savefig(path_retorno, dpi=300, bbox_inches='tight'); plt.close(fig)
    print(f"Gráfico de retorno (ajustes finos) salvo em: {path_retorno}")

    # --- 2. Gráfico de Volatilidade (sem alterações) ---
    fig, ax = plt.subplots(figsize=(12, 6))
    window=22
    ret_fundo = df_cota.loc[fundo_alvo_principal].pct_change()
    vol_fundo = ret_fundo.rolling(window=window).std() * np.sqrt(252)
    ax.plot(vol_fundo, color=COLOR_PALETTE['primary_light'], lw=2, label=nome_exib_principal)
    if benchmark in df_cota.index:
        ret_bench = df_cota.loc[benchmark].pct_change()
        vol_bench = ret_bench.rolling(window=window).std() * np.sqrt(252)
        ax.plot(vol_bench, color='k', linestyle='--', alpha=0.7, lw=2, label=benchmark)
    ax.spines[['top', 'right']].set_visible(False); ax.yaxis.set_major_formatter(ticker.PercentFormatter(1.0))
    ax.grid(axis='y', linestyle='--', alpha=0.7)
    ax.set_title(f'Volatilidade Anualizada (Janela Móvel de {window}d) - {nome_exib_principal}', fontsize=16, pad=20)
    ax.legend(loc='upper left', frameon=False, fontsize=12); plt.tight_layout()
    path_vol = os.path.join(PASTA_SAIDA_IMAGENS, f"grafico_volatilidade_{nome_arquivo_limpo}.png")
    fig.savefig(path_vol, dpi=300, bbox_inches='tight'); plt.close(fig)
    print(f"Gráfico de volatilidade salvo em: {path_vol}")

    # --- 3. Gráfico de Drawdown (sem alterações) ---
    fig, ax = plt.subplots(figsize=(12, 6))
    df_fund_dd = df_cota.loc[fundo_alvo_principal].dropna()
    daily_dd = df_fund_dd / df_fund_dd.cummax() - 1.0
    ax.plot(daily_dd, color=COLOR_PALETTE['primary_light'], lw=1)
    ax.fill_between(daily_dd.index, daily_dd, 0, color=COLOR_PALETTE['primary_light'], alpha=0.3)
    ax.spines[['top', 'right']].set_visible(False); ax.yaxis.set_major_formatter(ticker.PercentFormatter(1.0))
    ax.grid(axis='y', linestyle='--', alpha=0.7)
    ax.set_title(f'Drawdown Histórico - {nome_exib_principal}', fontsize=16, pad=20); plt.tight_layout()
    path_dd = os.path.join(PASTA_SAIDA_IMAGENS, f"grafico_drawdown_{nome_arquivo_limpo}.png")
    fig.savefig(path_dd, dpi=300, bbox_inches='tight'); plt.close(fig)
    print(f"Gráfico de drawdown salvo em: {path_dd}")

    # --- 4. Gráfico de Evolução do PL (sem alterações) ---
    fig, ax = plt.subplots(figsize=(12, 6))
    df_plot_pl = df_patr.loc[fundo_alvo_principal].dropna()
    ax.plot(df_plot_pl.index, df_plot_pl.values, color=COLOR_PALETTE['primary_light'], lw=2.5, label=nome_exib_principal)
    ax.fill_between(df_plot_pl.index, df_plot_pl.values, 0, color=COLOR_PALETTE['primary_light'], alpha=0.3)
    ax.spines[['top', 'right']].set_visible(False)
    ax.yaxis.set_major_formatter(ticker.FuncFormatter(lambda x, p: f'R$ {x/1e6:.1f}M'))
    ax.grid(axis='y', linestyle='--', alpha=0.7)
    ax.set_title(f'Evolução do Patrimônio Líquido (PL) - {nome_exib_principal}', fontsize=16, pad=20)
    ax.legend(loc='upper left', frameon=False, fontsize=12)
    pl_final = df_plot_pl.iloc[-1]
    ax.text(df_plot_pl.index[-1], pl_final, f' R$ {pl_final/1e6:.2f}M',
            color=COLOR_PALETTE['secondary'], fontsize=12, va='center', ha='left')
    plt.tight_layout()
    path_pl = os.path.join(PASTA_SAIDA_IMAGENS, f"grafico_pl_{nome_arquivo_limpo}.png")
    fig.savefig(path_pl, dpi=300, bbox_inches='tight'); plt.close(fig)
    print(f"Gráfico de PL (apenas Sênior) salvo em: {path_pl}")
    
    return {
        'retorno': path_retorno, 'volatilidade': path_vol,
        'drawdown': path_dd, 'pl': path_pl
    }

def plotar_graficos_estoque(df_estoque, data_referencia):
    """Cria e salva todos os gráficos relacionados ao estoque."""
    if df_estoque.empty:
        print("DataFrame de estoque vazio. Pulando geração de gráficos de estoque.")
        return {}

    print("\n--- Gerando imagens de Estoque ---")
    # ### ALTERAÇÃO ### - Usa a coluna 'Situacao' do novo DataFrame.
    df_avencer = df_estoque[df_estoque['Situacao'] == 'A vencer'].copy()
    paths = {}

    # Gerar e salvar cada gráfico de estoque
    paths['venc_mensal'] = plotar_vencimento_mensal(df_avencer, 'Valor Presente', 'VCTO_MES', data_referencia)
    paths['venc_anual'] = plotar_vencimento_anual(df_avencer, 'Valor Presente', 'VCTO_ANO')
    # ### ALTERAÇÃO ### - Gráficos de UF e CAPAG removidos.
    paths['aging_vencidos'] = plotar_aging_com_acumulado(df_estoque, data_referencia, 'Valor Presente')
    # ### ALTERAÇÃO ### - Usa a coluna 'SacadoCnpjCpf'.
    paths['conc_cumulativa_sacado'] = plotar_concentracao_cumulativa_sacado(df_avencer, 'Valor Presente', 'SacadoCnpjCpf')

    print("\nGráficos de estoque gerados com sucesso.")
    return paths

def plotar_vencimento_mensal(df_aberto, col_valor, col_vcto, data_referencia, cutoff=10):
    """Gera e salva o gráfico de vencimento mensal a partir da data de referência."""
    data_filtro = pd.Timestamp(data_referencia).normalize()
    df_futuro = df_aberto[df_aberto['Data Vencimento'] >= data_filtro].copy()

    if df_futuro.empty:
        print("Não há dados de vencimentos futuros para gerar o gráfico mensal.")
        return None

    df_futuro[col_vcto] = df_futuro['Data Vencimento'].dt.strftime('%Y-%m')
    df2 = df_futuro.groupby(col_vcto)[col_valor].sum()
    df3 = (df2 / df2.sum()).cumsum()
    df2 = df2[:cutoff]
    df3 = df3[:cutoff]

    ymin1, ymax1 = 0, df2.max() * 2
    ymax2 = 1.3
    ymin2 = df3.min() * 2.5 - 1.5 * ymax2

    fig, ax = plt.subplots(figsize=(8, 4.5))
    ax2 = ax.twinx()

    bars = ax.bar(df2.index, np.clip(df2.values, 0.01 * df2.max(), df2.max()), color=COLOR_PALETTE['primary_dark'])
    ax2.plot(df3.index, df3.values, color=COLOR_PALETTE['secondary'], lw=2, marker='o')

    month_map = {'01':'jan','02':'fev','03':'mar','04':'abr','05':'mai','06':'jun','07':'jul','08':'ago','09':'set','10':'out','11':'nov','12':'dez'}
    ax.set_xticks(df2.index)
    ax.set_xticklabels([f"{month_map[x[-2:]]}/{x[2:4]}" for x in df2.index])
    ax.set_ylim(ymin1, ymax1); ax2.set_ylim(ymin2, ymax2)
    ax.set_yticks([]); ax2.set_yticks([])
    ax.tick_params(bottom=False)
    ax.set_frame_on(False); ax2.spines[:].set_visible(False)

    max_height = max([bar.get_height() for bar in bars])
    for i, bar in enumerate(bars):
        h = bar.get_height() + max_height * 0.04
        value = df2.values[i] / 1000
        ax.text(bar.get_x() + bar.get_width()/2, h, f'{value:,.0f}'.replace(',', '.'), ha='center')

    for x, y in df3.items():
        value = f'{y:.0%}' if ((y < .990) or (y == 1)) else '>99%'
        ax2.text(x, y + (ymax2 - ymin2) * 0.04, value, ha='center')

    path_saida = os.path.join(PASTA_SAIDA_IMAGENS, "vencimento_mensal.png")
    fig.savefig(path_saida, bbox_inches='tight')
    plt.close(fig)
    print(f"Gráfico de vencimento mensal salvo em: {path_saida}")
    return path_saida

def plotar_vencimento_anual(df_aberto, col_valor, col_vcto):
    """Gera e salva o gráfico de vencimento anual."""
    df_aberto = df_aberto.copy()
    df_aberto[col_vcto] = df_aberto['Data Vencimento'].dt.year
    df2 = df_aberto.groupby(col_vcto)[col_valor].sum()
    df3 = df2.cumsum() / df2.sum()

    if df2.empty:
        print("Não há dados para gerar o gráfico de vencimento anual.")
        return None

    ymin1, ymax1 = 0, df2.max() * 2
    ymax2 = 1.15
    ymin2 = df3.min() * 2.5 - 1.5 * ymax2

    fig, ax = plt.subplots(figsize=(8, 4.5))
    ax2 = ax.twinx()

    bars = ax.bar(df2.index, np.clip(df2.values, 0.01 * df2.max(), df2.max()), color=COLOR_PALETTE['primary_dark'])
    ax2.plot(df3.index, df3.values, color=COLOR_PALETTE['secondary'], lw=2, marker='o')

    ax.set_xticks(df2.index.tolist())
    ax.set_ylim(ymin1, ymax1); ax2.set_ylim(ymin2, ymax2)
    ax.set_yticks([]); ax2.set_yticks([])
    ax.tick_params(bottom=False)
    ax.set_frame_on(False); ax2.spines[:].set_visible(False)

    max_height = max([bar.get_height() for bar in bars])
    for i, bar in enumerate(bars):
        h = bar.get_height() + max_height * 0.04
        value = df2.values[i] / 1000
        ax.text(bar.get_x() + bar.get_width()/2, h, f'{value:,.0f}'.replace(',', '.'), ha='center')

    for x, y in df3.items():
        value = f'{y:.0%}' if ((y < .990) or np.isclose(y, 1)) else '>99%'
        ax2.text(x, y + 0.06, value, ha='center')

    path_saida = os.path.join(PASTA_SAIDA_IMAGENS, "vencimento_anual.png")
    fig.savefig(path_saida, bbox_inches='tight')
    plt.close(fig)
    print(f"Gráfico de vencimento anual salvo em: {path_saida}")
    return path_saida

def plotar_concentracao_cumulativa_sacado(df_aberto, col_valor='Valor Presente', col_class='SacadoCnpjCpf'):
    """Gera um gráfico da concentração CUMULativa de valor por devedores."""
    df_agrupado = df_aberto.groupby(col_class)[col_valor].sum().sort_values(ascending=False)

    if df_agrupado.empty:
        print(f"Não há dados para gerar o gráfico de concentração cumulativa por '{col_class}'.")
        return None

    valor_total = df_agrupado.sum()
    num_sacados = len(df_agrupado)

    bins = [10, 20, 50, 100, 200, 500, 1000, 2000, 5000, 10000]
    resultados = {}
    for n in bins:
        if n < num_sacados:
            soma_top_n = df_agrupado.iloc[0:n].sum()
            pct_top_n = (soma_top_n / valor_total) * 100
            resultados[f'{n} maiores'] = pct_top_n

    resultados[f'{num_sacados} maiores'] = 100.0
    df_plot = pd.Series(resultados)

    fig, ax = plt.subplots(figsize=(8, 4.5)) # Mantém o mesmo figsize do gráfico de Aging
    bars = ax.barh(df_plot.index, df_plot.values, color=COLOR_PALETTE['primary_dark'])

    ax.invert_yaxis()
    ax.set_xticks([])
    ax.set_frame_on(False)
    ax.tick_params(left=False)
    
    # ### INÍCIO DA ALTERAÇÃO ###
    # 1. Fixa o limite do eixo X para padronizar o tamanho da imagem salva
    #    Isso garante espaço para os rótulos sem redimensionar a imagem.
    ax.set_xlim(0, 115) 

    # Adiciona os rótulos de percentual no final das barras
    for bar in bars:
        width = bar.get_width()
        ax.text(width + 2, bar.get_y() + bar.get_height()/2, f'{width:.0f}%',
                va='center', ha='left', fontsize=9)

    # 2. Remove o título da imagem, pois ele já existe no relatório.
    # ax.set_title("Concentração Cumulativa por Sacado", fontsize=12) # Linha removida
    # ### FIM DA ALTERAÇÃO ###

    path_saida = os.path.join(PASTA_SAIDA_IMAGENS, "concentracao_cumulativa_sacado.png")
    fig.savefig(path_saida, bbox_inches='tight')
    plt.close(fig)
    print(f"Gráfico de concentração cumulativa por sacado (sem título e com tamanho ajustado) salvo em: {path_saida}")
    return path_saida


def plotar_aging_com_acumulado(df_estoque, data_referencia, col_valor='Valor Presente'):
    """Gera um gráfico de Aging com linha de percentual acumulado."""
    df_vencidos = df_estoque[df_estoque['Situacao'] == 'Vencido'].copy()

    if df_vencidos.empty:
        print("Não há dados de carteira vencida para gerar o gráfico de Aging.")
        return None

    data_filtro = pd.Timestamp(data_referencia).normalize()
    df_vencidos['dias_atraso'] = (data_filtro - df_vencidos['Data Vencimento']).dt.days

    bins = [-1, 14, 30, 60, 90, 120, 150, 180, np.inf]
    labels = ['0-14 D', '15-30 D', '31-60 D', '61-90 D', '91-120 D', '121-150 D', '151-180 D', '180+ D']
    df_vencidos['Faixa Atraso'] = pd.cut(df_vencidos['dias_atraso'], bins=bins, labels=labels, right=True)

    df_aging = df_vencidos.groupby('Faixa Atraso', observed=False)[col_valor].sum()
    total_em_atraso = df_aging.sum()

    if total_em_atraso == 0:
        print("Valor total em atraso é zero. Pulando gráfico de Aging.")
        return None

    df_aging_pct_line = (df_aging.cumsum() / total_em_atraso)
    df_aging_bars = df_aging.copy()
    df_aging_bars['Total'] = total_em_atraso

    fig, ax = plt.subplots(figsize=(8, 4.5))
    ax2 = ax.twinx()

    bars = ax.bar(df_aging_bars.index, df_aging_bars.values, color=COLOR_PALETTE['primary_dark'], width=0.6)
    bars[-1].set_color(COLOR_PALETTE['secondary'])

    ax_ymax = df_aging_bars.max() * 1.40
    ax.set_ylim(0, ax_ymax)

    max_bar_height_ratio = (df_aging.max() / ax_ymax) * 1.15
    line_visual_start = min(max_bar_height_ratio + 0.1, 0.7)
    line_visual_end = 0.95
    line_data_start = df_aging_pct_line.iloc[0] if not df_aging_pct_line.empty else 0
    line_data_end = 1.0

    if abs(line_visual_end - line_visual_start) < 0.01:
        line_visual_start, line_visual_end = 0.6, 0.95

    y2_range = (line_data_end - line_data_start) / (line_visual_end - line_visual_start)
    y2_min = line_data_start - (y2_range * line_visual_start)
    y2_max = y2_min + y2_range
    ax2.set_ylim(y2_min, y2_max)

    ax2.plot(df_aging_pct_line.index, df_aging_pct_line.values, color=COLOR_PALETTE['secondary'], marker='o', lw=2)

    ax.set_yticks([])
    ax.set_frame_on(False)

    offset_barras = ax.get_ylim()[1] * 0.02
    for bar in bars:
        height = bar.get_height()
        # ### INÍCIO DA ALTERAÇÃO ###
        # Altera a unidade de Milhões (1e6) para Milhares (1e3) e formata como inteiro.
        label_text = f'{height/1e3:,.0f}'.replace(',', '.')
        # ### FIM DA ALTERAÇÃO ###
        ax.text(bar.get_x() + bar.get_width() / 2, height + offset_barras, label_text,
                ha='center', va='bottom', fontsize=9, color='#333333')

    ax2.spines[:].set_visible(False)
    ax2.set_yticks([])

    y2_axis_range = ax2.get_ylim()[1] - ax2.get_ylim()[0]
    label_offset = y2_axis_range * 0.05
    for i, pct in enumerate(df_aging_pct_line):
        ax2.text(i, pct + label_offset, f'{pct:.0%}', ha='center', fontsize=9, color=COLOR_PALETTE['secondary'])

    path_saida = os.path.join(PASTA_SAIDA_IMAGENS, "aging_com_acumulado.png")
    fig.savefig(path_saida, bbox_inches='tight')
    plt.close(fig)
    print(f"Gráfico de Aging (em milhares) salvo em: {path_saida}")
    return path_saida
# %%
# =============================================================================
# CÉLULA 4: FUNÇÕES DE MONTAGEM DO RELATÓRIO (DOCX)
# =============================================================================

def set_paragraph_horizontal_border(paragraph, color="76C6C5", size="4", space="1"):
    """Adiciona bordas na parte SUPERIOR e INFERIOR de um parágrafo."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for border_side in ['top', 'bottom']:
        border_element = OxmlElement(f'w:{border_side}')
        border_element.set(qn('w:val'), 'single')
        border_element.set(qn('w:sz'), size)
        border_element.set(qn('w:space'), space)
        border_element.set(qn('w:color'), color)
        pBdr.append(border_element)
    pPr.append(pBdr)

def configurar_cabecalho_rodape(doc, nome_relatorio, data_referencia, path_logo):
    """Configura o cabeçalho e as margens do documento."""
    section = doc.sections[0]
    section.top_margin, section.bottom_margin = Cm(1), Cm(1)
    section.left_margin, section.right_margin = Cm(1.6), Cm(1)

    header = section.header
    header_table = header.tables[0] if header.tables else header.add_table(rows=1, cols=2, width=Cm(18))

    cell_left = header_table.cell(0, 0)
    p_left = cell_left.paragraphs[0]; p_left.text = ""
    run_name = p_left.add_run(nome_relatorio.upper())
    run_name.font.name = "Gill Sans MT"
    run_name.font.size, run_name.bold = Pt(14), True
    run_name.font.color.rgb = RGBColor(16, 112, 130)
    p_left.add_run("\n")
    meses_pt = ["Janeiro", "Fevereiro", "Março", "Abril", "Maio", "Junho", "Julho", "Agosto", "Setembro", "Outubro", "Novembro", "Dezembro"]
    data_formatada = f"{meses_pt[data_referencia.month - 1]} {data_referencia.year}"
    run_month = p_left.add_run(data_formatada)
    run_month.font.name, run_month.font.size, run_month.italic = "Arial", Pt(12), True
    run_month.font.color.rgb = RGBColor(38, 38, 38)

    cell_right = header_table.cell(0, 1)
    p_right = cell_right.paragraphs[0]; p_right.text = ""
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_right.add_run().add_picture(path_logo, width=Cm(4.93))

def montar_corpo_relatorio(doc, paths_imagens):
    """Adiciona a tabela e os gráficos ao corpo do documento."""
    
    def add_section_title(cell, title, subtitle=""):
        p = cell.add_paragraph()
        p.paragraph_format.space_before, p.paragraph_format.space_after = Pt(6), Pt(2)
        run = p.add_run(title)
        run.font.name, run.font.size, run.bold = "Gill Sans MT", Pt(10), False
        run.font.color.rgb = RGBColor(89, 89, 89)
        if subtitle:
            p.add_run("\n")
            run_sub = p.add_run(subtitle)
            run_sub.font.name = "Gill Sans MT"
            run_sub.font.size = Pt(8)
            run_sub.font.color.rgb = RGBColor(89, 89, 89)
        set_paragraph_horizontal_border(p, color="76C6C5", space="2")

    # Tabela de Rentabilidade
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(12)
    run = title_p.add_run("\tRentabilidade Mensal")
    run.font.name, run.font.size, run.bold = "Gill Sans MT", Pt(11), False
    set_paragraph_horizontal_border(title_p, color="76C6C5", space="4")
    
    if paths_imagens.get('tabela_perf') and os.path.exists(paths_imagens['tabela_perf']):
        doc.add_picture(paths_imagens['tabela_perf'], width=Cm(17.8))

    # ### INÍCIO DA ALTERAÇÃO ###
    # Tabela de Gráficos agora com 4 linhas para acomodar o novo gráfico
    graphs_table = doc.add_table(rows=4, cols=2)
    # ### FIM DA ALTERAÇÃO ###
    graphs_table.autofit = False
    for col in graphs_table.columns:
        col.width = Cm(8.9)
    img_width = Cm(8.5)

    # ### INÍCIO DA ALTERAÇÃO ###
    # Mapa de gráficos atualizado com o novo "Índice de Subordinação"
    mapa_graficos = {
        (0, 0): ("Retorno Acumulado", "", 'retorno'),
        (0, 1): ("Evolução do Patrimônio Líquido (PL)", "(R$ MM)", 'pl'),
        (1, 0): ("Índice de Subordinação", "(% do PL Total)", 'subordinacao'), # NOVO GRÁFICO
        (1, 1): ("Volatilidade Anualizada", "(Janela de 22 dias)", 'volatilidade'),
        (2, 0): ("Vencimento Mensal", "(R$’000 e % acumulado)", 'venc_mensal'),
        (2, 1): ("Vencimento Anual", "(R$’000 e % acumulado)", 'venc_anual'),
        (3, 0): ("Concentração Cumulativa por Sacado", "(%)", 'conc_cumulativa_sacado'),
        (3, 1): ("Aging da Carteira Vencida", "(R$’000 e % acumulado)", 'aging_vencidos'),
    }
    # ### FIM DA ALTERAÇÃO ###
    
    for (row, col), (titulo, subtitulo, path_key) in mapa_graficos.items():
        path = paths_imagens.get(path_key)
        
        if path and os.path.exists(path):
            cell = graphs_table.cell(row, col)
            add_section_title(cell, f"\t{titulo}", f"\t{subtitulo}")
            p_img = cell.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.add_run().add_picture(path, width=img_width)

def add_page_field(run, field_type):
    """Adiciona um campo de numeração de página (PAGE ou NUMPAGES) a um 'run'."""
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve'); instrText.text = field_type
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)

def _set_cell_border(cell, **kwargs):
    """Função auxiliar para definir bordas de uma célula da tabela."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("start", "top", "end", "bottom"):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = f"w:{edge}"
            border = tcBorders.find(qn(tag))
            if border is None:
                border = OxmlElement(tag)
                tcBorders.append(border)
            for k, v in edge_data.items():
                border.set(qn(f"w:{k}"), str(v))

def adicionar_secao_logos(doc):
    """Adiciona a seção de logos de Gestão e Administração."""
    doc.add_paragraph()
    logo_data = [
        ("Gestão", PATH_LOGO),
        ("Gestão", PATH_LOGO_FICTOR),
        ("Administração", PATH_LOGO_VORTX)
    ]
    table = doc.add_table(rows=2, cols=3)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    border_style = {"sz": 4, "val": "single", "color": "76C6C5"}

    for i, (title, logo_path) in enumerate(logo_data):
        cell_title = table.cell(0, i)
        p_title = cell_title.paragraphs[0]
        p_title.text = title
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_title.runs[0]
        run.font.name = 'Arial'; run.font.size = Pt(10); run.font.color.rgb = RGBColor.from_string("595959")
        _set_cell_border(cell_title, top=border_style, bottom=border_style)

        cell_logo = table.cell(1, i)
        cell_logo.text = ''
        p_logo = cell_logo.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.add_run().add_picture(logo_path, height=Cm(1.2))
        _set_cell_border(cell_logo, top={"val": "nil"}, bottom={"val": "nil"}, start={"val": "nil"}, end={"val": "nil"})

def adicionar_secao_final(doc, df_patr, anbima_path, date_ref):
    """Adiciona a seção final do relatório (dados, disclaimer, rodapé)."""
    from dateutil.relativedelta import relativedelta
    
    p_dados_title = doc.add_paragraph()
    run_dados_title = p_dados_title.add_run("\tDados do Fundo")
    run_dados_title.font.name = "Gill Sans MT"
    run_dados_title.font.size = Pt(10)
    run_dados_title.font.color.rgb = RGBColor(89, 89, 89)
    # ### ALTERAÇÃO ### - Cor da borda alterada para azul da paleta.
    set_paragraph_horizontal_border(p_dados_title, "76C6C5", space="4")
    p_dados_title.paragraph_format.space_after = Pt(5)
    p_dados_title.paragraph_format.space_before = Pt(25)

    data_parent_tbl = doc.add_table(rows=1, cols=2)
    p_dados_title._element.addnext(data_parent_tbl._element)

    next_el = p_dados_title._element.getnext()
    if next_el is not None and next_el.tag.endswith("p") and not next_el.text:
        p_dados_title._p.getparent().remove(next_el)

    # ### INÍCIO DA ALTERAÇÃO ###
    # O cálculo de PL volta a somar todas as cotas para representar o PL TOTAL do fundo.
    if df_patr.empty or len(df_patr.columns) == 0:
        pl_ultimo, pl_12m = 0, 0
    else:
        all_dates = pd.to_datetime(df_patr.columns)
        # Soma o PL de todas as cotas na data de referência
        pl_ultimo = df_patr[date_ref].sum()

        date_start_12m = date_ref - relativedelta(years=1)
        cols_12m = all_dates[(all_dates > date_start_12m) & (all_dates <= date_ref)]
        
        # Soma o PL de cada dia e depois calcula a média desses totais diários
        pl_12m = df_patr[cols_12m].sum(axis=0).mean() if not cols_12m.empty else 0
    # ### FIM DA ALTERAÇÃO ###

    pl_ultimo_formatted = f"{pl_ultimo:,.0f}".replace(",", ".")
    pl_12m_formatted = f"{pl_12m:,.0f}".replace(",", ".")

    left_data_info = [
        "Política de Investimento\nO fundo investe em direitos creditórios oriundos empréstimos consignados concedidos de servidores públicos, aposentados e pensionistas.",
        "Denominação Social\nFICTOR CONSIGNADO II FUNDO DE INVESTIMENTO EM DIREITOS CREDITÓRIOS RESPONSABILIDADE LIMITADA",
        "CNPJ\n52.203.615/0001-19", "Data de início\n28/09/2023",
        "Rentabilidade alvo da cota Sênior\nCDI + 3%", "Forma de condomínio\nFechado",
        f"Patrimônio Líquido (Ref: {date_ref.strftime('%d/%m/%Y')})\nR$ {pl_ultimo_formatted}",
        f"Patrimônio Líquido Médio (12M)\nR$ {pl_12m_formatted}"
    ]
    right_data_info = [
        "Gestores\nPorto Real Investimentos\nFictor Asset", "Administrador\nVórtx DTVM",
        "Custodiante\nVórtx DTVM", "Taxa de gestão\n2,0%", "Taxa de administração\n0,18%",
        "Taxa de escrituração\nR$ 2.000,00 por mês", "Taxa máxima de custódia\nIncluída na taxa de gestão",
        "Taxa de performance\nNão há"
    ]

    def fill_data_cell(text_str, target_cell):
        p = target_cell.paragraphs[0] if target_cell.paragraphs else target_cell.add_paragraph(); p.clear()
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.2
        identifier, value = text_str.split("\n", 1)
        run_id = p.add_run(identifier.strip()); run_id.font.name = "Arial"; run_id.font.size = Pt(8.5); run_id.font.color.rgb = RGBColor(89, 89, 89); run_id.bold = True; run_id.add_break()
        value_lines = value.strip().split('\n')
        for i, line in enumerate(value_lines):
            run_val = p.add_run(line.strip()); run_val.font.name = "Arial"; run_val.font.size = Pt(8.5); run_val.font.color.rgb = RGBColor(89, 89, 89)
            if i < len(value_lines) - 1: run_val.add_break()

    left_tbl = data_parent_tbl.cell(0, 0).add_table(rows=len(left_data_info), cols=1)
    right_tbl = data_parent_tbl.cell(0, 1).add_table(rows=len(right_data_info), cols=1)
    for i, text in enumerate(left_data_info): fill_data_cell(text, left_tbl.cell(i, 0))
    for i, text in enumerate(right_data_info): fill_data_cell(text, right_tbl.cell(i, 0))

    adicionar_secao_logos(doc)
        
    p_disc_title = doc.add_paragraph()
    run_disc_title = p_disc_title.add_run("\t Disclaimer")
    run_disc_title.font.name = "Gill Sans MT"; run_disc_title.font.size = Pt(10)
    run_disc_title.font.color.rgb = RGBColor(89, 89, 89)
    # ### ALTERAÇÃO ### - Cor da borda alterada para azul da paleta.
    set_paragraph_horizontal_border(p_disc_title, "76C6C5")
    p_disc_body = doc.add_paragraph()
    p_disc_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run_disc_body = p_disc_body.add_run("As informações contidas neste material são de caráter exclusivamente informativo. A rentabilidade passada não representa garantia de rentabilidade futura. Fundos de investimento não contam com a garantia do administrador, do gestor da carteira, de qualquer mecanismo de seguro ou, ainda, do Fundo Garantidor de Créditos ‐ FGC. Os fundos de crédito privado estão sujeitos a risco de perda substancial de seu patrimônio líquido em caso de eventos que acarretem o não pagamento dos ativos integrantes de sua carteira, inclusive por força de intervenção, liquidação, regime de administração temporária, falência, recuperação judicial ou extrajudicial dos emissores responsáveis pelos ativos do fundo. Não há garantia de que este fundo terá o tratamento tributário para fundos de longo prazo. A rentabilidade divulgada não é líquida de impostos e taxa.")
    run_disc_body.font.name = "Arial"; run_disc_body.font.size = Pt(8.5); run_disc_body.font.color.rgb = RGBColor(166, 166, 166)

    footer = doc.sections[0].footer
    footer_table = footer.tables[0] if footer.tables else footer.add_table(rows=1, cols=3, width=Cm(18))
    cell_left = footer_table.cell(0, 0); cell_left.width = Cm(4); p_left = cell_left.paragraphs[0]; p_left.clear(); p_left.add_run().add_picture(anbima_path, height=Cm(1.0))
    cell_center = footer_table.cell(0, 1); cell_center.width = Cm(10); p_center = cell_center.paragraphs[0]; p_center.clear(); p_center.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_contato = p_center.add_run("PORTO REAL ASSET\nAv. São Gabriel 301 – 2º andar – Itaim Bibi – São Paulo (SP)\n"); run_contato.font.name = "Arial"; run_contato.font.size = Pt(9); run_contato.font.color.rgb = RGBColor(127, 127, 127)
    run_link = p_center.add_run("contato@portorealasset.com.br"); run_link.font.name = "Arial"; run_link.font.size = Pt(9); run_link.font.color.rgb = RGBColor(0, 0, 0); run_link.font.underline = True
    cell_right = footer_table.cell(0, 2); cell_right.width = Cm(4); p_right = cell_right.paragraphs[0]; p_right.clear(); p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_priv = p_right.add_run("Documento estritamente privado e confidencial"); run_priv.font.name = "Arial"; run_priv.font.size = Pt(8); run_priv.font.color.rgb = RGBColor(89, 89, 89)
    p_page = cell_right.add_paragraph(); p_page.alignment = WD_ALIGN_PARAGRAPH.RIGHT; font_page = {"name": "Arial", "size": Pt(8), "color": RGBColor(89, 89, 89)}
    run = p_page.add_run("Página "); run.font.name=font_page['name']; run.font.size=font_page['size']; run.font.color.rgb=font_page['color']
    run = p_page.add_run(); add_page_field(run, "PAGE"); run.font.name=font_page['name']; run.font.size=font_page['size']; run.font.color.rgb=font_page['color']
    run = p_page.add_run(" de "); run.font.name=font_page['name']; run.font.size=font_page['size']; run.font.color.rgb=font_page['color']
    run = p_page.add_run(); add_page_field(run, "NUMPAGES"); run.font.name=font_page['name']; run.font.size=font_page['size']; run.font.color.rgb=font_page['color']
# %%
# =============================================================================
# CÉLULA 5: ORQUESTRADOR PRINCIPAL - EXECUÇÃO DO FLUXO
# =============================================================================
def gerar_relatorio_completo():
    """Função principal que orquestra todo o processo."""
    from dateutil.relativedelta import relativedelta

    # --- ETAPA 1: PROCESSAMENTO DE DADOS DE RENTABILIDADE ---
    df_rentabilidade = carregar_dados_rentabilidade(PATH_RENTABILIDADE)
    if df_rentabilidade.empty:
        print("Processo interrompido: falha ao carregar dados de rentabilidade.")
        return

    df_cota = df_rentabilidade.pivot_table(index='data', columns='nomeFundo', values='valorCota').ffill().T
    df_patr = df_rentabilidade.pivot_table(index='data', columns='nomeFundo', values='pl').ffill().T

    df_cdi = obter_dados_bacen(12)
    if not df_cdi.empty:
        df_cota.loc['CDI'] = (1 + df_cdi['valor'] / 100).cumprod().ffill()

    df_retornos = preparar_df_retornos(df_rentabilidade, df_cota)

    if df_patr.empty or len(df_patr.columns) == 0:
        data_ref = datetime.now()
    else:
        all_dates = pd.to_datetime(df_patr.columns)
        data_ref = all_dates[-1]
    
    print(f"\nData de referência para o relatório definida como: {data_ref.strftime('%d/%m/%Y')}")

    # --- ETAPA 2: GERAÇÃO DAS IMAGENS ---
    print("\n--- Gerando imagens de Rentabilidade ---")
    paths_imagens = {}
    
    fundos_para_analise = {
        'FIDC FCT II SR2': 'Sênior'
    }

    paths_imagens['tabela_perf'] = criar_tabela_performance(df_rentabilidade, fundos_para_analise, df_cota)
    paths_rentabilidade = plotar_graficos_rentabilidade(df_cota, df_patr, df_retornos, fundos_para_analise)
    paths_imagens.update(paths_rentabilidade)
    
    # ### INÍCIO DA ALTERAÇÃO ###
    # Chamada para a nova função do gráfico de subordinação
    paths_imagens['subordinacao'] = plotar_grafico_subordinacao(df_patr, PASTA_SAIDA_IMAGENS)
    # ### FIM DA ALTERAÇÃO ###
    
    df_estoque = processar_dados_estoque(PATH_ESTOQUE)
    paths_estoque = plotar_graficos_estoque(df_estoque, data_ref)
    paths_imagens.update(paths_estoque)
    
    # --- ETAPA 3: MONTAGEM E SALVAMENTO DO RELATÓRIO ---
    print("\n--- Montando relatório ---")
    doc = Document(PATH_TEMPLATE_DOCX)
    
    configurar_cabecalho_rodape(doc, 'Lâmina de Performance e Estoque', data_ref, PATH_LOGO)
    montar_corpo_relatorio(doc, paths_imagens)
    adicionar_secao_final(doc, df_patr, PATH_ANBIMA_LOGO, data_ref)
    
    nome_arquivo = f"Lamina_Completa_{data_ref.strftime('%Y-%m-%d')}.docx"
    path_docx = os.path.join(PASTA_SAIDA_RELATORIOS, nome_arquivo)
    
    try:
        doc.save(path_docx)
        print(f"Relatório DOCX gerado com sucesso em: {path_docx}")

        try:
            print("Iniciando conversão para PDF...")
            path_pdf = path_docx.replace(".docx", ".pdf")
            convert(path_docx, path_pdf)
            print(f"Relatório PDF gerado com sucesso em: {path_pdf}")
        except Exception as e:
            print("\n--- ERRO NA CONVERSÃO PARA PDF ---")
            print(f"Ocorreu um erro ao tentar converter o DOCX para PDF. Mensagem: {e}")
            print("\nPASSOS PARA RESOLVER:")
            print("1. (MAIS COMUM) Abra o Gerenciador de Tarefas (Ctrl+Shift+Esc), vá em 'Detalhes' e finalize qualquer processo 'WINWORD.EXE'.")
            print("2. Abra o Microsoft Word sozinho e veja se aparece alguma janela pedindo sua atenção (ex: login, recuperação de arquivos).")
            print(f"3. Tente abrir o arquivo '{os.path.basename(path_docx)}' e salvá-lo como PDF manualmente para ver se o Word consegue.")
            print("4. Se nada funcionar, repare a sua instalação do Microsoft Office.")

    except Exception as e:
        print(f"Ocorreu um erro CRÍTICO ao gerar o arquivo DOCX: {e}")


if __name__ == "__main__":
    gerar_relatorio_completo()