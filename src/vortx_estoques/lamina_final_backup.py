# %%
# =============================================================================
# CÉLULA 1: BIBLIOTECAS E CONFIGURAÇÕES GERAIS
# =============================================================================
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import matplotlib.ticker as ticker
import os
import glob
from datetime import datetime
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert
from IPython.display import display
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from dateutil.parser import parse

#pip install python-dateutil


# --- PARÂMETROS DE ENTRADA E SAÍDA ---
# ! ATENÇÃO: Configure os caminhos e nomes de arquivos aqui.

# FUNDO PRINCIPAL PARA ANÁLISE
FUNDO_ALVO = 'FIDC FCT II SR2'

# CAMINHOS DOS DADOS DE ENTRADA
PATH_RENTABILIDADE = r"C:\Users\Leo\Desktop\Porto_Real\portoauto\src\vortx_estoques\data\135972-Rentabilidade_Sintetica.csv"
PATH_ESTOQUE = r"C:\Users\Leo\Desktop\Porto_Real\portoauto\src\vortx_estoques\data\estoque_consolidado_agosto"
PATH_TEMPLATE_DOCX = r"C:\Users\Leo\Desktop\Porto_Real\portoreal\notebooks\template.docx"
PATH_LOGO = r"C:\Users\Leo\Desktop\Porto_Real\portoauto\images\logo.png"
PATH_ANBIMA_LOGO = r"C:\Users\Leo\Desktop\Porto_Real\portoauto\images\anbima.png" 


# CAMINHOS DE SAÍDA
PASTA_SAIDA_IMAGENS = r"C:\Users\Leo\Desktop\Porto_Real\portoauto\src\vortx_estoques\output\lamina_imagens"
PASTA_SAIDA_RELATORIOS = r"C:\Users\Leo\Desktop\Porto_Real\portoauto\src\vortx_estoques\output"

# --- CONFIGURAÇÕES VISUAIS ---
pd.options.display.max_rows = 100
pd.options.display.max_columns = 50

# PALETA DE CORES
COLOR_PALETTE = {
    'primary_light': '#76C6C5',
    'primary_dark': '#0E5D5F',
    'secondary': '#163F3F',
    'header_bg': '#F1F9F9',
    'header_text': '#0E5D5F'
}

# Cria as pastas de saída se não existirem
os.makedirs(PASTA_SAIDA_IMAGENS, exist_ok=True)
os.makedirs(PASTA_SAIDA_RELATORIOS, exist_ok=True)
print(f"Imagens serão salvas em: {PASTA_SAIDA_IMAGENS}")
print(f"Relatórios serão salvos em: {PASTA_SAIDA_RELATORIOS}")


# %%
# =============================================================================
# CÉLULA 2: FUNÇÕES DE PROCESSAMENTO DE DADOS
# =============================================================================

def carregar_dados_rentabilidade(caminho_arquivo):
    """Lê, limpa e pré-processa o arquivo de rentabilidade sintética."""
    print("Carregando dados de rentabilidade...")
    try:
        colunas = [
            'carteira', 'nomeFundo', 'tipoCarteira', 'administrador', 'periodo',
            'emissao', 'data', 'valorCota', 'quantidade', 'numeroCotistas',
            'variacaoDia', 'variacaoPeriodo', 'captacoes', 'resgate', 'eventos',
            'pl', 'coluna_extra'
        ]
        df = pd.read_csv(
            caminho_arquivo, sep=';', encoding='latin1', header=None, names=colunas, skiprows=1
        )
        df.columns = df.columns.str.strip()

        colunas_numericas = ['valorCota', 'pl', 'quantidade', 'numeroCotistas', 'variacaoDia']
        for col in colunas_numericas:
            if col in df.columns:
                df[col] = df[col].astype(str).str.replace('.', '', regex=False).str.replace(',', '.', regex=False)
                df[col] = pd.to_numeric(df[col], errors='coerce')

        df['data'] = pd.to_datetime(df['data'], format='%d/%m/%Y', errors='coerce')
        df.dropna(subset=['data', 'valorCota'], inplace=True)
        df['nomeFundo'] = df['nomeFundo'].str.strip()

        print("Dados de rentabilidade carregados e limpos com sucesso.")
        return df
    except FileNotFoundError:
        print(f"ERRO: Arquivo de rentabilidade não encontrado em: {caminho_arquivo}")
        return pd.DataFrame()

def processar_dados_estoque(caminho_pasta):
    """Lê e consolida os arquivos de estoque de uma pasta."""
    print("\nCarregando dados de estoque...")
    arquivos_csv = glob.glob(os.path.join(caminho_pasta, "*.csv"))
    if not arquivos_csv:
        print(f"AVISO: Nenhum arquivo CSV de estoque encontrado em {caminho_pasta}")
        return pd.DataFrame()

    all_dfs = []
    float_cols = [
        'Valor Aquisicao', 'Valor Nominal', 'Valor Presente', 'PDD Vencido',
        'PDD Total', 'Taxa Operada Originador', 'CET Mensal', 'Taxa CCB',
        'Taxa Originador Split', 'Taxa Split FIDC'
    ]
    date_cols = ['Data Aquisicao', 'Data Vencimento', 'Data Referencia', 'Data de Nascimento']

    for file in arquivos_csv:
        try:
            df = pd.read_csv(file, encoding='utf-16', sep='\t', engine='python', on_bad_lines='warn', header=0)
            df.columns = df.columns.str.strip()

            if not df.empty:
                for col in float_cols:
                    if col in df.columns:
                        df[col] = df[col].astype(str).str.replace(',', '.').astype(float)
                # SUBSTITUA O LOOP ANTIGO POR ESTE AQUI:
                for col in date_cols:
                    if col in df.columns:
                        # Tenta converter cada valor na coluna de data, lidando com nulos (NaN, NaT)
                        # A mágica está no `parse(x, dayfirst=True)`
                        df[col] = df[col].apply(
                            lambda x: parse(str(x), dayfirst=True) if pd.notna(x) else pd.NaT
                        )
                all_dfs.append(df)
        except Exception as e:
            print(f"Erro ao processar o arquivo de estoque {os.path.basename(file)}: {e}")
            continue

    if not all_dfs:
        print("Nenhum dado de estoque foi carregado com sucesso.")
        return pd.DataFrame()

    df_final = pd.concat(all_dfs, ignore_index=True)
    print("Dados de estoque consolidados com sucesso.")
    return df_final

def obter_dados_bacen(codigo_bcb, data_inicio='28/03/2017'):
    """Busca uma série temporal da API do Banco Central do Brasil."""
    try:
        url = f'https://api.bcb.gov.br/dados/serie/bcdata.sgs.{codigo_bcb}/dados?formato=json&dataInicial={data_inicio}'
        df = pd.read_json(url)
        df['data'] = pd.to_datetime(df['data'], dayfirst=True)
        df.set_index('data', inplace=True)
        df['valor'] = pd.to_numeric(df['valor'], errors='coerce')
        print(f"Série {codigo_bcb} do Bacen carregada com sucesso!")
        return df
    except Exception as e:
        print(f"ERRO ao buscar dados do Bacen (Série {codigo_bcb}): {e}")
        return pd.DataFrame()

def preparar_df_retornos(df_rentabilidade, df_cota_com_cdi):
    """Cria o DataFrame de retornos diários a partir da 'variacaoDia' e adiciona o CDI."""
    df_retornos = df_rentabilidade.pivot_table(
        index='data',
        columns='nomeFundo',
        values='variacaoDia'
    ) / 100
    if 'CDI' in df_cota_com_cdi.index:
        df_retornos['CDI'] = df_cota_com_cdi.loc['CDI'].pct_change()
    return df_retornos


# %%
# =============================================================================
# CÉLULA 3: FUNÇÕES DE GERAÇÃO DE GRÁFICOS E TABELAS
# =============================================================================

def _render_mpl_table(data, ax, col_width=1.5, row_height=0.625, font_size=10):
    """Função auxiliar para renderizar uma tabela Matplotlib com estilo."""
    header_color = COLOR_PALETTE['header_bg']
    header_font_color = COLOR_PALETTE['header_text']
    edge_color = 'w'
    row_colors = ['#FFFFFF', '#FFFFFF']
    
    ax.axis('off')
    mpl_table = ax.table(cellText=data.values, cellLoc='center', bbox=[0, 0, 1, 1],
                         colLabels=data.columns, loc='center', colLoc='center')
    mpl_table.auto_set_font_size(False)
    mpl_table.set_fontsize(font_size)

    for k, cell in mpl_table._cells.items():
        cell.set_edgecolor(edge_color)
        cell.set_linewidth(0)
        if k[0] == 0:  # Linha de cabeçalho
            cell.set_text_props(weight='bold', color=header_font_color)
            cell.set_facecolor(header_color)
        else:
            if (k[0] - 1) % 3 == 2:
                 cell.set_facecolor('#F0F0F0') # Linha % CDI
            else:
                 cell.set_facecolor(row_colors[k[0] % len(row_colors)])

        if k[1] <= 0:  # Colunas de cabeçalho (Ano, Cota)
            cell.set_text_props(weight='bold', color=header_font_color)
            cell.set_facecolor(header_color)
    return ax

def criar_tabela_performance(df_rentabilidade_processada, fundos_a_exibir, df_cota_base, benchmark='CDI'):
    """
    Gera a tabela de performance calculando a rentabilidade mensal a partir da coluna 'variacaoDia',
    agrupada por ano e com estética aprimorada.
    """
    # --- ETAPA 1: CALCULAR PERFORMANCE MENSAL A PARTIR DA 'variacaoDia' ---
    df_calc = df_rentabilidade_processada[['data', 'nomeFundo', 'variacaoDia']].copy()
    df_calc['mes'] = df_calc['data'].dt.to_period('M')
    
    # Calcula o fator diário usando a coluna oficial, exatamente como no seu script de exemplo
    df_calc['fator_diario'] = 1 + (df_calc['variacaoDia'] / 100)
    
    # Agrupa por fundo e mês, e calcula o produto dos fatores diários (rentabilidade geométrica)
    fator_mensal = df_calc.groupby(['nomeFundo', 'mes'])['fator_diario'].prod()
    perf_mensal_fundos = (fator_mensal - 1).unstack('nomeFundo').T

    # Calcula a performance mensal para o benchmark (CDI)
    retornos_cdi = df_cota_base.loc[benchmark].pct_change().dropna()
    perf_mensal_cdi = (retornos_cdi + 1).resample('M').prod() - 1
    
    # --- ETAPA 2: MONTAR A ESTRUTURA DA TABELA ---
    all_rows = []
    years = sorted(df_calc['data'].dt.year.unique())

    def format_pct(x):
        return f'{x:.2%}'.replace('.', ',') if isinstance(x, (int, float)) and pd.notna(x) else '-'
    def format_pct_of_bench(x):
        return f'{x*100:.1f}%'.replace('.', ',') if isinstance(x, (int, float)) and pd.notna(x) else '-'

    for year in years:
        for nome_fundo, nome_exibicao in fundos_a_exibir.items():
            if nome_fundo not in perf_mensal_fundos.index:
                continue

            row_fundo = [year, nome_exibicao]
            row_pct_bench = [year, f"% {benchmark}"]
            
            # Rentabilidade YTD (Ano)
            start_of_year = pd.Timestamp(f'{year}-01-01')
            ret_fundo_ano = (perf_mensal_fundos.loc[nome_fundo][(perf_mensal_fundos.columns.year == year)] + 1).prod() - 1
            ret_bench_ano = (perf_mensal_cdi[perf_mensal_cdi.index.year == year] + 1).prod() - 1
            
            # Rentabilidade Acumulada (Desde o Início)
            start_date = df_cota_base.loc[nome_fundo].dropna().index[0]
            ret_fundo_inicio = (perf_mensal_fundos.loc[nome_fundo][perf_mensal_fundos.columns.to_timestamp() >= start_date] + 1).prod() - 1
            ret_bench_inicio = (perf_mensal_cdi[perf_mensal_cdi.index >= start_date] + 1).prod() - 1

            for month in range(1, 13):
                mes_period = pd.Period(f'{year}-{month}', freq='M')
                if mes_period in perf_mensal_fundos.columns:
                    ret_fundo = perf_mensal_fundos.loc[nome_fundo, mes_period]
                    ret_bench = perf_mensal_cdi.get(mes_period.to_timestamp('M'), pd.NA)
                    pct_cdi = ret_fundo / ret_bench if pd.notna(ret_fundo) and pd.notna(ret_bench) and ret_bench != 0 else '-'
                    row_fundo.append(format_pct(ret_fundo))
                    row_pct_bench.append(format_pct_of_bench(pct_cdi))
                else:
                    row_fundo.append('-')
                    row_pct_bench.append('-')

            pct_cdi_ano = ret_fundo_ano / ret_bench_ano if ret_bench_ano != 0 else '-'
            pct_cdi_inicio = ret_fundo_inicio / ret_bench_inicio if ret_bench_inicio != 0 else '-'
            
            row_fundo.extend([format_pct(ret_fundo_ano), format_pct(ret_fundo_inicio)])
            row_pct_bench.extend([format_pct_of_bench(pct_cdi_ano), format_pct_of_bench(pct_cdi_inicio)])
            all_rows.extend([row_fundo, row_pct_bench])

    # --- ETAPA 3: RENDERIZAR E SALVAR A TABELA (sem alterações na estética) ---
    if not all_rows: return None
    
    table_cols = ['Ano', 'Cota', 'Jan', 'Fev', 'Mar', 'Abr', 'Maio', 'Jun', 'Jul', 'Ago', 'Set', 'Out', 'Nov', 'Dez', 'YTD', 'Desde o\nInício']
    df_table = pd.DataFrame(all_rows, columns=table_cols)
    
    num_rows_per_year = len(fundos_a_exibir) * 2
    anos_formatados = [row['Ano'] if i % num_rows_per_year == 0 else '' for i, row in df_table.iterrows()]
    df_table['Ano'] = anos_formatados

    def render_mpl_table_final(data, ax, font_size=8.5):
        # (Esta função interna de renderização permanece a mesma da versão anterior)
        header_color = COLOR_PALETTE['header_bg']; header_font_color = COLOR_PALETTE['header_text']
        font_color = '#4F4F4F'; font_name = 'Gill Sans MT'
        ax.axis('off')
        mpl_table = ax.table(cellText=data.values, cellLoc='center', bbox=[0, 0, 1, 1], colLabels=data.columns, loc='center', colLoc='center')
        mpl_table.auto_set_font_size(False); mpl_table.set_fontsize(font_size)
        for k, cell in mpl_table._cells.items():
            cell.set_edgecolor('w'); cell.set_linewidth(0)
            props = {'fontfamily': font_name}
            if k[0] == 0:
                props.update({'weight': 'bold', 'color': header_font_color}); cell.set_facecolor(header_color)
            else:
                props.update({'color': font_color})
                if k[0] % 2 == 0: cell.set_facecolor('#F0F0F0')
                else: cell.set_facecolor('#FFFFFF')
            if k[1] <= 0:
                props.update({'weight': 'bold', 'color': header_font_color})
                if k[0] != 0: cell.set_facecolor(header_color)
            cell.set_text_props(**props)
        return ax

    fig, ax = plt.subplots(figsize=(12, len(all_rows) * 0.35))
    render_mpl_table_final(df_table, ax)
    
    path_tabela = os.path.join(PASTA_SAIDA_IMAGENS, f"tabela_rentabilidade_{list(fundos_a_exibir.keys())[0].replace(' ', '_')}.png")
    fig.savefig(path_tabela, dpi=300, bbox_inches='tight', pad_inches=0.1)
    plt.close(fig)
    print(f"Tabela de rentabilidade (cálculo via variacaoDia) salva em: {path_tabela}")
    return path_tabela

def plotar_graficos_rentabilidade(df_cota, df_patr, df_retornos, fundos_a_exibir, benchmark='CDI'):
    """
    Cria e salva todos os gráficos de performance.
    MODIFICADO: As linhas de retorno acumulado agora começam na mesma data (a mais recente entre os fundos).
    """
    fundo_alvo_principal = list(fundos_a_exibir.keys())[0]
    nome_exib_principal = list(fundos_a_exibir.values())[0]
    nome_arquivo_limpo = fundo_alvo_principal.replace(' ', '_')

    # --- 1. Gráfico de Retorno Acumulado (LÓGICA DE INÍCIO ATUALIZADA) ---
    fig, ax = plt.subplots(figsize=(12, 6))
    
    # --- MUDANÇA PRINCIPAL AQUI ---
    # 1. Encontrar a data de início de cada fundo.
    inception_dates = []
    for f in fundos_a_exibir.keys():
        first_valid_index = df_retornos[f].dropna().index.min()
        inception_dates.append(first_valid_index)

    # 2. Definir o ponto de partida comum como a data MAIS RECENTE (máxima).
    common_start_date = max(inception_dates)
    
    cores_fundos = [COLOR_PALETTE['primary_light'], COLOR_PALETTE['primary_dark']]

    # Loop para plotar cada fundo (Sênior e Subordinada)
    for i, (fundo_nome, fundo_label) in enumerate(fundos_a_exibir.items()):
        retornos_diarios_full = df_retornos[fundo_nome].dropna()
        
        # 3. Filtrar os retornos para começar APENAS a partir da data de início comum.
        retornos_diarios = retornos_diarios_full[retornos_diarios_full.index >= common_start_date]
        
        if retornos_diarios.empty:
            continue
        
        # 4. Calcular o produto acumulado SOBRE A SÉRIE JÁ FILTRADA.
        df_fund_evol_norm = (1 + retornos_diarios).cumprod() * 100
        
        cor = cores_fundos[i % len(cores_fundos)]
        ax.plot(df_fund_evol_norm, color=cor, lw=2.5, label=fundo_label)
        ax.text(df_fund_evol_norm.index[-1], df_fund_evol_norm.iloc[-1], f' {df_fund_evol_norm.iloc[-1]:.2f}', 
                color=cor, fontsize=12, va='center')

    # Plota o Benchmark (CDI) também a partir da data de início comum
    retornos_cdi = df_retornos[benchmark][df_retornos.index >= common_start_date].dropna()
    df_bench_evol_norm = (1 + retornos_cdi).cumprod() * 100
    ax.plot(df_bench_evol_norm, color='k', linestyle='--', alpha=0.7, lw=2, label=benchmark)
    ax.text(df_bench_evol_norm.index[-1], df_bench_evol_norm.iloc[-1], f' {df_bench_evol_norm.iloc[-1]:.2f}', 
            color='k', alpha=0.8, fontsize=12, va='center')

    # Formatação do Gráfico (sem alterações)
    ax.spines[['top', 'right']].set_visible(False)
    ax.yaxis.set_major_formatter(ticker.FormatStrFormatter('%.0f'))
    ax.grid(axis='y', linestyle='--', alpha=0.7)
    ax.set_title(f'Retorno Acumulado (Início Comum: {common_start_date.strftime("%d/%m/%Y")})', fontsize=16, pad=20)
    ax.legend(loc='upper left', frameon=False, fontsize=12)
    plt.tight_layout()
    path_retorno = os.path.join(PASTA_SAIDA_IMAGENS, f"grafico_retorno_{nome_arquivo_limpo}.png")
    fig.savefig(path_retorno, dpi=300, bbox_inches='tight'); plt.close(fig)
    print(f"Gráfico de retorno (início comum) salvo em: {path_retorno}")

    # --- O resto da função (Vol, DD, PL) permanece o mesmo ---
    # (Não é necessário colar o resto da função aqui, pois não há mudanças)
    # --- 2. Gráfico de Volatilidade (sem alterações, continua sendo gerado) ---
    fig, ax = plt.subplots(figsize=(12, 6))
    window=22
    ret_fundo = df_cota.loc[fundo_alvo_principal].pct_change()
    vol_fundo = ret_fundo.rolling(window=window).std() * np.sqrt(252)
    ax.plot(vol_fundo, color=COLOR_PALETTE['primary_light'], lw=2, label=nome_exib_principal)
    ret_bench = df_cota.loc[benchmark].pct_change(); vol_bench = ret_bench.rolling(window=window).std() * np.sqrt(252)
    ax.plot(vol_bench, color='k', linestyle='--', alpha=0.7, lw=2, label=benchmark)
    ax.spines[['top', 'right']].set_visible(False); ax.yaxis.set_major_formatter(ticker.PercentFormatter(1.0))
    ax.grid(axis='y', linestyle='--', alpha=0.7)
    ax.set_title(f'Volatilidade Anualizada (Janela Móvel de {window}d) - {nome_exib_principal}', fontsize=16, pad=20)
    ax.legend(loc='upper left', frameon=False, fontsize=12); plt.tight_layout()
    path_vol = os.path.join(PASTA_SAIDA_IMAGENS, f"grafico_volatilidade_{nome_arquivo_limpo}.png")
    fig.savefig(path_vol, dpi=300, bbox_inches='tight'); plt.close(fig)
    print(f"Gráfico de volatilidade salvo em: {path_vol}")

    # --- 3. Gráfico de Drawdown (sem alterações, continua sendo gerado) ---
    fig, ax = plt.subplots(figsize=(12, 6))
    df_fund_dd = df_cota.loc[fundo_alvo_principal].dropna()
    daily_dd = df_fund_dd / df_fund_dd.cummax() - 1.0
    ax.plot(daily_dd, color=COLOR_PALETTE['primary_light'], lw=1)
    ax.fill_between(daily_dd.index, daily_dd, 0, color=COLOR_PALETTE['primary_light'], alpha=0.3)
    ax.spines[['top', 'right']].set_visible(False); ax.yaxis.set_major_formatter(ticker.PercentFormatter(1.0))
    ax.grid(axis='y', linestyle='--', alpha=0.7)
    ax.set_title(f'Drawdown Histórico - {nome_exib_principal}', fontsize=16, pad=20); plt.tight_layout()
    path_dd = os.path.join(PASTA_SAIDA_IMAGENS, f"grafico_drawdown_{nome_arquivo_limpo}.png")
    fig.savefig(path_dd, dpi=300, bbox_inches='tight'); plt.close(fig)
    print(f"Gráfico de drawdown salvo em: {path_dd}")

    # --- 4. Gráfico de Evolução do PL (sem alterações) ---
    fig, ax = plt.subplots(figsize=(12, 6))
    fundos_originais = list(fundos_a_exibir.keys())
    df_plot_pl = df_patr.loc[fundos_originais].T.dropna(how='all').fillna(0)
    df_plot_pl = df_plot_pl[['FIDC FCT II', 'FIDC FCT II SR2']]
    cores_pl = [COLOR_PALETTE['primary_dark'], COLOR_PALETTE['primary_light']]
    ax.stackplot(df_plot_pl.index, df_plot_pl.T.values,
                labels=['Subordinada', 'Sênior'],
                colors=cores_pl, alpha=0.8)
    ax.spines[['top', 'right']].set_visible(False)
    ax.yaxis.set_major_formatter(ticker.FuncFormatter(lambda x, p: f'R$ {x/1e6:.1f}M'))
    ax.grid(axis='y', linestyle='--', alpha=0.7)
    ax.set_title(f'Evolução do Patrimônio Líquido (PL)', fontsize=16, pad=20)
    ax.legend(loc='upper left', frameon=False, fontsize=12)
    pl_total_final = df_plot_pl.iloc[-1].sum()
    ax.text(df_plot_pl.index[-1], pl_total_final, f' Total R$ {pl_total_final/1e6:.2f}M',
            color=COLOR_PALETTE['secondary'], fontsize=12, va='center', ha='left')
    plt.tight_layout()
    path_pl = os.path.join(PASTA_SAIDA_IMAGENS, f"grafico_pl_{nome_arquivo_limpo}.png")
    fig.savefig(path_pl, dpi=300, bbox_inches='tight'); plt.close(fig)
    print(f"Gráfico de PL empilhado salvo em: {path_pl}")
    
    return {
        'retorno': path_retorno, 'volatilidade': path_vol,
        'drawdown': path_dd, 'pl': path_pl
    }

def plotar_graficos_estoque(df_estoque):
    """Cria e salva todos os gráficos relacionados ao estoque."""
    if df_estoque.empty:
        print("DataFrame de estoque vazio. Pulando geração de gráficos de estoque.")
        return {}
    
    print("\n--- Gerando imagens de Estoque ---")
    df_avencer = df_estoque[df_estoque['Status'] == 'A vencer'].copy()
    paths = {}
    
    # Gerar e salvar cada gráfico de estoque
    paths['venc_mensal'] = plotar_vencimento_mensal(df_avencer, col_valor='Valor Presente', col_vcto='VCTO_MES')
    paths['venc_anual'] = plotar_vencimento_anual(df_avencer, col_valor='Valor Presente', col_vcto='VCTO_ANO')
    paths['conc_uf'] = plotar_concentracao_uf(df_avencer, col_valor='Valor Presente', col_class='UF')
    paths['dist_capag'] = plotar_distribuicao_capag(df_estoque, class_col='CAPAG', col_valor='Valor Presente')
    paths['aging_vencidos'] = plotar_aging_com_acumulado(df_estoque, col_valor='Valor Presente')
    paths['conc_cumulativa_sacado'] = plotar_concentracao_cumulativa_sacado(df_avencer, col_valor='Valor Presente', col_class='CPF Cliente')
    print("\nGráficos de estoque gerados com sucesso.")
    return paths

def plotar_vencimento_mensal(df_aberto, col_valor, col_vcto, cutoff=10):
    """Gera e salva o gráfico de vencimento mensal com a estética original."""
    df_aberto = df_aberto.copy()
    df_aberto[col_vcto] = df_aberto['Data Vencimento'].dt.strftime('%Y-%m')
    df2 = df_aberto.groupby(col_vcto)[col_valor].sum()
    df3 = (df2 / df2.sum()).cumsum()
    df2 = df2[:cutoff]
    df3 = df3[:cutoff]

    if df2.empty:
        print("Não há dados para gerar o gráfico de vencimento mensal.")
        return None

    # Restaura a lógica original de limites dinâmicos para espaçamento visual
    ymin1 = 0
    ymax1 = df2.max() * 2
    ymax2 = 1.3
    ymin2 = df3.min() * 2.5 - 1.5 * ymax2

    fig, ax = plt.subplots(figsize=(8, 4.5))
    ax2 = ax.twinx()

    # Restaura o np.clip para garantir que barras pequenas sejam visíveis
    bars = ax.bar(df2.index, np.clip(df2.values, 0.01 * df2.max(), df2.max()), color=COLOR_PALETTE['primary_dark'])
    ax2.plot(df3.index, df3.values, color=COLOR_PALETTE['secondary'], lw=2, marker='o')

    # Lógica de formatação dos eixos e rótulos do script original
    month_map = {'01':'jan','02':'fev','03':'mar','04':'abr','05':'mai','06':'jun','07':'jul','08':'ago','09':'set','10':'out','11':'nov','12':'dez'}
    ax.set_xticks(df2.index)
    ax.set_xticklabels([f"{month_map[x[-2:]]}/{x[2:4]}" for x in df2.index])
    ax.set_ylim(ymin1, ymax1)
    ax2.set_ylim(ymin2, ymax2)
    ax.set_yticks([])
    ax2.set_yticks([])
    ax.tick_params(bottom=False)
    ax.set_frame_on(False)
    ax2.spines[:].set_visible(False)

    max_height = max([bar.get_height() for bar in bars])
    for i, bar in enumerate(bars):
        h = bar.get_height() + max_height * 0.04
        value = df2.values[i] / 1000
        ax.text(bar.get_x() + bar.get_width()/2, h, f'{value:,.0f}'.replace(',', '.'), ha='center')

    for x, y in df3.items():
        value = f'{y:.0%}' if ((y < .990) or (y == 1)) else '>99%'
        ax2.text(x, y + (ymax2 - ymin2) * 0.04, value, ha='center')

    path_saida = os.path.join(PASTA_SAIDA_IMAGENS, "vencimento_mensal.png")
    fig.savefig(path_saida, bbox_inches='tight')
    plt.close(fig)
    print(f"Gráfico de vencimento mensal (estilo corrigido) salvo em: {path_saida}")
    return path_saida

def plotar_vencimento_anual(df_aberto, col_valor, col_vcto):
    """Gera e salva o gráfico de vencimento anual com a estética original."""
    df_aberto = df_aberto.copy()
    df_aberto[col_vcto] = df_aberto['Data Vencimento'].dt.year
    df2 = df_aberto.groupby(col_vcto)[col_valor].sum()
    df3 = df2.cumsum() / df2.sum()

    if df2.empty:
        print("Não há dados para gerar o gráfico de vencimento anual.")
        return None

    # Restaura a lógica original de limites dinâmicos para espaçamento visual
    ymin1 = 0
    ymax1 = df2.max() * 2
    ymax2 = 1.15
    ymin2 = df3.min() * 2.5 - 1.5 * ymax2

    fig, ax = plt.subplots(figsize=(8, 4.5))
    ax2 = ax.twinx()
    
    # Restaura o np.clip para garantir que barras pequenas sejam visíveis
    bars = ax.bar(df2.index, np.clip(df2.values, 0.01 * df2.max(), df2.max()), color=COLOR_PALETTE['primary_dark'])
    ax2.plot(df3.index, df3.values, color=COLOR_PALETTE['secondary'], lw=2, marker='o')
    
    # Lógica de formatação dos eixos e rótulos do script original
    ax.set_xticks(df2.index.tolist())
    ax.set_ylim(ymin1, ymax1)
    ax2.set_ylim(ymin2, ymax2)
    ax.set_yticks([])
    ax2.set_yticks([])
    ax.tick_params(bottom=False)
    ax.set_frame_on(False)
    ax2.spines[:].set_visible(False)

    max_height = max([bar.get_height() for bar in bars])
    for i, bar in enumerate(bars):
        h = bar.get_height() + max_height * 0.04
        value = df2.values[i] / 1000
        ax.text(bar.get_x() + bar.get_width()/2, h, f'{value:,.0f}'.replace(',', '.'), ha='center')

    for x, y in df3.items():
        value = f'{y:.0%}' if ((y < .990) or np.isclose(y, 1)) else '>99%'
        ax2.text(x, y + 0.06, value, ha='center')

    path_saida = os.path.join(PASTA_SAIDA_IMAGENS, "vencimento_anual.png")
    fig.savefig(path_saida, bbox_inches='tight')
    plt.close(fig)
    print(f"Gráfico de vencimento anual (estilo corrigido) salvo em: {path_saida}")
    return path_saida

def plotar_concentracao_uf(df_aberto, col_valor, col_class='UF', cutoff=15):
    """Gera e salva o gráfico de concentração por UF."""
    df11 = df_aberto.groupby(col_class)[col_valor].sum().sort_values(ascending=False)
    ncut = min(len(df11), cutoff)
    
    if ncut >= cutoff:
        x = df11[:cutoff].index.tolist() + ['Outros']
        y = list(df11[:cutoff].values) + [df11.iloc[cutoff:].sum()]
    else:
        x = df11.index.tolist()
        y = list(df11.values)
        
    fig, ax = plt.subplots(figsize=(6, 4))
    bars = ax.barh(x, y, color=COLOR_PALETTE['primary_dark'])
    plt.gca().invert_yaxis()
    ax.set_xticks([])
    ax.tick_params(left=False, bottom=False)
    ax.set_frame_on(False)

    for i, bar in enumerate(bars):
        value = y[i] / 1000
        ax.text(bar.get_width() * 1.01, bar.get_y() + bar.get_height()/2, f'{value:,.0f}'.replace(',', '.'), va='center', ha='left')

    path_saida = os.path.join(PASTA_SAIDA_IMAGENS, "concentracao_uf.png")
    fig.savefig(path_saida, bbox_inches='tight')
    plt.close(fig)
    print(f"Gráfico de concentração por UF salvo em: {path_saida}")
    return path_saida

def plotar_distribuicao_capag(df_in, class_col, col_valor):
    """Gera e salva o gráfico de distribuição por CAPAG."""
    df_capag = df_in.dropna(subset=[class_col]).groupby(class_col)[col_valor].sum()
    fig, ax = plt.subplots(figsize=(8, 4.5))
    bars = ax.bar(df_capag.index, df_capag.values, color=COLOR_PALETTE['primary_dark'])

    ax.set_xticks(df_capag.index.tolist())
    ax.set_yticks([])
    ax.tick_params(bottom=False)
    ax.set_frame_on(False)

    for i, bar in enumerate(bars):
        value = df_capag.values[i] / 1000
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() * 1.05, f'{value:,.0f}'.replace(',', '.'), ha='center')

    path_saida = os.path.join(PASTA_SAIDA_IMAGENS, "distribuicao_capag.png")
    fig.savefig(path_saida, bbox_inches='tight')
    plt.close(fig)
    print(f"Gráfico de distribuição por CAPAG salvo em: {path_saida}")
    return path_saida

def plotar_concentracao_cumulativa_sacado(df_aberto, col_valor='Valor Presente', col_class='CPF Cliente'):
    """
    Gera um gráfico de barras horizontal mostrando a concentração CUMULATIVA de valor
    em faixas de devedores (Top 10, Top 20, Top 50, etc.), como na imagem de exemplo.
    """
    df_agrupado = df_aberto.groupby(col_class)[col_valor].sum().sort_values(ascending=False)
    
    if df_agrupado.empty:
        print(f"Não há dados para gerar o gráfico de concentração cumulativa por '{col_class}'.")
        return None

    valor_total = df_agrupado.sum()
    num_sacados = len(df_agrupado)

    # Define as faixas cumulativas, como na imagem
    bins = [10, 20, 50, 100, 200, 500, 1000, 2000, 5000, 10000]
    
    resultados = {}
    # Calcula o percentual para cada faixa
    for n in bins:
        if n < num_sacados:
            soma_top_n = df_agrupado.iloc[0:n].sum()
            pct_top_n = (soma_top_n / valor_total) * 100
            resultados[f'{n} maiores'] = pct_top_n
            
    # Adiciona a linha final com o total de sacados
    resultados[f'{num_sacados} maiores'] = 100.0
    
    df_plot = pd.Series(resultados)
    
    # --- PLOTAGEM ---
    fig, ax = plt.subplots(figsize=(8, 4.5))
    bars = ax.barh(df_plot.index, df_plot.values, color=COLOR_PALETTE['primary_dark'])
    
    # Inverte o eixo para ter o '10 maiores' no topo
    ax.invert_yaxis()
    
    # Formatação dos eixos
    ax.set_xticks([])
    ax.set_frame_on(False)
    ax.tick_params(left=False) # Remove os tracinhos do eixo Y
    
    # Adiciona os rótulos de percentual no final das barras
    for bar in bars:
        width = bar.get_width()
        ax.text(width + 1, bar.get_y() + bar.get_height()/2, f'{width:.0f}%', 
                va='center', ha='left', fontsize=9)

    ax.set_title("Concentração Cumulativa por Sacado", fontsize=12)
    
    path_saida = os.path.join(PASTA_SAIDA_IMAGENS, "concentracao_cumulativa_sacado.png")
    fig.savefig(path_saida, bbox_inches='tight')
    plt.close(fig)
    print(f"Gráfico de concentração cumulativa por sacado salvo em: {path_saida}")
    return path_saida


def plotar_aging_com_acumulado(df_estoque, col_valor='Valor Presente'):
    """
    Gera um gráfico de Aging aprimorado, com mais faixas de atraso e
    rótulos de dados na linha de percentual acumulado, garantindo que
    a linha de acumulado comece visualmente acima das barras.
    """
    df_vencidos = df_estoque[df_estoque['Status'] == 'Vencido'].copy()

    if df_vencidos.empty:
        print("Não há dados de carteira vencida para gerar o gráfico de Aging.")
        return None

    hoje = datetime.now()
    df_vencidos['dias_atraso'] = (hoje - df_vencidos['Data Vencimento']).dt.days

    # Faixas de atraso
    bins = [-1, 14, 30, 60, 90, 120, 150, 180, np.inf]
    labels = ['0-14 D', '15-30 D', '31-60 D', '61-90 D', '91-120 D', '121-150 D', '151-180 D', '180+ D']
    df_vencidos['Faixa Atraso'] = pd.cut(df_vencidos['dias_atraso'], bins=bins, labels=labels, right=True)
    
    df_aging = df_vencidos.groupby('Faixa Atraso', observed=False)[col_valor].sum()
    total_em_atraso = df_aging.sum()

    df_aging_pct_acum = df_aging.cumsum() / total_em_atraso
    df_aging_pct_acum['Total'] = 1.0
    df_aging['Total'] = total_em_atraso
    
    fig, ax = plt.subplots(figsize=(8, 4.5))
    ax2 = ax.twinx()

    # Gráfico de Barras
    bars = ax.bar(df_aging.index, df_aging.values, color=COLOR_PALETTE['primary_dark'], width=0.6)
    bars[-1].set_color(COLOR_PALETTE['secondary'])
    
    # --- INÍCIO DA NOVA LÓGICA DE POSICIONAMENTO DA LINHA ---
    
    # 1. Definir o limite superior do eixo das barras para calcular a proporção
    ax_ymax = total_em_atraso * 1.40 # Aumenta um pouco o espaço para os rótulos
    ax.set_ylim(0, ax_ymax)

    # 2. Encontrar a altura relativa da maior barra (considerando o rótulo)
    max_bar_height = 0 if df_aging.empty else df_aging.max()
    # Adicionamos um buffer para o texto do rótulo da barra não tocar na linha
    max_bar_ratio = (max_bar_height / ax_ymax) * 1.15
    
    # 3. Calcular os limites do eixo da linha (ax2) para que ela comece acima das barras.
    # A linha ocupará o espaço visual entre (max_bar_ratio + 10%) e 95% da altura do gráfico.
    line_visual_start = min(max_bar_ratio + 0.1, 0.7) # Onde o primeiro ponto da linha aparecerá (em %)
    line_visual_end = 0.95                            # Onde o último ponto (100%) aparecerá

    # Valores de dados correspondentes ao início e fim da linha
    line_data_start = df_aging_pct_acum.iloc[0]
    line_data_end = 1.0

    # Evita divisão por zero se o espaço visual for muito pequeno
    if abs(line_visual_end - line_visual_start) < 0.01:
        line_visual_start = 0.6
        line_visual_end = 0.95

    # 4. Resolve um sistema de equações para encontrar os limites (y_min, y_max) do eixo ax2
    # que mapeiam os dados da linha para o espaço visual desejado.
    # (y_data - y_min) / (y_max - y_min) = y_visual_position
    y2_range = (line_data_end - line_data_start) / (line_visual_end - line_visual_start)
    y2_min = line_data_start - (y2_range * line_visual_start)
    y2_max = y2_min + y2_range
    
    # 5. Define os novos limites calculados para o eixo da linha
    ax2.set_ylim(y2_min, y2_max)
    
    # --- FIM DA NOVA LÓGICA ---

    # Gráfico de Linha
    line = ax2.plot(df_aging_pct_acum.index, df_aging_pct_acum.values, color=COLOR_PALETTE['secondary'], marker='o', lw=2)

    # Formatação do eixo das barras
    ax.set_yticks([])
    ax.set_frame_on(False)
    for bar in bars:
        height = bar.get_height()
        ax.text(bar.get_x() + bar.get_width() / 2, height, f'R$ {height/1e6:.2f}M'.replace('.',','),
                ha='center', va='bottom', fontsize=9, color='#333333')

    # Formatação do eixo da linha
    ax2.yaxis.set_major_formatter(ticker.PercentFormatter(1.0))
    ax2.spines[:].set_visible(False)
    
    # Adiciona rótulos na linha de forma dinâmica, usando a nova escala
    y2_axis_range = ax2.get_ylim()[1] - ax2.get_ylim()[0]
    label_offset = y2_axis_range * 0.05 # Offset de 5% do novo range do eixo
    
    for i, pct in enumerate(df_aging_pct_acum):
        if i < len(df_aging_pct_acum) - 1: # Não adiciona rótulo no ponto "Total"
            ax2.text(i, pct + label_offset, f'{pct:.0%}', ha='center', fontsize=9, color=COLOR_PALETTE['secondary'])
    
    # Salva a imagem
    path_saida = os.path.join(PASTA_SAIDA_IMAGENS, "aging_com_acumulado.png")
    fig.savefig(path_saida, bbox_inches='tight')
    plt.close(fig)
    print(f"Gráfico de Aging (versão aprimorada) salvo em: {path_saida}")
    return path_saida

# %%
# =============================================================================
# CÉLULA 4: FUNÇÕES DE MONTAGEM DO RELATÓRIO (DOCX)
# =============================================================================
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

### 1. FUNÇÃO DE BORDA ATUALIZADA (ADICIONA LINHA EM CIMA E EMBAIXO) ###
def set_paragraph_horizontal_border(paragraph, color="76C6C5", size="4", space="1"):
    """
    Adiciona bordas na parte SUPERIOR e INFERIOR de um parágrafo.
    """
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    
    # Loop para criar a borda 'top' (superior) e 'bottom' (inferior)
    for border_side in ['top', 'bottom']:
        border_element = OxmlElement(f'w:{border_side}')
        border_element.set(qn('w:val'), 'single')
        border_element.set(qn('w:sz'), size)
        border_element.set(qn('w:space'), space)
        border_element.set(qn('w:color'), color)
        pBdr.append(border_element)
        
    pPr.append(pBdr)

def configurar_cabecalho_rodape(doc, nome_relatorio, data_referencia, path_logo):
    """Configura o cabeçalho e as margens do documento."""
    section = doc.sections[0]
    section.top_margin, section.bottom_margin = Cm(1), Cm(1)
    section.left_margin, section.right_margin = Cm(1.6), Cm(1)
    
    header = section.header
    header_table = header.tables[0] if header.tables else header.add_table(rows=1, cols=2, width=Cm(18))
    
    # (O resto desta função permanece o mesmo)
    cell_left = header_table.cell(0, 0)
    p_left = cell_left.paragraphs[0]
    p_left.text = ""
    run_name = p_left.add_run(nome_relatorio.upper())
    run_name.font.name = "Gill Sans MT"
    run_name.font.size, run_name.bold = Pt(14), True
    run_name.font.color.rgb = RGBColor(16, 112, 130)
    p_left.add_run("\n")
    meses_pt = ["Janeiro", "Fevereiro", "Março", "Abril", "Maio", "Junho", "Julho", "Agosto", "Setembro", "Outubro", "Novembro", "Dezembro"]
    data_formatada = f"{meses_pt[data_referencia.month - 1]} {data_referencia.year}"
    run_month = p_left.add_run(data_formatada)
    run_month.font.name, run_month.font.size, run_month.italic = "Arial", Pt(12), True
    run_month.font.color.rgb = RGBColor(38, 38, 38)
    cell_right = header_table.cell(0, 1)
    p_right = cell_right.paragraphs[0]
    p_right.text = ""
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_right.add_run().add_picture(path_logo, width=Cm(4.93))

def montar_corpo_relatorio(doc, paths_imagens):
    """Adiciona a tabela e os gráficos ao corpo do documento (VERSÃO CORRIGIDA)."""
    
    def add_section_title(cell, title, subtitle=""):
        p = cell.add_paragraph()
        p.paragraph_format.space_before, p.paragraph_format.space_after = Pt(6), Pt(2)
        
        run = p.add_run(title)
        run.font.name, run.font.size, run.bold = "Gill Sans MT", Pt(10), False
        run.font.color.rgb = RGBColor(89, 89, 89)
        
        if subtitle:
            p.add_run("\n")
            run_sub = p.add_run(subtitle)
            run_sub.font.name = "Gill Sans MT"
            run_sub.font.size = Pt(8)
            run_sub.font.color.rgb = RGBColor(89, 89, 89)
            
        set_paragraph_horizontal_border(p, color="E2EEF0", space="2")

    # Tabela de Rentabilidade
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(12)
    run = title_p.add_run("\tRentabilidade Mensal")
    run.font.name, run.font.size, run.bold = "Gill Sans MT", Pt(11), False
    set_paragraph_horizontal_border(title_p, color="E2EEF0", space="4")
    
    # Adiciona a imagem da tabela de performance, se ela existir
    if paths_imagens.get('tabela_perf') and os.path.exists(paths_imagens['tabela_perf']):
        doc.add_picture(paths_imagens['tabela_perf'], width=Cm(17.8))

    # Tabela de Gráficos
    graphs_table = doc.add_table(rows=5, cols=2)
    graphs_table.autofit = False
    for col in graphs_table.columns:
        col.width = Cm(8.9)
    img_width = Cm(8.5)

    ### 1. MAPA DE GRÁFICOS CORRIGIDO (AGORA COM A CHAVE DE REFERÊNCIA DIRETA) ###
    # A estrutura agora é: (Título, Subtítulo, Chave_do_Dicionário_paths_imagens)

    
    mapa_graficos = {
        (0, 0): ("Retorno Acumulado", "", 'retorno'),
        (0, 1): ("Evolução do Patrimônio Líquido (PL)", "(R$ mm)", 'pl'),
        (1, 0): ("Vencimento Mensal", "(R$’000 e % acumulado)", 'venc_mensal'),
        (1, 1): ("Vencimento Anual", "(R$’000 e % acumulado)", 'venc_anual'),
        (2, 0): ("Concentração por UF", "(R$’000)", 'conc_uf'),
        (2, 1): ("Distribuição por CAPAG", "(R$’000)", 'dist_capag'),
        (3, 0): ("Concentração Cumulativa por Sacado", "(%)", 'conc_cumulativa_sacado'),
        (3, 1): ("Aging da Carteira Vencida", "(R$ mm e % acumulado)", 'aging_vencidos'),
        #(1, 0): ("Volatilidade Anualizada", "(Janela móvel 22d)", 'volatilidade'),
        #(1, 1): ("Drawdown Histórico", "", 'drawdown'),
    }
    
    ### 2. LOOP DE PLOTAGEM CORRIGIDO E SIMPLIFICADO ###
    # Este novo loop busca a imagem pela chave correta, sem "adivinhações"
    for (row, col), (titulo, subtitulo, path_key) in mapa_graficos.items():
        # Busca o caminho da imagem usando a chave de referência direta
        path = paths_imagens.get(path_key)
        
        # Verifica se o caminho foi encontrado e se o arquivo existe
        if path and os.path.exists(path):
            cell = graphs_table.cell(row, col)
            add_section_title(cell, f"\t{titulo}", f"\t{subtitulo}")
            p_img = cell.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.add_run().add_picture(path, width=img_width)

def add_page_field(run, field_type):
    """Adiciona um campo de numeração de página (PAGE ou NUMPAGES) a um 'run'."""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = field_type

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


### 4. FUNÇÃO PARA ADICIONAR SEÇÃO FINAL (DADOS, DISCLAIMER, RODAPÉ) ###
### 4. FUNÇÃO PARA ADICIONAR SEÇÃO FINAL (DADOS, DISCLAIMER, RODAPÉ) ###
def adicionar_secao_final(doc, df_patr, anbima_path):
    """
    Adiciona a seção final do relatório, contendo a tabela de dados do fundo,
    o disclaimer e o rodapé personalizado.
    """
    from dateutil.relativedelta import relativedelta
    
    # --- Adicionar o título "Dados do Fundo" ---
    p_dados_title = doc.add_paragraph()
    run_dados_title = p_dados_title.add_run("\tDados do Fundo")
    run_dados_title.font.name = "Gill Sans MT"
    run_dados_title.font.size = Pt(10)
    run_dados_title.font.color.rgb = RGBColor(89, 89, 89)
    set_paragraph_horizontal_border(p_dados_title, "E2EEF0", space="4")
    p_dados_title.paragraph_format.space_after = Pt(5)
    p_dados_title.paragraph_format.space_before = Pt(25)

    # --- Inserir a tabela mãe imediatamente após o título ---
    data_parent_tbl = doc.add_table(rows=1, cols=2)
    p_dados_title._element.addnext(data_parent_tbl._element)

    # Remove parágrafos vazios entre o título e a tabela
    next_el = p_dados_title._element.getnext()
    if next_el is not None and next_el.tag.endswith("p") and not next_el.text:
         p_dados_title._p.getparent().remove(next_el)

    # --- Calcular os dados de patrimônio (LÓGICA CORRIGIDA) ---
    if df_patr.empty or len(df_patr.columns) == 0:
        print("AVISO: df_patr está vazio. Dados de PL não serão preenchidos.")
        pl_ultimo, pl_12m, date_ref = 0, 0, datetime.now()
    else:
        # CORREÇÃO: As datas estão nas COLUNAS, não no índice.
        all_dates = pd.to_datetime(df_patr.columns)
        last_date = all_dates[-1]

        # Encontrar a data de referência (último dia útil do mês anterior) de forma robusta
        first_day_of_last_month = last_date.replace(day=1)
        last_day_of_prev_month = first_day_of_last_month - pd.Timedelta(days=1)
        
        # Filtra as colunas (datas) que estão no mês anterior ou antes
        prev_month_dates = all_dates[all_dates <= last_day_of_prev_month]
        
        if not prev_month_dates.empty:
            date_ref = prev_month_dates[-1]
        else:
            # Fallback: se não houver dados no mês anterior, usa a última data disponível
            date_ref = last_date

        # Calcula o PL na data de referência somando os valores da coluna daquela data
        pl_ultimo = df_patr[date_ref].sum()

        # Calcular PL médio dos últimos 12 meses
        date_start_12m = date_ref - relativedelta(years=1)
        cols_12m = all_dates[(all_dates > date_start_12m) & (all_dates <= date_ref)]
        
        if not cols_12m.empty:
            # Soma o PL de cada dia (axis=0) e depois calcula a média desses totais diários
            pl_12m = df_patr[cols_12m].sum(axis=0).mean()
        else:
            pl_12m = 0

    pl_ultimo_formatted = f"{pl_ultimo:,.0f}".replace(",", ".")
    pl_12m_formatted = f"{pl_12m:,.0f}".replace(",", ".")

    # --- Definir conteúdo das tabelas ---
    left_data_info = [
        "Política de Investimento\nO fundo investe em direitos creditórios relacionados a crédito consignado público nas esferas municipal, estadual e federal.",
        "Denominação Social\nFUNDO DE INVESTIMENTO EM DIREITOS CREDITÓRIOS FCT II", # Exemplo
        "CNPJ\n58.055.083/0001-04", # Exemplo
        "Data de início\n27/11/2024", # Exemplo
        "Rentabilidade alvo da cota Sênior\n16,0% a.a.",
        "Rentabilidade alvo da cota Mezanino\n16,5% a.a.",
        f"Patrimônio Líquido (Ref: {date_ref.strftime('%d/%m/%Y')})\nR$ {pl_ultimo_formatted}",
        f"Patrimônio Líquido Médio (12M)\nR$ {pl_12m_formatted}"
    ]
    right_data_info = [
        "Gestor\nPorto Real Investimentos", "Administrador\nBanvox DTVM", "Custodiante\nBanvox DTVM",
        "Taxa de gestão\n2,5%", "Taxa de administração\n0,11%", "Taxa máxima de custódia\n0,03%",
        "Taxa de performance\nNão há", "Forma de condomínio\nFechado"
    ]
    
    # --- Preencher as tabelas (o resto da função permanece igual) ---
    def fill_data_cell(text_str, target_cell):
        p = target_cell.paragraphs[0] if target_cell.paragraphs else target_cell.add_paragraph()
        p.clear()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.2
        
        identifier, value = text_str.split("\n", 1)
        run_id = p.add_run(identifier.strip())
        run_id.font.name = "Arial"; run_id.font.size = Pt(8.5)
        run_id.font.color.rgb = RGBColor(89, 89, 89); run_id.bold = True
        run_id.add_break()

        run_val = p.add_run(value.strip())
        run_val.font.name = "Arial"; run_val.font.size = Pt(8.5)
        run_val.font.color.rgb = RGBColor(89, 89, 89)

    left_tbl = data_parent_tbl.cell(0, 0).add_table(rows=len(left_data_info), cols=1)
    right_tbl = data_parent_tbl.cell(0, 1).add_table(rows=len(right_data_info), cols=1)

    for i, text in enumerate(left_data_info):
        fill_data_cell(text, left_tbl.cell(i, 0))
    for i, text in enumerate(right_data_info):
        fill_data_cell(text, right_tbl.cell(i, 0))
        
    # --- Disclaimer ao Final ---
    p_disc_title = doc.add_paragraph()
    run_disc_title = p_disc_title.add_run("\t Disclaimer")
    run_disc_title.font.name = "Gill Sans MT"; run_disc_title.font.size = Pt(10)
    run_disc_title.font.color.rgb = RGBColor(89, 89, 89)
    set_paragraph_horizontal_border(p_disc_title, "E2EEF0")

    p_disc_body = doc.add_paragraph()
    p_disc_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run_disc_body = p_disc_body.add_run(
        "As informações contidas neste material são de caráter exclusivamente informativo. "
        "A rentabilidade passada não representa garantia de rentabilidade futura. "
        "Fundos de investimento não contam com a garantia do administrador, do gestor da carteira, "
        "de qualquer mecanismo de seguro ou, ainda, do Fundo Garantidor de Créditos ‐ FGC. "
        "Os fundos de crédito privado estão sujeitos a risco de perda substancial de seu patrimônio líquido "
        "em caso de eventos que acarretem o não pagamento dos ativos integrantes de sua carteira, inclusive por força de intervenção, "
        "liquidação, regime de administração temporária, falência, recuperação judicial ou extrajudicial dos emissores responsáveis pelos ativos do fundo. "
        "Não há garantia de que este fundo terá o tratamento tributário para fundos de longo prazo. "
        "A rentabilidade divulgada não é líquida de impostos e taxa."
    )
    run_disc_body.font.name = "Arial"; run_disc_body.font.size = Pt(8.5)
    run_disc_body.font.color.rgb = RGBColor(166, 166, 166)

    # --- Rodapé ---
    section = doc.sections[0]
    footer = section.footer
    footer_table = footer.tables[0] if footer.tables else footer.add_table(rows=1, cols=3, width=Cm(18))

    # Esquerda: imagem Anbima
    cell_left = footer_table.cell(0, 0); cell_left.width = Cm(4)
    p_left = cell_left.paragraphs[0]; p_left.clear()
    p_left.add_run().add_picture(anbima_path, height=Cm(1.0))

    # Centro: informações de contato
    cell_center = footer_table.cell(0, 1); cell_center.width = Cm(10)
    p_center = cell_center.paragraphs[0]; p_center.clear()
    p_center.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_contato = p_center.add_run("PORTO REAL ASSET\nAv. São Gabriel 301 – 2º andar – Itaim Bibi – São Paulo (SP)\n")
    run_contato.font.name = "Arial"; run_contato.font.size = Pt(9)
    run_contato.font.color.rgb = RGBColor(127, 127, 127)
    run_link = p_center.add_run("contato@portorealasset.com.br")
    run_link.font.name = "Arial"; run_link.font.size = Pt(9)
    run_link.font.color.rgb = RGBColor(0, 0, 0); run_link.font.underline = True

    # Direita: numeração de páginas
    cell_right = footer_table.cell(0, 2); cell_right.width = Cm(4)
    p_right = cell_right.paragraphs[0]; p_right.clear()
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_priv = p_right.add_run("Documento estritamente privado e confidencial")
    run_priv.font.name = "Arial"; run_priv.font.size = Pt(8)
    run_priv.font.color.rgb = RGBColor(89, 89, 89)
    
    p_page = cell_right.add_paragraph(); p_page.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font_page = {"name": "Arial", "size": Pt(8), "color": RGBColor(89, 89, 89)}
    run_page_text = p_page.add_run("Página "); run_page_text.font.name = font_page['name']; run_page_text.font.size = font_page['size']; run_page_text.font.color.rgb = font_page['color']
    run_page_current = p_page.add_run(); add_page_field(run_page_current, "PAGE"); run_page_current.font.name = font_page['name']; run_page_current.font.size = font_page['size']; run_page_current.font.color.rgb = font_page['color']
    run_page_de = p_page.add_run(" de "); run_page_de.font.name = font_page['name']; run_page_de.font.size = font_page['size']; run_page_de.font.color.rgb = font_page['color']
    run_page_total = p_page.add_run(); add_page_field(run_page_total, "NUMPAGES"); run_page_total.font.name = font_page['name']; run_page_total.font.size = font_page['size']; run_page_total.font.color.rgb = font_page['color']

# %%
# =============================================================================
# CÉLULA 5: ORQUESTRADOR PRINCIPAL - EXECUÇÃO DO FLUXO
# =============================================================================

def gerar_relatorio_completo():
    """Função principal que orquestra todo o processo."""
    
    # --- ETAPA 1: PROCESSAMENTO DE DADOS DE RENTABILIDADE ---
    df_rentabilidade = carregar_dados_rentabilidade(PATH_RENTABILIDADE)
    if df_rentabilidade.empty:
        print("Processo interrompido: falha ao carregar dados de rentabilidade.")
        return

    df_cota = df_rentabilidade.pivot_table(index='data', columns='nomeFundo', values='valorCota').ffill().T
    df_patr = df_rentabilidade.pivot_table(index='data', columns='nomeFundo', values='pl').ffill().T
    
    df_cdi = obter_dados_bacen(12)
    if not df_cdi.empty:
        df_cota.loc['CDI'] = (1 + df_cdi['valor'] / 100).cumprod().ffill()
    
    df_retornos = preparar_df_retornos(df_rentabilidade, df_cota)
    
    # --- ETAPA 2: GERAÇÃO DAS IMAGENS ---
    print("\n--- Gerando imagens de Rentabilidade ---")
    paths_imagens = {}
    # Define os fundos a serem exibidos na tabela e seus novos nomes
    fundos_para_tabela = {
        'FIDC FCT II SR2': 'Sênior',
        'FIDC FCT II': 'Subordinada'
    }
    # Passa o DataFrame original processado para a função, que agora faz o cálculo internamente
    paths_imagens['tabela_perf'] = criar_tabela_performance(df_rentabilidade, fundos_para_tabela, df_cota)
    paths_rentabilidade = plotar_graficos_rentabilidade(df_cota, df_patr, df_retornos, fundos_para_tabela)
    
    paths_imagens.update(paths_rentabilidade)
    
    df_estoque = processar_dados_estoque(PATH_ESTOQUE)
    paths_estoque = plotar_graficos_estoque(df_estoque)
    paths_imagens.update(paths_estoque)
    
    # --- ETAPA 3: MONTAGEM E SALVAMENTO DO RELATÓRIO ---
    print("\n--- Montando relatório ---")
    ref_date = datetime.now()
    doc = Document(PATH_TEMPLATE_DOCX)
    
    configurar_cabecalho_rodape(doc, 'Lâmina de Performance e Estoque', ref_date, PATH_LOGO)
    montar_corpo_relatorio(doc, paths_imagens)
    
    # Note que usei df_patr (e não df_patrs) para corresponder ao que já existe no script
    adicionar_secao_final(doc, df_patr, PATH_ANBIMA_LOGO)
    
    nome_arquivo = f"Lamina_Completa_{ref_date.strftime('%Y-%m-%d')}.docx"
    path_docx = os.path.join(PASTA_SAIDA_RELATORIOS, nome_arquivo)
    
    try:
        doc.save(path_docx)
        print(f"Relatório DOCX gerado com sucesso em: {path_docx}")
        
        print("Iniciando conversão para PDF...")
        path_pdf = path_docx.replace(".docx", ".pdf")
        convert(path_docx, path_pdf)
        print(f"Relatório PDF gerado com sucesso em: {path_pdf}")
    except Exception as e:
        print(f"Ocorreu um erro ao salvar o relatório: {e}")

# --- Ponto de Entrada para Execução do Script ---
if __name__ == "__main__":
    gerar_relatorio_completo()