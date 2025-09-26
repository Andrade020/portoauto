# -*- coding: utf-8 -*-
"""
Pipeline completo para análise de performance e estoque de fundos.

Este script consolida as seguintes tarefas:
1.  Carrega e processa dados de rentabilidade de um fundo a partir de um arquivo CSV.
2.  Carrega e processa dados de estoque de múltiplos arquivos CSV.
3.  Gera uma série de visualizações (gráficos e tabelas) para as análises de performance e de estoque.
4.  Monta um relatório completo em formato DOCX e o converte para PDF, inserindo todas as visualizações geradas.
"""

# =============================================================================
# CÉLULA 1: BIBLIOTECAS E CONFIGURAÇÕES GERAIS
# =============================================================================
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import matplotlib.ticker as ticker
import os
import glob
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert
import io

# Configurações de visualização e estilo
pd.options.display.max_rows = 100
pd.options.display.max_columns = 50
plt.style.use('seaborn-v0_8-whitegrid')

# Paleta de cores para os gráficos
COLORS = {
    'primary_light': '#4BC1D7',
    'primary_dark': '#286D82',
    'primary_cyan': '#00B1DE',
    'neutral_gray': '#445D6D',
    'table_header_bg': '#e2eef0',
    'table_header_text': '#107082',
}

# =============================================================================
# CÉLULA 2: CLASSE DE CONFIGURAÇÃO
# =============================================================================
class Config:
    """
    Agrupa todas as configurações de caminhos e parâmetros da análise.
    ! ATENÇÃO: Modifique os caminhos e parâmetros nesta seção.
    """
    # --- Parâmetros da Análise ---
    FUND_NAME = 'FIDC FCT II'
    REFERENCE_DATE = datetime.now()
    
    # --- Caminhos Base ---
    BASE_DIR = Path(r"C:\Users\Leo\Desktop\Porto_Real\portoauto\src\vortx_estoques")
    OUTPUT_DIR = BASE_DIR / "output" / "relatorio_final"
    
    # --- Caminhos de Entrada ---
    # Dados de Rentabilidade
    PERFORMANCE_DATA_PATH = BASE_DIR / "data" / "135972-Rentabilidade_Sintetica.csv"
    # Dados de Estoque (pasta com múltiplos CSVs)
    INVENTORY_DATA_DIR = BASE_DIR / "data" / "estoque_consolidado"
    
    # Assets para o Relatório
    TEMPLATE_PATH = Path(r"C:\Users\Leo\Desktop\Porto_Real\portoreal\notebooks\template.docx")
    LOGO_PATH = Path(r"C:\Users\Leo\Desktop\Porto_Real\portoreal\images\logo.png")
    
    # --- Nomes dos Arquivos de Saída ---
    OUTPUT_FILENAME_BASE = f"Lamina_Completa_{FUND_NAME.replace(' ', '_')}_{REFERENCE_DATE.strftime('%Y-%m')}"
    
    def __init__(self):
        """Cria as pastas de saída se não existirem."""
        self.OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
        print(f"Relatórios serão salvos em: {self.OUTPUT_DIR}")

# =============================================================================
# MÓDULO 1: PROCESSAMENTO DE DADOS DE RENTABILIDADE
# =============================================================================

def load_performance_data(file_path):
    """Carrega e limpa os dados de rentabilidade do arquivo CSV."""
    print(f"Carregando dados de rentabilidade de: {file_path}")
    try:
        df_long = pd.read_csv(
            file_path,
            sep=';',
            encoding='latin1',
            header=None,
            names=[
                'carteira', 'nomeFundo', 'tipoCarteira', 'administrador', 'periodo', 
                'emissao', 'data', 'valorCota', 'quantidade', 'numeroCotistas', 
                'variacaoDia', 'variacaoPeriodo', 'captacoes', 'resgate', 'eventos', 
                'pl', 'coluna_extra'
            ]
        )
        
        df_long.columns = df_long.columns.str.strip()
        
        # Limpeza e conversão de tipos
        for col in ['valorCota', 'pl', 'quantidade', 'numeroCotistas']:
            if col in df_long.columns:
                df_long[col] = df_long[col].astype(str).str.replace('.', '', regex=False).str.replace(',', '.', regex=False)
                df_long[col] = pd.to_numeric(df_long[col], errors='coerce')

        df_long['data'] = pd.to_datetime(df_long['data'], format='%d/%m/%Y', errors='coerce')
        df_long.dropna(subset=['data', 'valorCota'], inplace=True)
        
        # Pivotagem para formato 'wide'
        df_cota = df_long.pivot_table(index='data', columns='nomeFundo', values='valorCota').ffill().T
        df_patr = df_long.pivot_table(index='data', columns='nomeFundo', values='pl').ffill().T
        df_cota.index = df_cota.index.str.strip()
        df_patr.index = df_patr.index.str.strip()

        # Adicionar Benchmarks (Placeholder)
        print("Adicionando dados de benchmark de exemplo.")
        datas_base = pd.date_range(start=df_cota.columns.min(), end=df_cota.columns.max(), freq='B')
        cdi_diario = 0.12 / 252
        ibov_diario_ret = np.random.normal(0.0005, 0.015, len(datas_base))
        cdi_series = (1 + cdi_diario) ** np.arange(len(datas_base)) * 1000
        ibov_series = (1 + ibov_diario_ret).cumprod() * 100000
        df_benchmarks = pd.DataFrame({'CDI': cdi_series, 'Ibovespa': ibov_series}, index=datas_base)
        df_cota = pd.concat([df_cota, df_benchmarks.T])

        print("Dados de rentabilidade processados com sucesso.")
        return df_cota, df_patr

    except FileNotFoundError:
        print(f"ERRO: Arquivo de rentabilidade não encontrado em: {file_path}")
        return pd.DataFrame(), pd.DataFrame()
    except Exception as e:
        print(f"ERRO inesperado ao processar dados de rentabilidade: {e}")
        return pd.DataFrame(), pd.DataFrame()

# =============================================================================
# MÓDULO 2: PROCESSAMENTO DE DADOS DE ESTOQUE
# =============================================================================

def load_inventory_data(folder_path):
    """Lê e processa todos os arquivos de estoque de uma pasta."""
    print(f"Carregando dados de estoque de: {folder_path}")
    list_files = list(Path(folder_path).glob("*.csv"))
    if not list_files:
        print(f"AVISO: Nenhum arquivo CSV encontrado em {folder_path}. Pulando análise de estoque.")
        return pd.DataFrame()

    all_dfs = []
    for file in list_files:
        try:
            df = pd.read_csv(file, encoding='utf-16', sep='\t', engine='python', on_bad_lines='warn', header=0)
            df.columns = df.columns.str.strip()

            if not df.empty:
                for col in ['Valor Aquisicao', 'Valor Nominal', 'Valor Presente', 'PDD Vencido', 'PDD Total']:
                    if col in df.columns:
                        df[col] = df[col].astype(str).str.replace(',', '.').astype(float)
                for col in ['Data Aquisicao', 'Data Vencimento', 'Data Referencia']:
                    if col in df.columns:
                        df[col] = pd.to_datetime(df[col], errors='coerce', dayfirst=True)
                all_dfs.append(df)
        except Exception as e:
            print(f"ERRO ao processar o arquivo de estoque {file.name}: {e}")
            continue

    if not all_dfs:
        print("Nenhum dado de estoque foi carregado com sucesso.")
        return pd.DataFrame()

    df_final_estoque = pd.concat(all_dfs, ignore_index=True)
    print("Dados de estoque consolidados e processados com sucesso!")
    return df_final_estoque

# =============================================================================
# MÓDULO 3: GERAÇÃO DE VISUALIZAÇÕES (GRÁFICOS E TABELAS)
# =============================================================================

### Funções Auxiliares de Plotagem (Estoque) ###

def plot_vencimento_mensal(df_aberto, col_valor='Valor Presente', col_vcto='VCTO_MES', cutoff=10):
    df_aberto[col_vcto] = df_aberto['Data Vencimento'].dt.strftime('%Y-%m')
    df2 = df_aberto.groupby(col_vcto)[col_valor].sum().sort_index()[:cutoff]
    if df2.empty: return None
    
    fig, ax = plt.subplots(figsize=(8, 4.5))
    bars = ax.bar(df2.index, df2.values, color=COLORS['primary_light'])
    ax.set_frame_on(False)
    ax.tick_params(bottom=False, left=False)
    ax.set_yticks([])
    ax.set_title(f"Vencimento Mensal ({col_valor})", fontsize=14, pad=20)
    for bar in bars:
        height = bar.get_height()
        ax.text(bar.get_x() + bar.get_width() / 2, height, f'R$ {height/1e3:,.0f}k'.replace(',', '.'), 
                ha='center', va='bottom', fontsize=10)
    plt.tight_layout()
    return fig

def plot_vencimento_anual(df_aberto, col_valor='Valor Presente', col_vcto='VCTO_ANO'):
    df_aberto[col_vcto] = df_aberto['Data Vencimento'].dt.year
    df2 = df_aberto.groupby(col_vcto)[col_valor].sum()
    if df2.empty: return None

    fig, ax = plt.subplots(figsize=(8, 4.5))
    bars = ax.bar(df2.index.astype(str), df2.values, color=COLORS['primary_light'])
    ax.set_frame_on(False)
    ax.tick_params(bottom=False, left=False)
    ax.set_yticks([])
    ax.set_title(f"Vencimento Anual ({col_valor})", fontsize=14, pad=20)
    for bar in bars:
        height = bar.get_height()
        ax.text(bar.get_x() + bar.get_width() / 2, height, f'R$ {height/1e6:,.1f}M'.replace(',', '.'), 
                ha='center', va='bottom', fontsize=10)
    plt.tight_layout()
    return fig

# ... (outras funções de plotagem de estoque podem ser adicionadas aqui) ...

### Funções Principais de Geração de Gráficos ###

def render_mpl_table(data):
    """Renderiza um DataFrame do pandas como uma imagem matplotlib."""
    fig, ax = plt.subplots(figsize=(8, 5))
    ax.axis('off')
    mpl_table = ax.table(
        cellText=data.values,
        bbox=[0, 0, 1, 1],
        colLabels=data.columns,
        rowLabels=[f"{idx[0]} {idx[1]}" for idx in data.index],
        loc='center'
    )
    mpl_table.auto_set_font_size(False)
    mpl_table.set_fontsize(10)
    for k, cell in mpl_table._cells.items():
        cell.set_edgecolor('w')
        if k[0] == 0: # Header
            cell.set_text_props(weight='bold', color=COLORS['table_header_text'])
            cell.set_facecolor(COLORS['table_header_bg'])
        if k[1] == -1: # Row labels
            cell.set_text_props(weight='bold', color=COLORS['table_header_text'])
            cell.get_text().set_ha('right')
    return fig

def generate_performance_visuals(df_cota, df_patr, fund_name):
    """Gera todos os gráficos e tabelas relacionados à performance."""
    if df_cota.empty or fund_name not in df_cota.index:
        print(f"ERRO: Fundo '{fund_name}' não encontrado nos dados de performance.")
        return {}

    visuals = {}
    comparables = [fund_name, 'CDI'] # Simplificado para o exemplo

    # 1. Tabela de Rentabilidade
    try:
        df_cota_fundo = df_cota.loc[fund_name].replace(0, np.nan).dropna()
        start_date = df_cota_fundo.index[0]
        years = sorted(list(set(df_cota_fundo.index.year)))
        rows = []
        for year in years:
            for comp in comparables:
                df_aux = df_cota.loc[comp].ffill()[start_date:]
                ret_mensal = df_aux.resample('M').last().pct_change().fillna(0)
                ytd_val = (df_aux[df_aux.index.year == year].iloc[-1] / df_aux[df_aux.index.year == year].iloc[0]) - 1
                row = ret_mensal[ret_mensal.index.year == year].to_frame().T
                row.index = pd.MultiIndex.from_tuples([(year, comp)])
                row.columns = row.columns.strftime('%b')
                row['Ano'] = ytd_val
                rows.append(row)
        df_return = pd.concat(rows).fillna(0)
        df_return_formatted = df_return.applymap(lambda x: f'{x:.2%}'.replace('.', ','))
        visuals['rentabilidade'] = render_mpl_table(df_return_formatted)
    except Exception as e:
        print(f"Erro ao gerar tabela de rentabilidade: {e}")


    # 2. Gráfico de Retorno Acumulado
    try:
        fig, ax = plt.subplots(figsize=(10, 5))
        for comp in comparables:
            series = df_cota.loc[comp].ffill()[start_date:]
            series_norm = series / series.iloc[0] * 100
            ax.plot(series_norm, label=comp, lw=2.5 if comp == fund_name else 2, 
                    color=COLORS['primary_cyan'] if comp == fund_name else 'k',
                    linestyle='-' if comp == fund_name else '--')
        ax.set_title(f'Retorno Acumulado - {fund_name}', fontsize=16)
        ax.legend(frameon=False)
        ax.spines[['top', 'right']].set_visible(False)
        visuals['retorno_acumulado'] = fig
    except Exception as e:
        print(f"Erro ao gerar gráfico de retorno: {e}")

    # 3. Gráfico de Volatilidade
    try:
        fig, ax = plt.subplots(figsize=(10, 5))
        vol = df_cota.loc[fund_name].pct_change().rolling(22).std() * np.sqrt(252)
        ax.plot(vol, label='Vol Fundo', color=COLORS['primary_cyan'])
        ax.yaxis.set_major_formatter(ticker.PercentFormatter(1.0))
        ax.set_title(f'Volatilidade Anualizada (Janela 22d) - {fund_name}', fontsize=16)
        ax.spines[['top', 'right']].set_visible(False)
        visuals['volatilidade'] = fig
    except Exception as e:
        print(f"Erro ao gerar gráfico de volatilidade: {e}")


    # 4. Gráfico de Drawdown
    try:
        fig, ax = plt.subplots(figsize=(10, 5))
        df_dd = df_cota.loc[fund_name].dropna()
        roll_max = df_dd.cummax()
        daily_dd = df_dd / roll_max - 1.0
        ax.fill_between(daily_dd.index, daily_dd, 0, color=COLORS['primary_cyan'], alpha=0.3)
        ax.plot(daily_dd, color=COLORS['primary_cyan'], lw=1)
        ax.yaxis.set_major_formatter(ticker.PercentFormatter(1.0))
        ax.set_title(f'Drawdown Histórico - {fund_name}', fontsize=16)
        ax.spines[['top', 'right']].set_visible(False)
        visuals['drawdown'] = fig
    except Exception as e:
        print(f"Erro ao gerar gráfico de drawdown: {e}")


    # 5. Gráfico de Evolução do PL
    if not df_patr.empty and fund_name in df_patr.index:
        try:
            fig, ax = plt.subplots(figsize=(10, 5))
            pl_fundo = df_patr.loc[fund_name].dropna().loc[lambda x: x > 0]
            ax.fill_between(pl_fundo.index, pl_fundo, 0, color=COLORS['primary_dark'], alpha=0.3)
            ax.plot(pl_fundo, color=COLORS['primary_dark'], lw=2)
            ax.yaxis.set_major_formatter(ticker.FuncFormatter(lambda x, p: f'R$ {x/1e6:.1f}M'))
            ax.set_title(f'Evolução do Patrimônio Líquido (PL) - {fund_name}', fontsize=16)
            ax.spines[['top', 'right']].set_visible(False)
            visuals['pl'] = fig
        except Exception as e:
            print(f"Erro ao gerar gráfico de PL: {e}")

    plt.close('all') # Fecha todas as figuras para não exibi-las no console
    return visuals

def generate_inventory_visuals(df_estoque):
    """Gera todos os gráficos relacionados ao estoque."""
    if df_estoque.empty:
        return {}
        
    visuals = {}
    df_avencer = df_estoque[df_estoque['Status'] == 'A vencer'].copy()

    # Mapeamento dos gráficos a serem gerados
    plot_map = {
        'vencimento_mensal': (plot_vencimento_mensal, {'df_aberto': df_avencer}),
        'vencimento_anual': (plot_vencimento_anual, {'df_aberto': df_avencer}),
        # Adicione outras funções de plotagem aqui
    }

    for name, (plot_func, kwargs) in plot_map.items():
        try:
            fig = plot_func(**kwargs)
            if fig:
                visuals[name] = fig
        except Exception as e:
            print(f"Erro ao gerar gráfico de estoque '{name}': {e}")
            
    plt.close('all')
    return visuals

# =============================================================================
# MÓDULO 4: MONTAGEM E GERAÇÃO DO RELATÓRIO
# =============================================================================

def add_image_from_fig(cell, fig, width):
    """Salva uma figura matplotlib em um buffer de memória e a insere em uma célula do Word."""
    if fig is None: return
    mem_image = io.BytesIO()
    fig.savefig(mem_image, format='png', dpi=300, bbox_inches='tight')
    mem_image.seek(0)
    cell.paragraphs[0].add_run().add_picture(mem_image, width=width)
    
def add_section_title(cell_or_doc, text):
    """Adiciona um parágrafo de título padronizado."""
    p = cell_or_doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Gill Sans MT"
    run.font.size = Pt(11)
    run.bold = True

def create_report(config, visuals):
    """Cria o documento DOCX e o preenche com os visuais."""
    print("Iniciando a criação do relatório...")
    doc = Document(config.TEMPLATE_PATH)
    
    # --- Configuração de Cabeçalho ---
    section = doc.sections[0]
    header = section.header
    header_table = header.tables[0] if header.tables else header.add_table(rows=1, cols=2, width=Cm(18))
    
    left_cell = header_table.cell(0, 0)
    left_cell.text = ""
    p_left = left_cell.paragraphs[0]
    run_name = p_left.add_run("LÂMINA DE PERFORMANCE E ESTOQUE")
    run_name.font.name = "Gill Sans MT"
    run_name.font.size = Pt(14)
    run_name.bold = True
    p_left.add_run("\n")
    meses_pt = ["Janeiro", "Fevereiro", "Março", "Abril", "Maio", "Junho", "Julho", "Agosto", "Setembro", "Outubro", "Novembro", "Dezembro"]
    data_formatada = f"{meses_pt[config.REFERENCE_DATE.month - 1]} {config.REFERENCE_DATE.year}"
    p_left.add_run(data_formatada).italic = True
    
    right_cell = header_table.cell(0, 1)
    right_cell.text = ""
    p_right = right_cell.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_right.add_run().add_picture(str(config.LOGO_PATH), width=Cm(4.93))

    # --- Corpo do Relatório ---
    add_section_title(doc, "Rentabilidade Mensal")
    add_image_from_fig(doc, visuals.get('rentabilidade'), width=Cm(17.8))

    # Tabela para organizar os gráficos
    graphs_table = doc.add_table(rows=3, cols=2)
    graphs_table.autofit = False
    graphs_table.allow_autofit = False
    col_width = Cm(8.9)
    img_width = Cm(8.5)
    for col in graphs_table.columns:
        col.width = col_width

    # Mapeamento de visuais para células
    cell_map = {
        (0, 0): ('Retorno Acumulado', 'retorno_acumulado'),
        (0, 1): ('Evolução do PL', 'pl'),
        (1, 0): ('Volatilidade Anualizada', 'volatilidade'),
        (1, 1): ('Drawdown Histórico', 'drawdown'),
        (2, 0): ('Vencimento Mensal do Estoque', 'vencimento_mensal'),
        (2, 1): ('Vencimento Anual do Estoque', 'vencimento_anual'),
    }

    for (row, col), (title, key) in cell_map.items():
        cell = graphs_table.cell(row, col)
        add_section_title(cell, title)
        add_image_from_fig(cell, visuals.get(key), width=img_width)
    
    # --- Salvamento ---
    try:
        docx_path = config.OUTPUT_DIR / f"{config.OUTPUT_FILENAME_BASE}.docx"
        pdf_path = config.OUTPUT_DIR / f"{config.OUTPUT_FILENAME_BASE}.pdf"
        
        doc.save(docx_path)
        print(f"Relatório DOCX gerado com sucesso em: {docx_path}")
        
        print("Iniciando conversão para PDF...")
        convert(str(docx_path), str(pdf_path))
        print(f"Relatório PDF gerado com sucesso em: {pdf_path}")
        
    except Exception as e:
        print(f"ERRO CRÍTICO ao salvar o relatório: {e}")

# =============================================================================
# CÉLULA 5: ORQUESTRADOR PRINCIPAL
# =============================================================================

def main():
    """
    Executa o pipeline completo de ponta a ponta.
    """
    print("="*50)
    print("INICIANDO PIPELINE DE GERAÇÃO DE RELATÓRIO")
    print("="*50)
    
    # 1. Carregar configurações
    config = Config()
    
    # 2. Processar dados de performance e gerar visuais
    df_cota, df_patr = load_performance_data(config.PERFORMANCE_DATA_PATH)
    performance_visuals = generate_performance_visuals(df_cota, df_patr, config.FUND_NAME)
    
    # 3. Processar dados de estoque e gerar visuais
    df_estoque = load_inventory_data(config.INVENTORY_DATA_DIR)
    inventory_visuals = generate_inventory_visuals(df_estoque)
    
    # 4. Combinar todos os visuais
    all_visuals = {**performance_visuals, **inventory_visuals}
    
    # 5. Gerar o relatório final
    create_report(config, all_visuals)
    
    print("="*50)
    print("PIPELINE FINALIZADO COM SUCESSO!")
    print("="*50)


if __name__ == "__main__":
    main()